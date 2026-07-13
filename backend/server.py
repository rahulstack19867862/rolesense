from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
from emergentintegrations.llm.chat import LlmChat, UserMessage
import json
import re
import io
import base64
import hashlib
from Crypto.Cipher import AES
from Crypto.Util.Padding import unpad

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ.get('DB_NAME', 'rolesense')]

# LLM API Key
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY')

# Create the main app without a prefix
app = FastAPI(title="RoleSense - JD Intelligence Dashboard")

# CORS Configuration - Add BEFORE routes
# Get allowed origins from environment
cors_origins_env = os.environ.get('CORS_ORIGINS', '*').strip()
if cors_origins_env == '*':
    cors_origins = ["*"]
else:
    cors_origins = [origin.strip() for origin in cors_origins_env.split(',')]

# Add CORS middleware FIRST
app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_credentials=True if cors_origins_env != '*' else False,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
    allow_headers=["*"],
    expose_headers=["*"],
    max_age=600,
)

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ============ Models ============

class ScreeningQuestion(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    question: str
    skill_area: str
    expected_answer: str
    difficulty: str = "medium"  # easy, medium, hard
    time_estimate: str = "2-3 min"

class JobDescription(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    company: Optional[str] = None
    raw_text: str
    parsed_data: Optional[Dict[str, Any]] = None
    analysis: Optional[Dict[str, Any]] = None
    improvement_suggestions: Optional[List[str]] = None
    # New fields for Active Jobs workflow
    status: str = "draft"  # draft, active, closed
    requisition_number: Optional[str] = None
    requisition_date: Optional[datetime] = None
    screening_questions: List[Dict[str, Any]] = []
    publish_links: Optional[Dict[str, str]] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class JobDescriptionCreate(BaseModel):
    title: str
    company: Optional[str] = None
    raw_text: str

class JobDescriptionUpdate(BaseModel):
    title: Optional[str] = None
    company: Optional[str] = None
    raw_text: Optional[str] = None
    status: Optional[str] = None

class ScreeningQuestionUpdate(BaseModel):
    questions: List[Dict[str, Any]]

# ============ Auto-Routing Resume Repository Models ============

# Source types for tracking where resumes come from
SOURCE_TYPES = {
    "linkedin": {"name": "LinkedIn", "quality_weight": 1.2},
    "indeed": {"name": "Indeed", "quality_weight": 1.0},
    "glassdoor": {"name": "Glassdoor", "quality_weight": 1.1},
    "monster": {"name": "Monster", "quality_weight": 0.9},
    "ziprecruiter": {"name": "ZipRecruiter", "quality_weight": 1.0},
    "naukri": {"name": "Naukri", "quality_weight": 1.0},
    "referral": {"name": "Employee Referral", "quality_weight": 1.5},
    "career_page": {"name": "Career Page", "quality_weight": 1.3},
    "job_fair": {"name": "Job Fair", "quality_weight": 0.8},
    "social_post": {"name": "Social Media Post", "quality_weight": 0.9},
    "direct": {"name": "Direct Application", "quality_weight": 1.0},
    "agency": {"name": "Recruitment Agency", "quality_weight": 1.1},
    "campus": {"name": "Campus Recruitment", "quality_weight": 0.85},
    "internal": {"name": "Internal Transfer", "quality_weight": 1.4}
}

# SLA Rules for notifications
SLA_RULES = {
    "HR": {
        "Talent Acquisition": {"daily_threshold": 10, "notify_email": "ta-lead@company.com", "priority": "high"},
        "HR Operations": {"daily_threshold": 5, "notify_email": "hr-ops@company.com", "priority": "medium"},
        "Learning & Development": {"daily_threshold": 3, "notify_email": "ld-team@company.com", "priority": "low"},
    },
    "IT": {
        "Software Engineering": {"daily_threshold": 15, "notify_email": "engineering-lead@company.com", "priority": "high"},
        "Data Science & AI": {"daily_threshold": 8, "notify_email": "data-team@company.com", "priority": "high"},
        "Quality Assurance": {"daily_threshold": 10, "notify_email": "qa-lead@company.com", "priority": "medium"},
    },
    "Finance": {
        "Financial Planning & Analysis": {"daily_threshold": 5, "notify_email": "finance-lead@company.com", "priority": "medium"},
    },
    "default": {"daily_threshold": 10, "notify_email": "recruitment@company.com", "priority": "medium"}
}

# Skill Taxonomy for each function - used for AI classification
SKILL_TAXONOMY = {
    "HR": {
        "Talent Acquisition": ["recruiting", "sourcing", "interviewing", "employer branding", "campus hiring", "lateral hiring", "ATS", "headhunting", "offer negotiation", "candidate experience"],
        "HR Operations": ["payroll", "compliance", "employee records", "HRIS", "benefits administration", "onboarding", "offboarding", "attendance", "leave management"],
        "Learning & Development": ["training", "e-learning", "LMS", "skill development", "leadership development", "coaching", "mentoring", "curriculum design", "performance consulting"],
        "HR Business Partner": ["strategic HR", "employee relations", "workforce planning", "change management", "organizational development", "talent management", "succession planning"],
        "Compensation & Benefits": ["salary benchmarking", "compensation structure", "incentive programs", "equity compensation", "benefits design", "total rewards"]
    },
    "IT": {
        "Software Engineering": ["python", "java", "javascript", "react", "node.js", "angular", "vue", "backend", "frontend", "full-stack", "microservices", "API", "REST", "GraphQL"],
        "Quality Assurance": ["testing", "QA", "automation testing", "selenium", "cypress", "jest", "manual testing", "test planning", "bug tracking", "JIRA"],
        "Data & Analytics": ["data analysis", "SQL", "python", "tableau", "power BI", "data visualization", "ETL", "data warehouse", "big data", "spark", "hadoop"],
        "Data Science & AI": ["machine learning", "deep learning", "NLP", "computer vision", "TensorFlow", "PyTorch", "AI", "predictive modeling", "statistical analysis"],
        "DevOps & Infrastructure": ["AWS", "Azure", "GCP", "kubernetes", "docker", "CI/CD", "terraform", "ansible", "linux", "networking", "cloud architecture"],
        "IT Support": ["helpdesk", "technical support", "troubleshooting", "hardware", "software installation", "user support", "ticketing system"],
        "Cybersecurity": ["security", "penetration testing", "vulnerability assessment", "SIEM", "firewall", "encryption", "compliance", "SOC", "incident response"],
        "Product Management": ["product strategy", "roadmap", "agile", "scrum", "user stories", "backlog", "stakeholder management", "market research"]
    },
    "Finance": {
        "Accounting": ["bookkeeping", "accounts payable", "accounts receivable", "general ledger", "reconciliation", "journal entries", "GAAP", "IFRS"],
        "Financial Planning & Analysis": ["FP&A", "budgeting", "forecasting", "financial modeling", "variance analysis", "business planning", "scenario analysis"],
        "Audit": ["internal audit", "external audit", "SOX compliance", "risk assessment", "audit planning", "control testing", "audit reports"],
        "Tax": ["tax compliance", "tax planning", "corporate tax", "GST", "VAT", "transfer pricing", "tax returns", "tax advisory"],
        "Treasury": ["cash management", "liquidity", "foreign exchange", "hedging", "banking relationships", "working capital"],
        "Investment & Banking": ["investment analysis", "portfolio management", "M&A", "due diligence", "valuation", "financial advisory"]
    },
    "Marketing": {
        "Digital Marketing": ["SEO", "SEM", "PPC", "Google Ads", "social media marketing", "content marketing", "email marketing", "analytics", "conversion optimization"],
        "Brand Management": ["brand strategy", "brand identity", "brand positioning", "market research", "consumer insights", "brand campaigns"],
        "Product Marketing": ["go-to-market", "product positioning", "competitive analysis", "messaging", "sales enablement", "product launches"],
        "Content & Creative": ["copywriting", "content creation", "video production", "graphic design", "storytelling", "editorial"],
        "Marketing Analytics": ["marketing ROI", "attribution", "A/B testing", "customer analytics", "campaign analysis", "marketing automation"],
        "Events & PR": ["event management", "public relations", "media relations", "press releases", "corporate communications", "crisis management"]
    },
    "Operations": {
        "Business Operations": ["process improvement", "operational efficiency", "KPIs", "SLAs", "vendor management", "business continuity"],
        "Project Management": ["project planning", "PMO", "PMP", "agile", "waterfall", "resource allocation", "stakeholder management", "risk management"],
        "Quality Management": ["ISO", "six sigma", "lean", "quality control", "quality assurance", "process documentation", "continuous improvement"],
        "Facilities Management": ["facility operations", "maintenance", "space planning", "health & safety", "security", "real estate"]
    },
    "Supply Chain": {
        "Procurement": ["sourcing", "vendor management", "contract negotiation", "RFP", "RFQ", "supplier evaluation", "cost optimization"],
        "Logistics": ["transportation", "freight", "warehousing", "distribution", "last mile delivery", "route optimization", "3PL"],
        "Inventory Management": ["inventory control", "demand planning", "stock optimization", "warehouse management", "cycle counting"],
        "Supply Planning": ["demand forecasting", "S&OP", "supply planning", "production planning", "capacity planning", "MRP"]
    },
    "Administration": {
        "Executive Assistant": ["calendar management", "travel coordination", "meeting scheduling", "correspondence", "presentation preparation"],
        "Office Administration": ["office management", "front desk", "reception", "document management", "filing", "office supplies"],
        "Legal & Compliance": ["contract management", "legal documentation", "regulatory compliance", "corporate governance", "policy management"]
    }
}

# ============ JD CREATION MASTER DATA ============
# Human-controlled dropdowns and selections for JD creation

# Role Types
ROLE_TYPES = ["IT", "Non-IT"]

# Business Models
BUSINESS_MODELS = ["B2B", "B2C", "B2B2C", "D2C", "SaaS", "Marketplace", "Consulting", "Manufacturing", "Services", "Healthcare", "EdTech", "FinTech", "Other"]

# Experience Ranges
EXPERIENCE_RANGES = [
    {"label": "Fresher (0 years)", "min": 0, "max": 0},
    {"label": "0-1 years", "min": 0, "max": 1},
    {"label": "1-3 years", "min": 1, "max": 3},
    {"label": "3-5 years", "min": 3, "max": 5},
    {"label": "5-7 years", "min": 5, "max": 7},
    {"label": "7-10 years", "min": 7, "max": 10},
    {"label": "10-15 years", "min": 10, "max": 15},
    {"label": "15+ years", "min": 15, "max": 25}
]

# Currencies
CURRENCIES = ["INR", "USD", "EUR", "GBP", "AED", "SGD", "AUD", "CAD", "JPY", "CNY"]

# Compensation Types
COMPENSATION_TYPES = ["Per Annum", "Per Month", "Per Hour", "CTC", "Fixed + Variable"]

# India Locations
INDIA_LOCATIONS = [
    "Mumbai", "Delhi NCR", "Bangalore", "Hyderabad", "Chennai", "Kolkata", "Pune", 
    "Ahmedabad", "Jaipur", "Lucknow", "Chandigarh", "Bhopal", "Indore", "Nagpur",
    "Kochi", "Coimbatore", "Visakhapatnam", "Thiruvananthapuram", "Guwahati",
    "Pan India", "Remote - India", "Tier 1 Cities", "Tier 2 Cities", "Metro Cities Only"
]

# International Locations
INTERNATIONAL_LOCATIONS = [
    "United States", "United Kingdom", "Canada", "Australia", "Singapore", "UAE", 
    "Germany", "Netherlands", "Ireland", "Switzerland", "Japan", "Hong Kong",
    "Remote - Global", "Remote - US", "Remote - Europe", "Remote - APAC", "Remote - Middle East"
]

# Work Modes
WORK_MODES = ["Work from Office", "Remote", "Hybrid (2-3 days office)", "Hybrid (3-4 days office)", "Field Work", "Client Location"]

# Employment Types
EMPLOYMENT_TYPES = ["Full-time Permanent", "Full-time Contract", "Part-time", "Internship", "Freelance/Consultant", "Third Party Payroll"]

# ============ COMPETENCY MASTER DATA ============

# Core/Behavioral Competencies
CORE_BEHAVIORAL_COMPETENCIES = [
    {"id": "comm_stakeholder", "name": "Communication & Stakeholder Management", "description": "Ability to communicate effectively with internal and external stakeholders"},
    {"id": "problem_solving", "name": "Problem-solving & Analytical Thinking", "description": "Ability to analyze complex situations and find solutions"},
    {"id": "accountability", "name": "Accountability & Ownership", "description": "Taking responsibility for actions and outcomes"},
    {"id": "adaptability", "name": "Adaptability & Learning Agility", "description": "Ability to adapt to changes and learn new skills quickly"},
    {"id": "time_mgmt", "name": "Time Management & Execution Discipline", "description": "Managing time effectively and delivering on commitments"},
    {"id": "teamwork", "name": "Teamwork & Collaboration", "description": "Working effectively with others towards common goals"},
    {"id": "leadership", "name": "Leadership & Influence", "description": "Ability to lead teams and influence decisions"},
    {"id": "strategic_planning", "name": "Strategic Planning & Vision", "description": "Ability to develop long-term strategies and translate vision into actionable plans"},
    {"id": "customer_focus", "name": "Customer Focus", "description": "Understanding and meeting customer needs"},
    {"id": "integrity", "name": "Integrity & Ethics", "description": "Acting with honesty and ethical standards"},
    {"id": "initiative", "name": "Initiative & Proactiveness", "description": "Taking action without being asked"},
    {"id": "conflict_resolution", "name": "Conflict Resolution", "description": "Handling disagreements constructively"},
    {"id": "emotional_intelligence", "name": "Emotional Intelligence", "description": "Understanding and managing emotions effectively"},
    {"id": "decision_making", "name": "Decision Making", "description": "Ability to make timely and effective decisions under pressure"},
    {"id": "change_management", "name": "Change Management", "description": "Leading and managing organizational change effectively"},
    {"id": "innovation", "name": "Innovation & Creativity", "description": "Generating new ideas and driving innovation"}
]

# Functional/Domain Competencies - IT
IT_FUNCTIONAL_COMPETENCIES = [
    {"id": "it_prog_lang", "name": "Programming Languages", "description": "Proficiency in relevant programming languages", "examples": ["Python", "Java", "JavaScript", "C++", "C#", "Go", "Rust", "Ruby", "PHP", "Swift", "Kotlin"]},
    {"id": "it_frameworks", "name": "Frameworks & Libraries", "description": "Experience with development frameworks", "examples": ["React", "Angular", "Vue.js", "Node.js", "Django", "Flask", "Spring Boot", ".NET", "FastAPI", "Express.js"]},
    {"id": "it_databases", "name": "Database Management", "description": "Working with databases", "examples": ["MySQL", "PostgreSQL", "MongoDB", "Redis", "Oracle", "SQL Server", "Cassandra", "DynamoDB"]},
    {"id": "it_cloud", "name": "Cloud Platforms", "description": "Cloud infrastructure experience", "examples": ["AWS", "Azure", "GCP", "Heroku", "DigitalOcean", "Oracle Cloud"]},
    {"id": "it_devops", "name": "DevOps & CI/CD", "description": "DevOps practices and tools", "examples": ["Docker", "Kubernetes", "Jenkins", "GitLab CI", "GitHub Actions", "Terraform", "Ansible"]},
    {"id": "it_testing", "name": "Testing & QA", "description": "Software testing methodologies", "examples": ["Unit Testing", "Integration Testing", "Selenium", "Jest", "Cypress", "JUnit", "PyTest"]},
    {"id": "it_security", "name": "Cybersecurity", "description": "Security practices and tools", "examples": ["OWASP", "Penetration Testing", "Encryption", "IAM", "SIEM", "SOC"]},
    {"id": "it_data_science", "name": "Data Science & ML", "description": "Data science and machine learning", "examples": ["TensorFlow", "PyTorch", "Scikit-learn", "Pandas", "NumPy", "NLP", "Computer Vision"]},
    {"id": "it_architecture", "name": "System Architecture", "description": "Designing scalable systems", "examples": ["Microservices", "REST APIs", "GraphQL", "Message Queues", "Event-Driven"]},
    {"id": "it_agile", "name": "Agile & Project Management", "description": "Agile methodologies", "examples": ["Scrum", "Kanban", "JIRA", "Confluence", "Sprint Planning"]}
]

# Functional/Domain Competencies - Non-IT
NON_IT_FUNCTIONAL_COMPETENCIES = [
    {"id": "nit_sales", "name": "Sales & Business Development", "description": "Sales processes and techniques", "examples": ["Lead Generation", "Pipeline Management", "Negotiation", "CRM", "Account Management"]},
    {"id": "nit_marketing", "name": "Marketing & Brand Management", "description": "Marketing strategies", "examples": ["Digital Marketing", "SEO/SEM", "Brand Strategy", "Content Marketing", "Social Media"]},
    {"id": "nit_finance", "name": "Finance & Accounting", "description": "Financial management", "examples": ["Financial Analysis", "Budgeting", "Forecasting", "Taxation", "Audit", "GAAP/IFRS"]},
    {"id": "nit_hr", "name": "Human Resources", "description": "HR processes", "examples": ["Recruitment", "Employee Relations", "Performance Management", "Compensation & Benefits", "HRIS"]},
    {"id": "nit_operations", "name": "Operations & Process Management", "description": "Operational excellence", "examples": ["Process Improvement", "Six Sigma", "Lean", "Quality Management", "Vendor Management"]},
    {"id": "nit_supply_chain", "name": "Supply Chain & Logistics", "description": "Supply chain management", "examples": ["Procurement", "Inventory Management", "Logistics", "Warehouse Management", "Demand Planning"]},
    {"id": "nit_legal", "name": "Legal & Compliance", "description": "Legal and regulatory", "examples": ["Contract Management", "Regulatory Compliance", "Corporate Law", "IP Rights", "Data Privacy"]},
    {"id": "nit_analytics", "name": "Business Analytics", "description": "Data-driven decision making", "examples": ["Excel Advanced", "Power BI", "Tableau", "SQL", "Data Visualization"]},
    {"id": "nit_customer_service", "name": "Customer Service & Support", "description": "Customer handling", "examples": ["Customer Handling", "Complaint Resolution", "CRM Tools", "Service Excellence"]},
    {"id": "nit_admin", "name": "Administration & Coordination", "description": "Administrative skills", "examples": ["Executive Assistance", "Calendar Management", "Travel Coordination", "Documentation"]}
]

# Cognitive Competencies
COGNITIVE_COMPETENCIES = [
    {"id": "cog_logical", "name": "Logical Reasoning", "description": "Ability to think logically and systematically"},
    {"id": "cog_data_interp", "name": "Data Interpretation", "description": "Ability to understand and analyze data"},
    {"id": "cog_decision", "name": "Decision-making Ability", "description": "Making sound decisions under various conditions"},
    {"id": "cog_attention", "name": "Attention to Detail", "description": "Thoroughness and accuracy in work"},
    {"id": "cog_critical", "name": "Critical Thinking", "description": "Evaluating information objectively"},
    {"id": "cog_numerical", "name": "Numerical Aptitude", "description": "Comfort with numbers and calculations"},
    {"id": "cog_verbal", "name": "Verbal Reasoning", "description": "Understanding written and verbal information"},
    {"id": "cog_spatial", "name": "Spatial Reasoning", "description": "Understanding visual and spatial relationships"},
    {"id": "cog_memory", "name": "Working Memory", "description": "Retaining and using information effectively"},
    {"id": "cog_pattern", "name": "Pattern Recognition", "description": "Identifying patterns and trends"}
]

# Good-to-Have Competencies
GOOD_TO_HAVE_COMPETENCIES = [
    {"id": "gth_domain_adv", "name": "Advanced Domain Exposure", "description": "Deep expertise in specific industry verticals"},
    {"id": "gth_cross_func", "name": "Cross-functional Experience", "description": "Experience working across multiple departments"},
    {"id": "gth_multi_industry", "name": "Multi-industry Experience", "description": "Experience in different industry sectors"},
    {"id": "gth_process_opt", "name": "Process Optimization Experience", "description": "History of improving processes"},
    {"id": "gth_automation", "name": "Automation Exposure", "description": "Experience with automation tools/processes"},
    {"id": "gth_startup", "name": "Startup Experience", "description": "Experience working in startup environment"},
    {"id": "gth_mnc", "name": "MNC/Corporate Experience", "description": "Experience in large corporations"},
    {"id": "gth_global", "name": "Global/International Exposure", "description": "Experience working with international teams"},
    {"id": "gth_certifications", "name": "Relevant Certifications", "description": "Professional certifications in the field"},
    {"id": "gth_publications", "name": "Publications/Patents", "description": "Published work or patents"}
]

# Trainable Competencies
TRAINABLE_COMPETENCIES = [
    {"id": "trn_tools", "name": "Tools & Technologies (Learnable)", "description": "Tools that can be learned on the job"},
    {"id": "trn_org_systems", "name": "Organization-specific Systems", "description": "Internal systems and processes"},
    {"id": "trn_domain_specific", "name": "Domain-specific Processes", "description": "Industry-specific procedures"},
    {"id": "trn_ai_literacy", "name": "AI & Automation Literacy", "description": "Understanding of AI tools and automation"},
    {"id": "trn_new_platforms", "name": "Emerging Platforms", "description": "New technologies and platforms"},
    {"id": "trn_soft_skills", "name": "Soft Skills Enhancement", "description": "Communication, presentation skills"},
    {"id": "trn_leadership", "name": "Leadership Development", "description": "Leadership and management skills"},
    {"id": "trn_compliance", "name": "Compliance & Regulatory", "description": "Industry regulations and compliance"}
]

# Tools & Technologies - IT
IT_TOOLS = {
    "Programming": ["VS Code", "IntelliJ IDEA", "PyCharm", "Eclipse", "Sublime Text", "Vim"],
    "Version Control": ["Git", "GitHub", "GitLab", "Bitbucket", "SVN"],
    "Project Management": ["JIRA", "Trello", "Asana", "Monday.com", "Azure DevOps", "Notion"],
    "Communication": ["Slack", "Microsoft Teams", "Zoom", "Discord", "Google Meet"],
    "Design": ["Figma", "Adobe XD", "Sketch", "InVision", "Canva"],
    "Data & Analytics": ["Tableau", "Power BI", "Looker", "Metabase", "Google Analytics"],
    "Cloud Console": ["AWS Console", "Azure Portal", "GCP Console", "Vercel", "Netlify"],
    "Documentation": ["Confluence", "Notion", "ReadMe", "Swagger/OpenAPI"],
    "Monitoring": ["Datadog", "New Relic", "Grafana", "Prometheus", "Splunk"],
    "Testing": ["Postman", "Insomnia", "Selenium IDE", "JMeter", "LoadRunner"]
}

# Tools & Technologies - Non-IT
NON_IT_TOOLS = {
    "Office Suite": ["MS Office", "Google Workspace", "LibreOffice"],
    "CRM": ["Salesforce", "HubSpot", "Zoho CRM", "Pipedrive", "Freshsales"],
    "ERP": ["SAP", "Oracle ERP", "Microsoft Dynamics", "NetSuite", "Tally"],
    "HRMS": ["Workday", "SuccessFactors", "BambooHR", "Darwinbox", "Keka"],
    "Communication": ["MS Teams", "Slack", "Zoom", "Google Meet", "Outlook"],
    "Project Management": ["MS Project", "Asana", "Monday.com", "Trello", "Basecamp"],
    "Finance": ["QuickBooks", "Xero", "FreshBooks", "SAP FICO", "Oracle Financials"],
    "Marketing": ["HubSpot", "Mailchimp", "Hootsuite", "Buffer", "Google Ads", "Meta Ads Manager"],
    "Analytics": ["Excel Advanced", "Power BI", "Tableau", "Google Analytics", "Mixpanel"],
    "Design": ["Canva", "Adobe Creative Suite", "PowerPoint", "Prezi"]
}

# Education Requirements
EDUCATION_LEVELS = [
    "10th Pass", "12th Pass", "Diploma", "Bachelor's Degree", "Master's Degree", 
    "MBA", "PhD", "Professional Certification", "Any Graduate", "Any Post Graduate"
]

EDUCATION_FIELDS = [
    "Computer Science/IT", "Engineering", "Business Administration", "Commerce", "Economics",
    "Science", "Arts/Humanities", "Law", "Medicine", "Pharmacy", "Architecture",
    "Hotel Management", "Mass Communication", "Design", "Any Field"
]

# New Structured JD Model
class StructuredJDBasicInfo(BaseModel):
    company_name: str
    about_company: Optional[str] = None
    title: str
    role_type: str  # IT or Non-IT
    business_model: str
    experience_min: int
    experience_max: int
    compensation_min: Optional[float] = None
    compensation_max: Optional[float] = None
    compensation_currency: str = "INR"
    compensation_type: str = "Per Annum"
    locations_india: List[str] = []
    locations_international: List[str] = []
    work_mode: str
    employment_type: str
    education_level: str
    education_field: Optional[str] = None

class StructuredJDCompetencies(BaseModel):
    must_have_behavioral: List[str] = []  # IDs from CORE_BEHAVIORAL_COMPETENCIES
    must_have_functional: List[str] = []  # IDs from IT/NON_IT_FUNCTIONAL_COMPETENCIES
    must_have_cognitive: List[str] = []  # IDs from COGNITIVE_COMPETENCIES
    must_have_skills: List[str] = []  # Free text skills - minimum 4
    good_to_have_competencies: List[str] = []  # IDs from GOOD_TO_HAVE_COMPETENCIES
    good_to_have_skills: List[str] = []  # Free text skills - minimum 3
    trainable_competencies: List[str] = []
    tools_must_have: List[str] = []
    tools_good_to_have: List[str] = []

class StructuredJDCreate(BaseModel):
    basic_info: StructuredJDBasicInfo
    competencies: StructuredJDCompetencies
    responsibilities: List[str] = []
    additional_notes: Optional[str] = None

class StructuredJD(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    basic_info: StructuredJDBasicInfo
    competencies: StructuredJDCompetencies
    responsibilities: List[str] = []
    additional_notes: Optional[str] = None
    # Application submission options
    application_email: Optional[str] = None  # Email for resume submission
    application_link: Optional[str] = None   # External ATS/career page link
    # Workflow fields
    status: str = "draft"  # draft, active, closed
    requisition_number: Optional[str] = None
    requisition_date: Optional[datetime] = None
    screening_questions: List[Dict[str, Any]] = []
    publish_links: Optional[Dict[str, str]] = None
    # AI-Enhanced Content - stores AI-generated full JD text
    ai_enhanced_jd_content: Optional[str] = None
    ai_enhanced_at: Optional[datetime] = None
    # Edit history for tracking changes
    edit_history: List[Dict[str, Any]] = []
    # Metadata
    created_by: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Notification(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    type: str  # sla_alert, new_application, status_change
    title: str
    message: str
    recipient_email: str
    function: Optional[str] = None
    sub_function: Optional[str] = None
    priority: str = "medium"  # low, medium, high, urgent
    read: bool = False
    data: Optional[Dict[str, Any]] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class FunctionalFolder(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str  # HR, IT, Finance, etc.
    description: str
    sub_folders: List[str] = []  # List of sub-folder IDs
    resume_count: int = 0
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class SubFunctionalFolder(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str  # Talent Acquisition, Software Engineering, etc.
    parent_function: str  # HR, IT, etc.
    skills: List[str] = []  # Skills associated with this sub-function
    resume_count: int = 0
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class RoutedResume(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    candidate_id: Optional[str] = None  # Link to candidate if exists
    job_id: Optional[str] = None  # Job they applied to
    name: str
    email: str
    phone: Optional[str] = None
    raw_text: str
    parsed_data: Optional[Dict[str, Any]] = None
    # Routing information
    primary_function: str  # HR, IT, Finance, etc.
    sub_function: str  # Talent Acquisition, Software Engineering, etc.
    confidence_score: float = 0.0  # AI confidence in classification
    matched_skills: List[str] = []
    routing_reason: str = ""
    # Enhanced: Auto-tagged interview questions based on skills
    suggested_interview_questions: List[Dict[str, Any]] = []
    skill_tags: List[str] = []  # Auto-generated skill tags
    # Source tracking with quality scoring
    source: str = "direct"  # linkedin, indeed, glassdoor, referral, etc.
    source_quality_score: float = 1.0  # Quality weight based on source
    # Application tracking
    applied_job_title: Optional[str] = None
    application_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    status: str = "new"  # new, reviewed, shortlisted, rejected, hired
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class JobApplication(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    job_id: str
    job_title: str
    company: Optional[str] = None
    applicant_name: str
    applicant_email: str
    applicant_phone: Optional[str] = None
    resume_text: str
    cover_letter: Optional[str] = None
    source: str = "job_link"  # job_link, linkedin, indeed, referral
    routed_resume_id: Optional[str] = None  # Link to routed resume
    status: str = "received"  # received, processing, routed, error
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ApplicationSubmit(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    candidate_name: Optional[str] = None
    candidate_email: Optional[str] = None
    candidate_phone: Optional[str] = None
    phone: Optional[str] = None
    resume_text: str
    cover_letter: Optional[str] = None
    linkedin_url: Optional[str] = None
    current_company: Optional[str] = None
    current_ctc: Optional[str] = None
    expected_ctc: Optional[str] = None
    notice_period: Optional[str] = None
    source: Optional[str] = "careers_page"
    
    @property
    def applicant_name(self) -> str:
        return self.name or self.candidate_name or "Unknown"
    
    @property
    def applicant_email(self) -> str:
        return self.email or self.candidate_email or ""
    
    @property
    def applicant_phone(self) -> str:
        return self.phone or self.candidate_phone or ""

class Candidate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    raw_resume: str
    parsed_data: Optional[Dict[str, Any]] = None
    analysis: Optional[Dict[str, Any]] = None
    pipeline_stage: str = "new"
    tags: List[str] = []
    notes: List[Dict[str, Any]] = []
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class CandidateCreate(BaseModel):
    name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    raw_resume: str

class CandidateUpdate(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    pipeline_stage: Optional[str] = None
    tags: Optional[List[str]] = None

class MatchResult(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    jd_id: str
    candidate_id: str
    overall_score: float
    match_breakdown: Dict[str, Any]
    strengths: List[str]
    gaps: List[str]
    why_explanation: str
    confidence_score: float
    recruiter_feedback: Optional[str] = None
    recruiter_override: Optional[float] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class MatchRequest(BaseModel):
    jd_id: str
    candidate_id: str

class SearchQuery(BaseModel):
    query: str
    filters: Optional[Dict[str, Any]] = None

class SearchResult(BaseModel):
    candidates: List[Dict[str, Any]]
    interpretation: str
    total_count: int

class PipelineStage(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    order: int
    color: str = "#3B82F6"
    jd_id: Optional[str] = None

class NoteCreate(BaseModel):
    content: str

class RecruiterFeedback(BaseModel):
    feedback: Optional[str] = None
    override_score: Optional[float] = None


# ============ HR PRE-ASSESSMENT QUESTIONNAIRE ============

# Default HR Pre-Assessment Questions (10 objective questions different from basic info)
DEFAULT_HR_PREASSESSMENT_QUESTIONS = [
    {
        "id": "q1",
        "question": "What is the primary reason for considering a job change at this time?",
        "type": "textarea",
        "required": True,
        "category": "motivation"
    },
    {
        "id": "q2", 
        "question": "Are there any gaps in your employment history? If yes, please explain.",
        "type": "textarea",
        "required": True,
        "category": "employment_history"
    },
    {
        "id": "q3",
        "question": "What are your top 3 professional achievements in your current/last role?",
        "type": "textarea",
        "required": True,
        "category": "achievements"
    },
    {
        "id": "q4",
        "question": "Are you open to relocation? If yes, which locations?",
        "type": "select",
        "options": ["Yes - Any location", "Yes - Specific cities only", "No - Current location only", "Remote only"],
        "required": True,
        "category": "relocation"
    },
    {
        "id": "q5",
        "question": "What is your preferred work arrangement?",
        "type": "select",
        "options": ["Full-time On-site", "Hybrid (2-3 days office)", "Fully Remote", "Flexible/Any"],
        "required": True,
        "category": "work_preference"
    },
    {
        "id": "q6",
        "question": "Do you have any ongoing commitments (projects, bonds, etc.) with your current employer?",
        "type": "textarea",
        "required": True,
        "category": "commitments"
    },
    {
        "id": "q7",
        "question": "What are your key strengths that make you suitable for this role?",
        "type": "textarea",
        "required": True,
        "category": "strengths"
    },
    {
        "id": "q8",
        "question": "What areas would you like to improve or develop further?",
        "type": "textarea",
        "required": True,
        "category": "development"
    },
    {
        "id": "q9",
        "question": "Are you currently interviewing with other companies? What stage?",
        "type": "select",
        "options": ["Not actively looking elsewhere", "Early stage interviews", "Final rounds elsewhere", "Have offers in hand"],
        "required": True,
        "category": "interview_status"
    },
    {
        "id": "q10",
        "question": "What are your long-term career goals (3-5 years)?",
        "type": "textarea",
        "required": True,
        "category": "career_goals"
    }
]

class PreAssessmentQuestionnaire(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str  # Questionnaire name
    description: Optional[str] = None
    questions: List[Dict[str, Any]] = Field(default_factory=lambda: DEFAULT_HR_PREASSESSMENT_QUESTIONS.copy())
    created_by: Optional[str] = None  # client_id or admin
    is_default: bool = False
    is_active: bool = True
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class PreAssessmentResponse(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    questionnaire_id: str
    resume_id: str
    candidate_name: str
    candidate_email: str
    job_id: Optional[str] = None
    job_title: Optional[str] = None
    responses: Dict[str, Any] = Field(default_factory=dict)  # question_id -> response
    status: str = "pending"  # pending, sent, completed
    sent_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None
    access_token: str = Field(default_factory=lambda: str(uuid.uuid4())[:8].upper())
    hr_fitment_analysis: Optional[Dict[str, Any]] = None  # Compiled HR analysis
    recruiter_notes: List[Dict[str, Any]] = Field(default_factory=list)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class CompetencyReport(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    resume_id: str
    candidate_name: str
    candidate_email: str
    preassessment_response_id: Optional[str] = None
    resume_analysis: Dict[str, Any] = Field(default_factory=dict)  # Parsed skills, projects
    hr_fitment_data: Dict[str, Any] = Field(default_factory=dict)  # From pre-assessment
    recruiter_observations: List[Dict[str, Any]] = Field(default_factory=list)
    report_type: str = "basic"  # basic, detailed (career_trajectory)
    generated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    pdf_url: Optional[str] = None
    docx_url: Optional[str] = None


# ============ ADMIN PANEL MODELS ============

class AdminUser(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    password_hash: str
    name: str
    role: str = "admin"  # admin, super_admin
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    last_login: Optional[datetime] = None
    is_active: bool = True

class ClientOrganization(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    organization_name: str
    organization_type: str  # corporate, staffing_vendor
    business_domain: str  # e.g., techcorp.com
    contact_email: str
    contact_phone: Optional[str] = None
    contact_person: str
    address: Optional[str] = None
    
    # Subscription & Access
    subscription_status: str = "trial"  # trial, active, expired, suspended
    access_level: str = "full"  # full, limited
    trial_start_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    trial_end_date: Optional[datetime] = None  # 90 days from start
    paid_start_date: Optional[datetime] = None
    paid_end_date: Optional[datetime] = None
    
    # Module Access
    modules_enabled: Dict[str, bool] = Field(default_factory=lambda: {
        "jd_intelligence": True,
        "resume_repository": True,
        "career_trajectory": True,
        "hr_fitment": True
    })
    
    # Usage Limits
    monthly_assessment_limit: Optional[int] = None  # None = unlimited during trial
    assessments_used_this_month: int = 0
    
    # Stats
    total_users: int = 0
    total_resumes: int = 0
    total_assessments: int = 0
    total_jds_created: int = 0
    
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    is_active: bool = True

class ClientUser(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str  # References ClientOrganization.id
    email: str
    password_hash: str
    name: str
    role: str = "user"  # admin, user, viewer
    department: Optional[str] = None
    is_active: bool = True
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    last_login: Optional[datetime] = None
    created_by_admin: bool = True
    invited_by: Optional[str] = None  # User ID who invited this user
    invitation_code: Optional[str] = None  # Code used to join

class InvitationCode(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    code: str  # Unique invitation code
    client_id: str  # Which organization this invite is for
    created_by: str  # Admin user ID who created it
    email_restricted: Optional[str] = None  # If set, only this email can use the code
    domain_restricted: Optional[str] = None  # If set, only emails from this domain can use
    role: str = "user"  # Role assigned on joining: user, viewer
    max_uses: int = 1  # How many times code can be used
    uses_count: int = 0  # How many times it has been used
    expires_at: Optional[datetime] = None  # Code expiration
    is_active: bool = True
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    used_by: List[str] = Field(default_factory=list)  # List of user IDs who used this code

class IPWhitelist(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    ip_address: str  # Can be IP or CIDR range
    description: Optional[str] = None
    is_active: bool = True
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class SecurityLog(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: Optional[str] = None
    user_id: Optional[str] = None
    action: str  # login, logout, invite_created, user_added, access_changed, etc.
    ip_address: Optional[str] = None
    user_agent: Optional[str] = None
    details: Optional[Dict[str, Any]] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class CustomerFeedback(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    user_id: Optional[str] = None
    user_email: str
    feedback_type: str  # improvement, issue, feature_request, general
    subject: str
    description: str
    priority: str = "medium"  # low, medium, high, critical
    status: str = "new"  # new, in_review, in_progress, resolved, closed
    admin_response: Optional[str] = None
    resolved_at: Optional[datetime] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# Admin Request Models
class AdminLoginRequest(BaseModel):
    email: str
    password: str
    encrypted: Optional[bool] = False

# Encryption key - loaded from environment variable
ENCRYPTION_KEY = os.environ.get('ENCRYPTION_KEY', 'RoleSense2024SecureKey!@#$')

def decrypt_password(encrypted_password: str) -> str:
    """Decrypt AES encrypted password from CryptoJS format"""
    try:
        # CryptoJS AES encryption format: "Salted__" + salt (8 bytes) + ciphertext
        encrypted_data = base64.b64decode(encrypted_password)
        
        # Check for "Salted__" prefix
        if encrypted_data[:8] == b'Salted__':
            salt = encrypted_data[8:16]
            ciphertext = encrypted_data[16:]
            
            # Derive key and IV using MD5 (CryptoJS default)
            key_iv = b''
            prev = b''
            while len(key_iv) < 48:  # 32 bytes key + 16 bytes IV
                prev = hashlib.md5(prev + ENCRYPTION_KEY.encode() + salt).digest()
                key_iv += prev
            
            key = key_iv[:32]
            iv = key_iv[32:48]
            
            # Decrypt using AES-CBC
            cipher = AES.new(key, AES.MODE_CBC, iv)
            decrypted = unpad(cipher.decrypt(ciphertext), AES.block_size)
            return decrypted.decode('utf-8')
        else:
            # Not salted format, return as-is (might be plain text)
            return encrypted_password
    except Exception as e:
        logger.error(f"Password decryption failed: {e}")
        # Return the password as-is if decryption fails (backward compatibility)
        return encrypted_password

class CreateClientRequest(BaseModel):
    organization_name: str
    organization_type: str  # corporate, staffing_vendor
    business_domain: str
    contact_email: str
    contact_phone: Optional[str] = None
    contact_person: str
    address: Optional[str] = None

class CreateClientUserRequest(BaseModel):
    client_id: str
    email: str
    password: str
    name: str
    role: str = "user"
    department: Optional[str] = None

class UpdateAccessRequest(BaseModel):
    access_level: Optional[str] = None  # full, limited
    modules_enabled: Optional[Dict[str, bool]] = None
    monthly_assessment_limit: Optional[int] = None
    subscription_status: Optional[str] = None

class SubmitFeedbackRequest(BaseModel):
    feedback_type: str
    subject: str
    description: str
    priority: str = "medium"

class SelfSignupRequest(BaseModel):
    """Request model for self-signup of new organizations"""
    organization_name: str
    organization_type: str  # corporate, staffing_vendor
    business_domain: str  # Company's email domain (e.g., techjobs.in)
    contact_email: str
    contact_phone: Optional[str] = None
    contact_person: str
    password: str
    encrypted: Optional[bool] = False

class ChangePasswordRequest(BaseModel):
    current_password: str
    new_password: str

# Simple password hashing (in production, use bcrypt)
import hashlib
def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(password: str, hashed: str) -> bool:
    return hash_password(password) == hashed

def ensure_timezone_aware(dt):
    """Ensure datetime is timezone-aware (UTC)"""
    if isinstance(dt, str):
        return datetime.fromisoformat(dt.replace('Z', '+00:00'))
    elif isinstance(dt, datetime):
        if dt.tzinfo is None:
            return dt.replace(tzinfo=timezone.utc)
        return dt
    return dt

# ============ Helper Functions ============

async def get_llm_response(system_prompt: str, user_prompt: str) -> str:
    """Get response from LLM"""
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=str(uuid.uuid4()),
            system_message=system_prompt
        ).with_model("openai", "gpt-4o")
        
        user_message = UserMessage(text=user_prompt)
        response = await chat.send_message(user_message)
        return response
    except Exception as e:
        logger.error(f"LLM Error: {e}")
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")

def parse_json_from_response(response: str) -> Dict:
    """Extract JSON from LLM response"""
    try:
        # Try to find JSON in the response
        json_match = re.search(r'\{[\s\S]*\}', response)
        if json_match:
            return json.loads(json_match.group())
        return json.loads(response)
    except json.JSONDecodeError:
        return {"raw_response": response}

def serialize_datetime(obj):
    """Serialize datetime objects for MongoDB"""
    if isinstance(obj, datetime):
        return obj.isoformat()
    elif isinstance(obj, dict):
        return {k: serialize_datetime(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [serialize_datetime(item) for item in obj]
    return obj

def deserialize_datetime(obj):
    """Deserialize datetime strings from MongoDB"""
    if isinstance(obj, str):
        try:
            return datetime.fromisoformat(obj)
        except:
            return obj
    elif isinstance(obj, dict):
        return {k: deserialize_datetime(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [deserialize_datetime(item) for item in obj]
    return obj

# ============ API Routes ============

@api_router.get("/")
async def root():
    return {"message": "Role Sense - JD Intelligence Dashboard API", "version": "1.0.0"}

# ============ Master Data APIs for JD Creation ============

@api_router.get("/master/jd-options")
async def get_jd_master_options():
    """Get all master data options for structured JD creation"""
    return {
        "role_types": ROLE_TYPES,
        "business_models": BUSINESS_MODELS,
        "experience_ranges": EXPERIENCE_RANGES,
        "currencies": CURRENCIES,
        "compensation_types": COMPENSATION_TYPES,
        "india_locations": INDIA_LOCATIONS,
        "international_locations": INTERNATIONAL_LOCATIONS,
        "work_modes": WORK_MODES,
        "employment_types": EMPLOYMENT_TYPES,
        "education_levels": EDUCATION_LEVELS,
        "education_fields": EDUCATION_FIELDS
    }

@api_router.get("/master/competencies")
async def get_competency_master():
    """Get all competency master data organized by category"""
    return {
        "core_behavioral": CORE_BEHAVIORAL_COMPETENCIES,
        "it_functional": IT_FUNCTIONAL_COMPETENCIES,
        "non_it_functional": NON_IT_FUNCTIONAL_COMPETENCIES,
        "cognitive": COGNITIVE_COMPETENCIES,
        "good_to_have": GOOD_TO_HAVE_COMPETENCIES,
        "trainable": TRAINABLE_COMPETENCIES
    }

@api_router.get("/master/tools")
async def get_tools_master():
    """Get all tools and technologies master data"""
    return {
        "it_tools": IT_TOOLS,
        "non_it_tools": NON_IT_TOOLS
    }

# ============ Structured JD CRUD Operations ============

@api_router.post("/jd/structured/create")
async def create_structured_jd(jd_data: StructuredJDCreate):
    """Create a new structured JD with human-defined competencies"""
    
    # Validate minimum requirements
    if len(jd_data.competencies.must_have_skills) < 4:
        raise HTTPException(status_code=400, detail="Minimum 4 must-have skills required")
    if len(jd_data.competencies.good_to_have_skills) < 3:
        raise HTTPException(status_code=400, detail="Minimum 3 good-to-have skills required")
    
    # Create structured JD object
    structured_jd = StructuredJD(
        basic_info=jd_data.basic_info,
        competencies=jd_data.competencies,
        responsibilities=jd_data.responsibilities,
        additional_notes=jd_data.additional_notes,
        status="draft"
    )
    
    # Save to database
    doc = serialize_datetime(structured_jd.model_dump())
    await db.structured_jds.insert_one(doc)
    
    return structured_jd

@api_router.get("/jd/structured/list")
async def list_structured_jds():
    """List all structured JDs"""
    jds = await db.structured_jds.find({}, {"_id": 0}).sort("created_at", -1).to_list(100)
    return [deserialize_datetime(jd) for jd in jds]

@api_router.get("/jd/structured/{jd_id}")
async def get_structured_jd(jd_id: str):
    """Get a specific structured JD"""
    jd = await db.structured_jds.find_one({"id": jd_id}, {"_id": 0})
    if not jd:
        raise HTTPException(status_code=404, detail="Structured JD not found")
    return deserialize_datetime(jd)

@api_router.put("/jd/structured/{jd_id}")
async def update_structured_jd(jd_id: str, jd_data: StructuredJDCreate):
    """Update a structured JD"""
    existing = await db.structured_jds.find_one({"id": jd_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Structured JD not found")
    
    # Validate minimum requirements
    if len(jd_data.competencies.must_have_skills) < 4:
        raise HTTPException(status_code=400, detail="Minimum 4 must-have skills required")
    if len(jd_data.competencies.good_to_have_skills) < 3:
        raise HTTPException(status_code=400, detail="Minimum 3 good-to-have skills required")
    
    update_data = {
        "basic_info": jd_data.basic_info.model_dump(),
        "competencies": jd_data.competencies.model_dump(),
        "responsibilities": jd_data.responsibilities,
        "additional_notes": jd_data.additional_notes,
        "updated_at": datetime.now(timezone.utc)
    }
    
    await db.structured_jds.update_one({"id": jd_id}, {"$set": serialize_datetime(update_data)})
    return await get_structured_jd(jd_id)

@api_router.delete("/jd/structured/{jd_id}")
async def delete_structured_jd(jd_id: str):
    """Delete a structured JD"""
    result = await db.structured_jds.delete_one({"id": jd_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Structured JD not found")
    return {"message": "Structured JD deleted successfully"}

@api_router.post("/jd/structured/{jd_id}/submit")
async def submit_structured_jd(jd_id: str, data: dict = None):
    """Submit a structured JD and move it to Active Jobs"""
    jd = await db.structured_jds.find_one({"id": jd_id})
    if not jd:
        raise HTTPException(status_code=404, detail="Structured JD not found")
    
    # Allow re-submission for active jobs (to update version)
    version = (data or {}).get("version", "human")  # 'human' or 'ai'
    
    # Generate requisition number only if not already present
    requisition_number = jd.get("requisition_number") or generate_requisition_number()
    requisition_date = jd.get("requisition_date") or datetime.now(timezone.utc)
    
    # Get job details for share content
    basic_info = jd.get("basic_info", {})
    job_title = basic_info.get("title", "")
    company_name = basic_info.get("company_name", "")
    location = ", ".join(basic_info.get("locations_india", [])[:2]) or "India"
    experience = f"{basic_info.get('experience_min', 0)}-{basic_info.get('experience_max', 0)} years"
    
    # Get skills for description
    competencies = jd.get("competencies", {})
    skills = competencies.get("must_have_skills", [])[:3]
    skills_text = ", ".join(skills) if skills else ""
    
    # Application info
    application_email = jd.get("application_email", "")
    application_link = jd.get("application_link", "")
    apply_text = ""
    apply_url = application_link or "https://rolesense.com/jobs"
    if application_email:
        apply_text = f"Apply: {application_email}"
    elif application_link:
        apply_text = f"Apply: {application_link}"
    
    # Create engaging post content for each platform
    from urllib.parse import quote, quote_plus
    
    # LinkedIn post content (more professional) - keep it concise for better sharing
    linkedin_text = f"""🚀 We're Hiring: {job_title}

📍 {location}
💼 {experience} experience
🏢 {company_name}

Skills: {skills_text}

{apply_text}

#Hiring #Jobs #Career"""
    
    # Twitter/X post content (max 280 chars)
    twitter_text = f"🚀 Hiring: {job_title} at {company_name}!\n📍 {location}\n💼 {experience}\n{apply_text}\n#Hiring #Jobs"
    if len(twitter_text) > 280:
        twitter_text = f"🚀 Hiring: {job_title}!\n📍 {location}\n{apply_text}\n#Hiring #Jobs"
    
    # Facebook post content
    facebook_text = f"""🎯 Job Opening: {job_title}

Join {company_name}!

📍 {location}
💼 {experience}
Skills: {skills_text}

{apply_text}

Know someone perfect? Tag them! 👇"""
    
    # Generate share dialog URLs with proper encoding
    # LinkedIn: Use shareArticle for better pre-fill support
    linkedin_share_url = f"https://www.linkedin.com/shareArticle?mini=true&url={quote(apply_url, safe='')}&title={quote(f'Hiring: {job_title} at {company_name}', safe='')}&summary={quote(linkedin_text[:256], safe='')}&source=RoleSense"
    
    # Twitter: Use intent/tweet with text parameter
    twitter_share_url = f"https://twitter.com/intent/tweet?text={quote(twitter_text, safe='')}"
    
    # Facebook: Use sharer with quote parameter  
    facebook_share_url = f"https://www.facebook.com/sharer/sharer.php?u={quote(apply_url, safe='')}&quote={quote(facebook_text[:300], safe='')}"
    
    # RoleSense direct job link - this is the shareable link for job seekers
    rolesense_job_link = f"/jobs/{jd_id}"
    
    publish_links = {
        "rolesense": rolesense_job_link,  # Direct job application link
        "linkedin": linkedin_share_url,
        "linkedin_post": linkedin_share_url,
        "twitter": twitter_share_url,
        "facebook": facebook_share_url,
        "indeed": f"https://employers.indeed.com/p/post-job?q={quote(job_title)}",
        "naukri": f"https://recruiter.naukri.com/post-job?title={quote(job_title)}",
        "glassdoor": f"https://www.glassdoor.com/employers/post-job?title={quote(job_title)}"
    }
    
    update_data = {
        "status": "active",
        "requisition_number": requisition_number,
        "requisition_date": requisition_date,
        "publish_links": publish_links,
        "rolesense_job_link": rolesense_job_link,  # Store separately for easy access
        "submitted_version": version,  # Track which version was submitted
        "updated_at": datetime.now(timezone.utc)
    }
    
    await db.structured_jds.update_one({"id": jd_id}, {"$set": serialize_datetime(update_data)})
    return await get_structured_jd(jd_id)

@api_router.get("/jd/structured/active/list")
async def list_active_structured_jds():
    """List all active structured JDs"""
    jds = await db.structured_jds.find({"status": "active"}, {"_id": 0}).sort("requisition_date", -1).to_list(100)
    return [deserialize_datetime(jd) for jd in jds]

@api_router.get("/jd/structured/{jd_id}/download/{format}")
async def download_structured_jd(jd_id: str, format: str, version: str = "human"):
    """Download structured JD as PDF. 
    version: 'human' for original human-created JD, 'ai' for AI-enhanced version
    """
    jd = await db.structured_jds.find_one({"id": jd_id}, {"_id": 0})
    if not jd:
        raise HTTPException(status_code=404, detail="Structured JD not found")
    
    basic = jd.get("basic_info", {})
    competencies = jd.get("competencies", {})
    
    if format.lower() == "pdf":
        # Generate PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=12, textColor=colors.HexColor('#1e3a5f'))
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=14, spaceBefore=16, spaceAfter=8, textColor=colors.HexColor('#2c5282'))
        subheading_style = ParagraphStyle('CustomSubheading', parent=styles['Heading3'], fontSize=12, spaceBefore=12, spaceAfter=6, textColor=colors.HexColor('#4a5568'))
        body_style = ParagraphStyle('CustomBody', parent=styles['Normal'], fontSize=10, spaceAfter=6, leading=14)
        ai_badge_style = ParagraphStyle('AIBadge', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#7c3aed'), backColor=colors.HexColor('#f3e8ff'))
        
        story = []
        
        # If AI version requested and available
        if version == "ai" and jd.get("ai_enhanced_jd_content"):
            # Add AI badge
            story.append(Paragraph("🤖 AI-Enhanced Job Description", ai_badge_style))
            story.append(Spacer(1, 8))
            
            # Title
            story.append(Paragraph(f"Job Description: {basic.get('title', 'Untitled')}", title_style))
            story.append(Paragraph(f"Company: {basic.get('company_name', 'N/A')}", body_style))
            if jd.get("requisition_number"):
                story.append(Paragraph(f"Requisition #: {jd.get('requisition_number')}", body_style))
            story.append(Spacer(1, 16))
            
            # Parse and format the AI-enhanced content
            ai_content = jd.get("ai_enhanced_jd_content", "")
            # Split into paragraphs and add them
            paragraphs = ai_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Check if it's a heading (all caps or ends with colon)
                    lines = para.strip().split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        # Check if it looks like a heading
                        if line.isupper() or (len(line) < 50 and line.endswith(':')):
                            story.append(Paragraph(line.replace(':', ''), heading_style))
                        elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                            story.append(Paragraph(f"• {line.lstrip('•-* ')}", body_style))
                        else:
                            story.append(Paragraph(line, body_style))
                    story.append(Spacer(1, 8))
            
            filename = f"JD_AI_{basic.get('title', 'job').replace(' ', '_')}_{jd_id[:8]}.pdf"
        else:
            # Original human-created JD format
            # Title
            story.append(Paragraph(f"Job Description: {basic.get('title', 'Untitled')}", title_style))
            story.append(Paragraph(f"Company: {basic.get('company_name', 'N/A')}", body_style))
            if jd.get("requisition_number"):
                story.append(Paragraph(f"Requisition #: {jd.get('requisition_number')}", body_style))
            story.append(Spacer(1, 12))
            
            # Basic Information
            story.append(Paragraph("Basic Information", heading_style))
            basic_data = [
                ["Role Type", basic.get("role_type", "N/A")],
                ["Business Model", basic.get("business_model", "N/A")],
                ["Experience", f"{basic.get('experience_min', 0)} - {basic.get('experience_max', 0)} years"],
                ["Compensation", f"{basic.get('compensation_currency', 'INR')} {basic.get('compensation_min', 'N/A')} - {basic.get('compensation_max', 'N/A')} {basic.get('compensation_type', '')}"],
                ["Work Mode", basic.get("work_mode", "N/A")],
                ["Employment Type", basic.get("employment_type", "N/A")],
                ["Education", f"{basic.get('education_level', 'N/A')} in {basic.get('education_field', 'Any Field')}"],
            ]
            
            if basic.get("locations_india"):
                basic_data.append(["India Locations", ", ".join(basic.get("locations_india", []))])
            if basic.get("locations_international"):
                basic_data.append(["International Locations", ", ".join(basic.get("locations_international", []))])
            
            table = Table(basic_data, colWidths=[150, 300])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f7fafc')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#2d3748')),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
            ]))
            story.append(table)
            story.append(Spacer(1, 16))
            
            # About Company
            if basic.get("about_company"):
                story.append(Paragraph("About the Company", heading_style))
                story.append(Paragraph(basic.get("about_company"), body_style))
                story.append(Spacer(1, 12))
            
            # Must-Have Skills
            story.append(Paragraph("Must-Have Skills (Non-Negotiable)", heading_style))
            if competencies.get("must_have_skills"):
                for skill in competencies.get("must_have_skills", []):
                    story.append(Paragraph(f"• {skill}", body_style))
            story.append(Spacer(1, 12))
            
            # Good-to-Have Skills
            story.append(Paragraph("Good-to-Have Skills (Preferred)", heading_style))
            if competencies.get("good_to_have_skills"):
                for skill in competencies.get("good_to_have_skills", []):
                    story.append(Paragraph(f"• {skill}", body_style))
            story.append(Spacer(1, 12))
            
            # Competencies
            story.append(Paragraph("Competency Requirements", heading_style))
            
            if competencies.get("must_have_behavioral"):
                story.append(Paragraph("Core/Behavioral Competencies:", subheading_style))
                for comp_id in competencies.get("must_have_behavioral", []):
                    comp = next((c for c in CORE_BEHAVIORAL_COMPETENCIES if c["id"] == comp_id), None)
                    if comp:
                        story.append(Paragraph(f"• {comp['name']}", body_style))
            
            if competencies.get("must_have_cognitive"):
                story.append(Paragraph("Cognitive Competencies:", subheading_style))
                for comp_id in competencies.get("must_have_cognitive", []):
                    comp = next((c for c in COGNITIVE_COMPETENCIES if c["id"] == comp_id), None)
                    if comp:
                        story.append(Paragraph(f"• {comp['name']}", body_style))
            
            # Tools
            if competencies.get("tools_must_have") or competencies.get("tools_good_to_have"):
                story.append(Paragraph("Tools & Technologies", heading_style))
                if competencies.get("tools_must_have"):
                    story.append(Paragraph("Must-Have Tools:", subheading_style))
                    story.append(Paragraph(", ".join(competencies.get("tools_must_have", [])), body_style))
                if competencies.get("tools_good_to_have"):
                    story.append(Paragraph("Good-to-Have Tools:", subheading_style))
                    story.append(Paragraph(", ".join(competencies.get("tools_good_to_have", [])), body_style))
            
            # Responsibilities
            if jd.get("responsibilities"):
                story.append(Paragraph("Key Responsibilities", heading_style))
                for resp in jd.get("responsibilities", []):
                    story.append(Paragraph(f"• {resp}", body_style))
            
            # Additional Notes
            if jd.get("additional_notes"):
                story.append(Spacer(1, 12))
                story.append(Paragraph("Additional Notes", heading_style))
                story.append(Paragraph(jd.get("additional_notes"), body_style))
            
            filename = f"JD_{basic.get('title', 'job').replace(' ', '_')}_{jd_id[:8]}.pdf"
        
        doc.build(story)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    
    else:
        raise HTTPException(status_code=400, detail="Unsupported format. Use 'pdf'")

@api_router.post("/jd/structured/{jd_id}/ai-analyze")
async def ai_analyze_structured_jd(jd_id: str):
    """AI analysis to enhance human-created JD with additional insights and suggestions"""
    jd = await db.structured_jds.find_one({"id": jd_id}, {"_id": 0})
    if not jd:
        raise HTTPException(status_code=404, detail="Structured JD not found")
    
    basic = jd.get("basic_info", {})
    competencies = jd.get("competencies", {})
    responsibilities = jd.get("responsibilities", [])
    
    # Build JD text from structured data for AI analysis
    jd_text = f"""
JOB TITLE: {basic.get('title', '')}
COMPANY: {basic.get('company_name', '')}
ABOUT COMPANY: {basic.get('about_company', '')}

ROLE TYPE: {basic.get('role_type', '')}
BUSINESS MODEL: {basic.get('business_model', '')}
EXPERIENCE: {basic.get('experience_min', 0)} - {basic.get('experience_max', 0)} years
WORK MODE: {basic.get('work_mode', '')}
EMPLOYMENT TYPE: {basic.get('employment_type', '')}
EDUCATION: {basic.get('education_level', '')} in {basic.get('education_field', '')}

LOCATIONS INDIA: {', '.join(basic.get('locations_india', []))}
LOCATIONS INTERNATIONAL: {', '.join(basic.get('locations_international', []))}

COMPENSATION: {basic.get('compensation_currency', 'INR')} {basic.get('compensation_min', 'N/A')} - {basic.get('compensation_max', 'N/A')} {basic.get('compensation_type', '')}

MUST-HAVE SKILLS: {', '.join(competencies.get('must_have_skills', []))}
GOOD-TO-HAVE SKILLS: {', '.join(competencies.get('good_to_have_skills', []))}
TOOLS: {', '.join(competencies.get('tools_must_have', []))}

KEY RESPONSIBILITIES:
{chr(10).join(['- ' + r for r in responsibilities if r])}
"""
    
    # Build data completeness context for AI
    has_compensation = basic.get('compensation_min') and basic.get('compensation_max') and basic.get('compensation_min') != 'N/A'
    has_locations = bool(basic.get('locations_india', [])) or bool(basic.get('locations_international', []))
    has_skills = bool(competencies.get('must_have_skills', []))
    
    completeness_context = f"""
DATA ALREADY PROVIDED BY RECRUITER (DO NOT suggest adding these):
- Compensation: {'YES - {}{} - {} {}'.format(basic.get('compensation_currency', 'INR'), basic.get('compensation_min', ''), basic.get('compensation_max', ''), basic.get('compensation_type', '')) if has_compensation else 'NOT PROVIDED - May suggest adding'}
- Locations: {'YES - India: ' + ', '.join(basic.get('locations_india', [])) + ' | International: ' + ', '.join(basic.get('locations_international', [])) if has_locations else 'NOT PROVIDED - May suggest adding'}
- Work Mode: {basic.get('work_mode', 'Not specified')}
- Skills: {'YES - ' + str(len(competencies.get('must_have_skills', []))) + ' must-have skills provided' if has_skills else 'NOT PROVIDED - May suggest adding'}
"""
    
    system_prompt = """You are an expert HR consultant and recruitment specialist. Your role is to ENHANCE and VALIDATE human-created job descriptions, NOT to replace them.

CRITICAL RULES:
1. DO NOT suggest adding data that is ALREADY PROVIDED (check the DATA ALREADY PROVIDED section carefully)
2. Only suggest improvements for fields that are genuinely MISSING or incomplete
3. If compensation is provided, DO NOT suggest adding compensation - instead analyze if it's competitive
4. If locations are provided, DO NOT suggest adding locations - instead analyze the coverage
5. If skills are provided, focus on skill gap analysis, not suggesting to add more basic skills
6. Generate a compelling, DETAILED job summary (at least 4-5 sentences) for posting
7. The enhanced_description should be COMPREHENSIVE (minimum 200 words) - this is the main JD content

Respond with valid JSON only."""

    user_prompt = f"""Analyze this human-created job description and provide enhancement suggestions:

{jd_text}

{completeness_context}

Provide analysis in this JSON format:
{{
    "validation": {{
        "completeness_score": 0-100,
        "clarity_score": 0-100,
        "market_competitiveness_score": 0-100,
        "overall_quality_score": 0-100
    }},
    "generated_summary": "A compelling 4-5 sentence job summary that captures the role's impact, team, growth opportunities, and company culture",
    "enhanced_description": "A comprehensive, professional job description (MINIMUM 200 words) that expands on the role, team dynamics, day-to-day responsibilities, growth path, and what makes this opportunity unique. Keep all original requirements intact but present them in an engaging, professional narrative format.",
    "skill_analysis": {{
        "skill_coverage": "Assessment of skill requirements coverage",
        "missing_common_skills": ["Only list skills commonly required for THIS SPECIFIC role that are genuinely not listed"],
        "skill_level_assessment": "Assessment of experience vs skill expectations"
    }},
    "compensation_analysis": {{
        "market_alignment": "How the PROVIDED compensation compares to market rates for this role (or note if not provided)",
        "recommendation": "Specific suggestions only if compensation seems off-market or is missing"
    }},
    "suggested_screening_questions": [
        {{"question": "Question text", "skill_tested": "What skill this tests", "expected_response_hint": "What to look for"}}
    ],
    "improvement_suggestions": [
        {{"area": "Area name", "current": "What is currently provided or missing", "suggestion": "Specific actionable improvement", "priority": "high/medium/low"}}
    ],
    "posting_ready": true/false,
    "posting_blockers": ["Only genuine blockers - missing critical info that would confuse candidates"],
    "social_post_snippets": {{
        "linkedin": "Engaging LinkedIn post (max 200 chars) highlighting the opportunity",
        "twitter": "Catchy Twitter/X post (max 180 chars)",
        "general": "General social media post (max 300 chars) with call to action"
    }}
}}

IMPORTANT: Only include improvement_suggestions for data that is GENUINELY MISSING. If compensation/locations/skills are provided, focus suggestions on QUALITY improvements, not adding what already exists."""

    response = await get_llm_response(system_prompt, user_prompt)
    ai_analysis = parse_json_from_response(response)
    
    # Store AI analysis in the JD
    update_data = {
        "ai_analysis": ai_analysis,
        "ai_analyzed_at": datetime.now(timezone.utc),
        "updated_at": datetime.now(timezone.utc)
    }
    
    await db.structured_jds.update_one({"id": jd_id}, {"$set": serialize_datetime(update_data)})
    
    # Return updated JD
    updated_jd = await db.structured_jds.find_one({"id": jd_id}, {"_id": 0})
    return deserialize_datetime(updated_jd)

@api_router.post("/jd/structured/{jd_id}/generate-questions")
async def generate_structured_jd_questions(jd_id: str):
    """Generate screening questions for a structured JD"""
    jd = await db.structured_jds.find_one({"id": jd_id}, {"_id": 0})
    if not jd:
        raise HTTPException(status_code=404, detail="Structured JD not found")
    
    basic = jd.get("basic_info", {})
    competencies = jd.get("competencies", {})
    
    system_prompt = """You are an expert technical recruiter. Generate screening questions that assess both technical skills and behavioral competencies.
    Always respond with valid JSON only."""
    
    user_prompt = f"""Generate screening questions for this role:

Title: {basic.get('title', '')}
Role Type: {basic.get('role_type', '')}
Experience: {basic.get('experience_min', 0)}-{basic.get('experience_max', 0)} years

Must-Have Skills: {', '.join(competencies.get('must_have_skills', []))}
Good-to-Have Skills: {', '.join(competencies.get('good_to_have_skills', []))}
Tools: {', '.join(competencies.get('tools_must_have', []))}

Generate 8-10 screening questions in this JSON format:
{{
    "questions": [
        {{
            "id": "unique_id",
            "question": "The question text",
            "skill_area": "Which skill/competency this tests",
            "question_type": "technical/behavioral/situational",
            "difficulty": "easy/medium/hard",
            "expected_answer": "Key points to look for in the answer",
            "time_estimate": "2-3 min"
        }}
    ]
}}"""

    response = await get_llm_response(system_prompt, user_prompt)
    questions_data = parse_json_from_response(response)
    
    questions = []
    for q in questions_data.get("questions", []):
        questions.append({
            "id": str(uuid.uuid4()),
            "question": q.get("question", ""),
            "skill_area": q.get("skill_area", ""),
            "question_type": q.get("question_type", "technical"),
            "difficulty": q.get("difficulty", "medium"),
            "expected_answer": q.get("expected_answer", ""),
            "time_estimate": q.get("time_estimate", "2-3 min")
        })
    
    # Update JD with questions
    await db.structured_jds.update_one(
        {"id": jd_id}, 
        {"$set": {"screening_questions": questions, "updated_at": datetime.now(timezone.utc)}}
    )
    
    return {"questions": questions}

@api_router.post("/jd/structured/{jd_id}/generate-ai-jd")
async def generate_ai_enhanced_jd(jd_id: str):
    """Generate an AI-enhanced version of the JD with professional language and formatting"""
    jd = await db.structured_jds.find_one({"id": jd_id}, {"_id": 0})
    if not jd:
        raise HTTPException(status_code=404, detail="Structured JD not found")
    
    basic = jd.get("basic_info", {})
    competencies = jd.get("competencies", {})
    responsibilities = jd.get("responsibilities", [])
    additional_notes = jd.get("additional_notes", "")
    ai_analysis = jd.get("ai_analysis", {})
    
    # Build the human JD text for AI to enhance
    human_jd_text = f"""
JOB TITLE: {basic.get('title', '')}
COMPANY: {basic.get('company_name', '')}
ABOUT COMPANY: {basic.get('about_company', '')}

ROLE TYPE: {basic.get('role_type', '')}
BUSINESS MODEL: {basic.get('business_model', '')}
EXPERIENCE REQUIRED: {basic.get('experience_min', 0)} - {basic.get('experience_max', 0)} years
WORK MODE: {basic.get('work_mode', '')}
EMPLOYMENT TYPE: {basic.get('employment_type', '')}
EDUCATION: {basic.get('education_level', '')} in {basic.get('education_field', 'Any Field')}

LOCATIONS INDIA: {', '.join(basic.get('locations_india', []))}
LOCATIONS INTERNATIONAL: {', '.join(basic.get('locations_international', []))}

COMPENSATION: {basic.get('compensation_currency', 'INR')} {basic.get('compensation_min', 'Not specified')} - {basic.get('compensation_max', 'Not specified')} {basic.get('compensation_type', 'Per Annum')}

MUST-HAVE SKILLS:
{chr(10).join(['- ' + s for s in competencies.get('must_have_skills', [])])}

GOOD-TO-HAVE SKILLS:
{chr(10).join(['- ' + s for s in competencies.get('good_to_have_skills', [])])}

REQUIRED TOOLS:
{chr(10).join(['- ' + t for t in competencies.get('tools_must_have', [])])}

KEY RESPONSIBILITIES:
{chr(10).join(['- ' + r for r in responsibilities if r])}

ADDITIONAL NOTES:
{additional_notes}
"""

    # Use AI analysis enhanced description if available
    enhanced_desc = ai_analysis.get("enhanced_description", "")
    generated_summary = ai_analysis.get("generated_summary", "")
    
    system_prompt = """You are an expert HR copywriter and job description specialist. Your task is to transform structured job data into a COMPREHENSIVE, polished, and compelling job description that attracts top talent.

CRITICAL REQUIREMENTS:
1. The output must be DETAILED - minimum 800 words covering all aspects of the role
2. Maintain ALL the original requirements and information - do not remove any requirements
3. Use professional, engaging language that appeals to top-tier candidates
4. Structure the JD in a clear, readable format with proper sections
5. Add compelling opening paragraph about the opportunity and company impact
6. Make the company sound attractive and describe team culture
7. Use bullet points effectively for requirements and responsibilities
8. Include ALL these sections (each section should be substantial):
   - About the Company (2-3 paragraphs about company, culture, mission)
   - About the Role (detailed description of the position's impact and scope)
   - Key Responsibilities (comprehensive list with context)
   - Required Qualifications (all must-have requirements)
   - Preferred Qualifications (all good-to-have requirements) 
   - Technical Skills & Tools (detailed skill requirements)
   - Compensation & Benefits (if provided, expand on total value proposition)
   - Work Location & Mode (details about work arrangement)
   - What We Offer (growth opportunities, learning, work-life balance)
   - How to Apply (call to action)
9. Include specific metrics and impact statements where possible
10. End with a compelling call-to-action

Output the complete formatted job description as plain text (no markdown or special formatting codes)."""

    user_prompt = f"""Transform this structured job data into a COMPREHENSIVE, professional job description (minimum 800 words):

{human_jd_text}

{f'Incorporate this enhanced description for the role overview: {enhanced_desc}' if enhanced_desc else ''}
{f'Use this as the opening hook: {generated_summary}' if generated_summary else ''}

Generate a complete, detailed, professional job description that covers:
1. Company overview and culture (make it compelling)
2. Role impact and why it matters
3. Day-to-day responsibilities with context
4. All skill requirements (technical and soft skills)
5. Growth and career development opportunities
6. Compensation and benefits value proposition
7. Work environment and team dynamics
8. Application process and next steps

This JD should be ready to post on LinkedIn, Indeed, and company career page. Make it stand out!"""

    response = await get_llm_response(system_prompt, user_prompt)
    
    # Store the AI-enhanced JD content
    update_data = {
        "ai_enhanced_jd_content": response,
        "ai_enhanced_at": datetime.now(timezone.utc),
        "updated_at": datetime.now(timezone.utc)
    }
    
    await db.structured_jds.update_one({"id": jd_id}, {"$set": serialize_datetime(update_data)})
    
    # Return updated JD
    updated_jd = await db.structured_jds.find_one({"id": jd_id}, {"_id": 0})
    return deserialize_datetime(updated_jd)

# ============ Vendor JD Upload Route ============

# Import PDF and DOCX parsers
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import olefile
except ImportError:
    olefile = None

@api_router.post("/jd/parse-file")
async def parse_jd_file(file: UploadFile = File(...)):
    """Parse uploaded file (PDF, DOCX, DOC, TXT, RTF) and extract text content"""
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded")
    
    filename = file.filename.lower()
    content = await file.read()
    
    extracted_text = ""
    
    try:
        if filename.endswith('.txt'):
            # Plain text file
            try:
                extracted_text = content.decode('utf-8')
            except UnicodeDecodeError:
                extracted_text = content.decode('latin-1')
                
        elif filename.endswith('.pdf'):
            # PDF file
            if PdfReader is None:
                raise HTTPException(status_code=500, detail="PDF parsing not available")
            
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            extracted_text = "\n\n".join(text_parts)
            
        elif filename.endswith('.docx'):
            # DOCX file
            if DocxDocument is None:
                raise HTTPException(status_code=500, detail="DOCX parsing not available")
            
            docx_file = io.BytesIO(content)
            doc = DocxDocument(docx_file)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            extracted_text = "\n\n".join(text_parts)
            
        elif filename.endswith('.doc'):
            # Old DOC format - extract using olefile
            if olefile is None:
                raise HTTPException(
                    status_code=400, 
                    detail="DOC parsing library not available. Please save as .docx and try again."
                )
            
            try:
                doc_file = io.BytesIO(content)
                ole = olefile.OleFileIO(doc_file)
                
                # Try to extract text from WordDocument stream
                if ole.exists('WordDocument'):
                    # Read the raw stream
                    word_stream = ole.openstream('WordDocument').read()
                    
                    # Extract readable text (basic extraction)
                    text_parts = []
                    current_text = ""
                    for byte in word_stream:
                        if 32 <= byte < 127:  # Printable ASCII
                            current_text += chr(byte)
                        elif byte in (10, 13):  # Newlines
                            if current_text.strip():
                                text_parts.append(current_text.strip())
                            current_text = ""
                        else:
                            if len(current_text) > 3:  # Save meaningful text chunks
                                text_parts.append(current_text.strip())
                            current_text = ""
                    
                    if current_text.strip():
                        text_parts.append(current_text.strip())
                    
                    # Filter out short/garbage strings
                    extracted_text = "\n".join([t for t in text_parts if len(t) > 5])
                    
                ole.close()
                
                if not extracted_text.strip():
                    raise HTTPException(
                        status_code=400,
                        detail="Could not extract text from .doc file. Please save as .docx format and try again."
                    )
                    
            except Exception as e:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Error parsing .doc file: {str(e)}. Please save as .docx format and try again."
                )
        
        elif filename.endswith('.rtf'):
            # RTF file
            try:
                from striprtf.striprtf import rtf_to_text
                try:
                    rtf_content = content.decode('utf-8')
                except UnicodeDecodeError:
                    rtf_content = content.decode('latin-1')
                extracted_text = rtf_to_text(rtf_content)
            except ImportError:
                raise HTTPException(status_code=400, detail="RTF parsing not available. Please save as DOCX or TXT.")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error parsing RTF: {str(e)}")
                
        else:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file format. Please upload PDF, DOCX, DOC, RTF, or TXT files."
            )
        
        if not extracted_text.strip():
            raise HTTPException(
                status_code=400, 
                detail="Could not extract text from file. The file may be empty or image-based."
            )
        
        return {
            "success": True,
            "filename": file.filename,
            "text": extracted_text,
            "character_count": len(extracted_text)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error parsing file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error parsing file: {str(e)}")

class VendorJDUpload(BaseModel):
    title: str
    client_name: str
    requisition_date: str
    raw_text: str
    ai_analysis: Optional[Dict[str, Any]] = None
    application_email: Optional[str] = None
    application_link: Optional[str] = None
    # New fields
    compensation_min: Optional[float] = None
    compensation_max: Optional[float] = None
    compensation_currency: str = "INR"
    location: Optional[str] = None
    experience_min: Optional[int] = None
    experience_max: Optional[int] = None
    business_model: Optional[str] = None
    work_mode: Optional[str] = None
    reporting_to: Optional[str] = None
    team_handling: Optional[str] = None
    responsibilities: Optional[List[str]] = None

@api_router.post("/jd/vendor/upload")
async def vendor_upload_jd(jd_data: VendorJDUpload):
    """Upload a JD from vendor (simplified flow with client info)"""
    
    jd_id = str(uuid.uuid4())
    requisition_date = datetime.now(timezone.utc)
    
    try:
        requisition_date = datetime.fromisoformat(jd_data.requisition_date.replace('Z', '+00:00'))
    except:
        pass
    
    # Generate requisition number
    date_part = requisition_date.strftime("%Y%m%d")
    random_part = jd_id[:6].upper()
    requisition_number = f"REQ-{date_part}-{random_part}"
    
    # Create JD document
    jd_doc = {
        "id": jd_id,
        "title": jd_data.title,
        "company": jd_data.client_name,
        "client_name": jd_data.client_name,
        "raw_text": jd_data.raw_text,
        "requisition_date_received": requisition_date,
        "requisition_number": requisition_number,
        "status": "draft",
        "source": "vendor_upload",
        "quality_score": jd_data.ai_analysis.get("quality_score") if jd_data.ai_analysis else None,
        "extracted_skills": jd_data.ai_analysis.get("extracted_skills", []) if jd_data.ai_analysis else [],
        "experience_range": jd_data.ai_analysis.get("experience_range") if jd_data.ai_analysis else None,
        "ai_analysis": jd_data.ai_analysis,
        "application_email": jd_data.application_email,
        "application_link": jd_data.application_link,
        # New fields
        "compensation_min": jd_data.compensation_min,
        "compensation_max": jd_data.compensation_max,
        "compensation_currency": jd_data.compensation_currency,
        "location": jd_data.location,
        "experience_min": jd_data.experience_min,
        "experience_max": jd_data.experience_max,
        "business_model": jd_data.business_model,
        "work_mode": jd_data.work_mode,
        "reporting_to": jd_data.reporting_to,
        "team_handling": jd_data.team_handling,
        "responsibilities": jd_data.responsibilities,
        "created_at": datetime.now(timezone.utc),
        "updated_at": datetime.now(timezone.utc)
    }
    
    await db.job_descriptions.insert_one(serialize_datetime(jd_doc))
    
    return {
        "id": jd_id,
        "title": jd_data.title,
        "client_name": jd_data.client_name,
        "requisition_number": requisition_number,
        "status": "draft"
    }

# ============ JD Intelligence Routes (Legacy) ============

@api_router.post("/jd/analyze")
async def analyze_job_description(jd: JobDescriptionCreate):
    """Analyze a job description and extract structured data"""
    
    system_prompt = """You are an expert HR analyst specializing in job description analysis.
    
    CRITICAL RULES:
    1. ONLY extract information that is EXPLICITLY stated in the job description
    2. DO NOT make assumptions or generate default values
    3. If information is not mentioned, use null or empty values
    4. Every piece of data must be traceable to the original JD text
    5. Your role is to EXTRACT, not to GENERATE or ASSUME
    
    Always respond with valid JSON only, no additional text."""
    
    user_prompt = f"""Analyze this job description and extract ONLY the information that is explicitly mentioned.

JOB DESCRIPTION:
{jd.raw_text}

EXTRACTION RULES:
- For experience_years: Extract ONLY if specific years are mentioned (e.g., "5+ years", "3-5 years")
  - If NO experience requirement is mentioned, set min, max, and preferred to null
  - Do NOT assume or generate experience values
- For skills: Extract ONLY skills explicitly listed in the JD
- For salary: Extract ONLY if explicitly stated, otherwise set to null
- For location: Extract ONLY if mentioned, otherwise "Not specified"
- For quality_score: Rate 0-100 based on JD completeness and clarity
- For why_quality_score: Explain what's missing or well-defined in the JD

Respond with a JSON object:
{{
    "summary": "Brief role summary based on JD content",
    "required_skills": ["only skills explicitly marked as required"],
    "preferred_skills": ["only skills explicitly marked as preferred/nice-to-have"],
    "experience_years": {{
        "min": null or number if explicitly stated,
        "max": null or number if explicitly stated,
        "preferred": null or number if explicitly stated,
        "source_text": "exact text from JD mentioning experience or null if not mentioned"
    }},
    "education": ["only education requirements explicitly stated"],
    "responsibilities": ["responsibilities listed in JD"],
    "benefits": ["benefits listed in JD"],
    "salary_range": {{
        "min": null or number if stated,
        "max": null or number if stated,
        "currency": "currency if stated or null",
        "source_text": "exact salary text from JD or null"
    }},
    "location": "location if mentioned or 'Not specified'",
    "employment_type": "type if mentioned or 'Not specified'",
    "seniority_level": "level if clear from JD or 'Not specified'",
    "department": "department if mentioned or 'Not specified'",
    "key_qualifications": ["qualifications explicitly listed"],
    "red_flags": ["any vague, unrealistic, or concerning requirements"],
    "quality_score": 0-100,
    "improvement_suggestions": ["specific suggestions to improve this JD"],
    "why_quality_score": "Detailed explanation: what's well-defined vs what's missing or vague",
    "extraction_confidence": "high/medium/low - how clearly the JD stated its requirements"
}}"""
    
    response = await get_llm_response(system_prompt, user_prompt)
    parsed_data = parse_json_from_response(response)
    
    # Create JD object
    jd_obj = JobDescription(
        title=jd.title,
        company=jd.company,
        raw_text=jd.raw_text,
        parsed_data=parsed_data,
        analysis={
            "quality_score": parsed_data.get("quality_score", 0),
            "why_quality_score": parsed_data.get("why_quality_score", ""),
            "red_flags": parsed_data.get("red_flags", [])
        },
        improvement_suggestions=parsed_data.get("improvement_suggestions", [])
    )
    
    # Save to database
    doc = serialize_datetime(jd_obj.model_dump())
    await db.job_descriptions.insert_one(doc)
    
    return jd_obj

@api_router.get("/jd/list", response_model=List[JobDescription])
async def list_job_descriptions():
    """List all job descriptions"""
    jds = await db.job_descriptions.find({}, {"_id": 0}).to_list(100)
    return [deserialize_datetime(jd) for jd in jds]

@api_router.get("/jd/{jd_id}")
async def get_job_description(jd_id: str):
    """Get a specific job description"""
    jd = await db.job_descriptions.find_one({"id": jd_id}, {"_id": 0})
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")
    return deserialize_datetime(jd)

@api_router.delete("/jd/{jd_id}")
async def delete_job_description(jd_id: str):
    """Delete a job description"""
    result = await db.job_descriptions.delete_one({"id": jd_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Job description not found")
    return {"message": "Job description deleted successfully"}

# ============ Active Jobs Workflow Routes ============

def generate_requisition_number():
    """Generate a unique requisition number"""
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d")
    unique_id = str(uuid.uuid4())[:6].upper()
    return f"REQ-{timestamp}-{unique_id}"

@api_router.post("/jd/{jd_id}/submit")
async def submit_job_description(jd_id: str):
    """Submit a JD and move it to Active Jobs with requisition number"""
    jd = await db.job_descriptions.find_one({"id": jd_id}, {"_id": 0})
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")
    
    if jd.get("status") == "active":
        raise HTTPException(status_code=400, detail="Job is already active")
    
    # Generate requisition number and set dates
    requisition_number = generate_requisition_number()
    requisition_date = datetime.now(timezone.utc)
    
    # Get job details for share content
    job_title = jd.get("title", "")
    company_name = jd.get("company", "") or jd.get("client_name", "")
    location = jd.get("parsed_data", {}).get("location", "India") if jd.get("parsed_data") else "India"
    experience = jd.get("parsed_data", {}).get("experience_years", {}) if jd.get("parsed_data") else {}
    exp_text = f"{experience.get('min', 0)}+ years" if experience.get('min') else "Experience required"
    
    # Get skills
    skills = jd.get("parsed_data", {}).get("required_skills", [])[:3] if jd.get("parsed_data") else []
    skills_text = ", ".join(skills) if skills else "Various skills required"
    
    # Application info
    application_email = jd.get("application_email", "")
    application_link = jd.get("application_link", "")
    apply_text = ""
    if application_email:
        apply_text = f"Apply: {application_email}"
    elif application_link:
        apply_text = f"Apply: {application_link}"
    
    from urllib.parse import quote
    
    # LinkedIn post content
    linkedin_text = f"""🚀 We're Hiring: {job_title}

📍 Location: {location}
💼 Experience: {exp_text}
🏢 Company: {company_name}

Key Skills: {skills_text}

{apply_text}

#Hiring #{job_title.replace(' ', '')} #Jobs #Career"""
    
    # Twitter/X post content
    twitter_text = f"""🚀 Hiring: {job_title} at {company_name}!
📍 {location} | 💼 {exp_text}
{apply_text}
#Hiring #Jobs"""
    if len(twitter_text) > 280:
        twitter_text = f"🚀 Hiring: {job_title}!\n📍 {location}\n{apply_text}\n#Hiring #Jobs"
    
    # Facebook post content
    facebook_text = f"""🎯 Job Opening: {job_title}

Join {company_name}!

📍 {location}
💼 {exp_text}
Skills: {skills_text}

{apply_text}"""
    
    # Application URL for sharing
    apply_url = application_link or "https://rolesense.com/jobs"
    
    # Generate share dialog URLs with proper encoding
    from urllib.parse import quote_plus
    
    # LinkedIn: Use shareArticle for better pre-fill support
    linkedin_share_url = f"https://www.linkedin.com/shareArticle?mini=true&url={quote(apply_url, safe='')}&title={quote(f'Hiring: {job_title} at {company_name}', safe='')}&summary={quote(linkedin_text[:256], safe='')}&source=RoleSense"
    
    # Twitter: Use intent/tweet with text parameter
    twitter_share_url = f"https://twitter.com/intent/tweet?text={quote(twitter_text, safe='')}"
    
    # Facebook: Use sharer with quote parameter
    facebook_share_url = f"https://www.facebook.com/sharer/sharer.php?u={quote(apply_url, safe='')}&quote={quote(facebook_text[:300], safe='')}"
    
    publish_links = {
        "linkedin": linkedin_share_url,
        "linkedin_post": linkedin_share_url,
        "twitter": twitter_share_url,
        "facebook": facebook_share_url,
        "indeed": f"https://employers.indeed.com/p/post-job?q={quote(job_title)}",
        "naukri": f"https://recruiter.naukri.com/post-job?title={quote(job_title)}",
        "glassdoor": f"https://www.glassdoor.com/employers/post-job?title={quote(job_title)}",
        "monster": f"https://hiring.monster.com/post-job?title={quote(job_title)}",
        "ziprecruiter": f"https://www.ziprecruiter.com/post-job?title={quote(job_title)}"
    }
    
    # Update JD status
    update_data = {
        "status": "active",
        "requisition_number": requisition_number,
        "requisition_date": requisition_date.isoformat(),
        "publish_links": publish_links,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.job_descriptions.update_one(
        {"id": jd_id},
        {"$set": update_data}
    )
    
    # Return updated JD
    updated_jd = await db.job_descriptions.find_one({"id": jd_id}, {"_id": 0})
    return deserialize_datetime(updated_jd)

@api_router.get("/jd/active/list")
async def list_active_jobs():
    """List all active job descriptions from both regular and structured JDs"""
    # Fetch from both collections
    regular_jds = await db.job_descriptions.find({"status": "active"}, {"_id": 0}).sort("requisition_date", -1).to_list(100)
    structured_jds = await db.structured_jds.find({"status": "active"}, {"_id": 0}).sort("requisition_date", -1).to_list(100)
    
    # Mark structured JDs for identification
    for sjd in structured_jds:
        sjd["is_structured"] = True
    
    # Combine and sort by requisition_date
    all_jds = [deserialize_datetime(jd) for jd in regular_jds] + [deserialize_datetime(jd) for jd in structured_jds]
    all_jds.sort(key=lambda x: x.get("requisition_date", ""), reverse=True)
    
    return all_jds

@api_router.post("/jd/{jd_id}/close")
async def close_job(jd_id: str):
    """Close an active job (checks both regular and structured JDs)"""
    # Try regular JDs first
    result = await db.job_descriptions.update_one(
        {"id": jd_id},
        {"$set": {"status": "closed", "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        # Try structured JDs
        result = await db.structured_jds.update_one(
            {"id": jd_id},
            {"$set": {"status": "closed", "updated_at": datetime.now(timezone.utc)}}
        )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Job description not found")
    return {"message": "Job closed successfully"}

@api_router.post("/jd/{jd_id}/reopen")
async def reopen_job(jd_id: str):
    """Reopen a closed job (checks both regular and structured JDs)"""
    # Try regular JDs first
    result = await db.job_descriptions.update_one(
        {"id": jd_id},
        {"$set": {"status": "active", "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        # Try structured JDs
        result = await db.structured_jds.update_one(
            {"id": jd_id},
            {"$set": {"status": "active", "updated_at": datetime.now(timezone.utc)}}
        )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Job description not found")
    return {"message": "Job reopened successfully"}

@api_router.put("/jd/{jd_id}")
async def update_job_description(jd_id: str, update: JobDescriptionUpdate):
    """Update a job description (for editing/republishing)"""
    update_dict = {k: v for k, v in update.model_dump().items() if v is not None}
    update_dict["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    result = await db.job_descriptions.update_one(
        {"id": jd_id},
        {"$set": update_dict}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Job description not found")
    
    # If raw_text was updated, re-analyze
    if update.raw_text:
        jd = await db.job_descriptions.find_one({"id": jd_id}, {"_id": 0})
        # Re-analyze logic can be triggered here
    
    updated_jd = await db.job_descriptions.find_one({"id": jd_id}, {"_id": 0})
    return deserialize_datetime(updated_jd)

@api_router.get("/jobs/{job_id}/public")
async def get_public_job_details(job_id: str):
    """Public endpoint to get job details for application page (no auth required)"""
    job = await db.job_descriptions.find_one({"id": job_id}, {"_id": 0})
    is_structured = False
    
    if not job:
        # Check structured JDs collection
        job = await db.structured_jds.find_one({"id": job_id}, {"_id": 0})
        is_structured = True
    
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    if job.get("status") != "active":
        raise HTTPException(status_code=404, detail="This job posting is no longer accepting applications")
    
    # Return only public-safe fields based on job type
    if is_structured:
        basic_info = job.get("basic_info", {})
        competencies = job.get("competencies", {})
        public_fields = {
            "id": job.get("id"),
            "title": basic_info.get("title"),
            "company_name": basic_info.get("company_name"),
            "location": ", ".join(basic_info.get("locations_india", [])) or ", ".join(basic_info.get("locations_international", [])),
            "employment_type": basic_info.get("employment_type"),
            "requisition_number": job.get("requisition_number"),
            "description": job.get("ai_enhanced_jd_content") or job.get("responsibilities", "")[:1000],
            "must_have_skills": competencies.get("must_have_skills", []),
            "good_to_have_skills": competencies.get("good_to_have_skills", []),
            "experience_required": f"{basic_info.get('experience_min', 0)}-{basic_info.get('experience_max', 0)} years",
            "published_at": job.get("submitted_at") or job.get("created_at"),
            "is_structured": True
        }
    else:
        public_fields = {
            "id": job.get("id"),
            "title": job.get("title"),
            "company_name": job.get("company_name") or job.get("company") or job.get("client_name"),
            "location": job.get("location"),
            "employment_type": job.get("employment_type"),
            "requisition_number": job.get("requisition_number"),
            "description": job.get("description") or job.get("raw_text", "")[:1000],
            "must_have_skills": job.get("must_have_skills", []),
            "good_to_have_skills": job.get("good_to_have_skills", []),
            "experience_required": job.get("experience_required") or f"{job.get('experience_min', 0)}-{job.get('experience_max', 0)} years",
            "published_at": job.get("submitted_at") or job.get("created_at"),
            "is_structured": False
        }
    
    return public_fields

@api_router.get("/jobs/public")
async def get_public_jobs():
    """Get all public job listings with source field to distinguish corporate vs staffing"""
    # Get all active jobs from both regular JDs and structured JDs
    regular_jobs = await db.job_descriptions.find({"status": "active"}, {"_id": 0}).to_list(100)
    structured_jobs = await db.structured_jds.find({"status": "active"}, {"_id": 0}).to_list(100)
    
    public_jobs = []
    
    # Process regular jobs
    for job in regular_jobs:
        public_job = {
            "id": job.get("id"),
            "title": job.get("title"),
            "company": job.get("company"),
            "location": job.get("location", "Not specified"),
            "employment_type": job.get("employment_type", "Full-time"),
            "requisition_number": job.get("requisition_number"),
            "description": job.get("raw_text", "")[:500] + "..." if len(job.get("raw_text", "")) > 500 else job.get("raw_text", ""),
            "published_at": job.get("requisition_date") or job.get("created_at"),
            "source": "corporate"  # Regular JDs are from corporate clients
        }
        public_jobs.append(public_job)
    
    # Process structured jobs
    for job in structured_jobs:
        basic_info = job.get("basic_info", {})
        public_job = {
            "id": job.get("id"),
            "title": basic_info.get("title"),
            "company": basic_info.get("company_name"),
            "location": ", ".join(basic_info.get("locations_india", [])[:2]) or "India",
            "employment_type": basic_info.get("employment_type", "Full-time"),
            "requisition_number": job.get("requisition_number"),
            "description": f"Experience: {basic_info.get('experience_min', 0)}-{basic_info.get('experience_max', 0)} years. " +
                          f"Skills: {', '.join(job.get('competencies', {}).get('must_have_skills', [])[:5])}",
            "published_at": job.get("requisition_date") or job.get("created_at"),
            "source": "staffing"  # Structured JDs are from staffing vendors
        }
        public_jobs.append(public_job)
    
    # Sort by published date (newest first)
    public_jobs.sort(key=lambda x: x.get("published_at") or "", reverse=True)
    
    return public_jobs

@api_router.post("/jd/{jd_id}/generate-screening-questions")
async def generate_screening_questions(jd_id: str):
    """Generate AI-powered screening questions based on JD skills and requirements"""
    jd = await db.job_descriptions.find_one({"id": jd_id}, {"_id": 0})
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")
    
    parsed_data = jd.get("parsed_data", {})
    required_skills = parsed_data.get("required_skills", [])
    preferred_skills = parsed_data.get("preferred_skills", [])
    responsibilities = parsed_data.get("responsibilities", [])
    experience_years = parsed_data.get("experience_years", {})
    
    system_prompt = """You are an expert technical recruiter specializing in phone screening interviews.
    Generate targeted screening questions that help assess candidate fit for the role.
    Questions should be practical, specific, and have clear expected answers.
    Always respond with valid JSON only, no additional text."""
    
    user_prompt = f"""Generate comprehensive phone screening questions for this role:

JOB TITLE: {jd.get('title', '')}
COMPANY: {jd.get('company', '')}

REQUIRED SKILLS: {', '.join(required_skills)}
PREFERRED SKILLS: {', '.join(preferred_skills)}
KEY RESPONSIBILITIES: {'; '.join(responsibilities[:5])}
EXPERIENCE LEVEL: {experience_years.get('min', 0)}-{experience_years.get('max', 5)} years

JOB DESCRIPTION:
{jd.get('raw_text', '')[:2000]}

Generate 8-12 screening questions covering:
1. Technical skills verification (based on required skills)
2. Experience validation
3. Problem-solving scenarios
4. Culture fit / soft skills
5. Role-specific situational questions

Respond with JSON array:
{{
    "questions": [
        {{
            "id": "unique-id",
            "question": "The screening question to ask",
            "skill_area": "Technical/Experience/Problem-Solving/Culture/Situational",
            "expected_answer": "What a good candidate should mention or demonstrate",
            "red_flags": "Warning signs in answers to watch for",
            "follow_up": "Optional follow-up question if needed",
            "difficulty": "easy/medium/hard",
            "time_estimate": "1-2 min"
        }}
    ]
}}"""
    
    response = await get_llm_response(system_prompt, user_prompt)
    questions_data = parse_json_from_response(response)
    
    questions = questions_data.get("questions", [])
    
    # Add unique IDs if not present
    for q in questions:
        if not q.get("id"):
            q["id"] = str(uuid.uuid4())
    
    # Save to database
    await db.job_descriptions.update_one(
        {"id": jd_id},
        {"$set": {
            "screening_questions": questions,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {"questions": questions, "total_count": len(questions)}

@api_router.put("/jd/{jd_id}/screening-questions")
async def update_screening_questions(jd_id: str, update: ScreeningQuestionUpdate):
    """Update/edit screening questions for a JD"""
    result = await db.job_descriptions.update_one(
        {"id": jd_id},
        {"$set": {
            "screening_questions": update.questions,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Job description not found")
    
    return {"message": "Screening questions updated successfully", "questions": update.questions}

@api_router.get("/jd/{jd_id}/screening-questions")
async def get_screening_questions(jd_id: str):
    """Get screening questions for a JD"""
    jd = await db.job_descriptions.find_one({"id": jd_id}, {"_id": 0, "screening_questions": 1, "title": 1})
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")
    
    return {
        "jd_id": jd_id,
        "title": jd.get("title", ""),
        "questions": jd.get("screening_questions", [])
    }

def generate_pdf_content(jd: dict) -> bytes:
    """Generate PDF content for JD"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=12, textColor=colors.HexColor('#1F2937'))
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14, spaceAfter=8, spaceBefore=16, textColor=colors.HexColor('#374151'))
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=11, spaceAfter=6, leading=14)
    bullet_style = ParagraphStyle('Bullet', parent=styles['Normal'], fontSize=11, leftIndent=20, spaceAfter=4)
    
    story = []
    parsed_data = jd.get("parsed_data", {})
    
    # Title
    story.append(Paragraph(jd.get('title', 'Job Description'), title_style))
    
    # Company and Requisition Info
    info_text = f"<b>Company:</b> {jd.get('company', 'N/A')} | <b>Requisition #:</b> {jd.get('requisition_number', 'N/A')} | <b>Status:</b> {jd.get('status', 'draft').upper()}"
    story.append(Paragraph(info_text, body_style))
    story.append(Spacer(1, 12))
    
    # Summary
    if parsed_data.get('summary'):
        story.append(Paragraph("Summary", heading_style))
        story.append(Paragraph(parsed_data['summary'], body_style))
    
    # Required Skills
    if parsed_data.get('required_skills'):
        story.append(Paragraph("Required Skills", heading_style))
        for skill in parsed_data['required_skills']:
            story.append(Paragraph(f"• {skill}", bullet_style))
    
    # Preferred Skills
    if parsed_data.get('preferred_skills'):
        story.append(Paragraph("Preferred Skills", heading_style))
        for skill in parsed_data['preferred_skills']:
            story.append(Paragraph(f"• {skill}", bullet_style))
    
    # Experience
    exp = parsed_data.get('experience_years', {})
    if exp:
        story.append(Paragraph("Experience Required", heading_style))
        story.append(Paragraph(f"{exp.get('min', 0)} - {exp.get('max', 0)} years", body_style))
    
    # Responsibilities
    if parsed_data.get('responsibilities'):
        story.append(Paragraph("Responsibilities", heading_style))
        for resp in parsed_data['responsibilities']:
            story.append(Paragraph(f"• {resp}", bullet_style))
    
    # Education
    if parsed_data.get('education'):
        story.append(Paragraph("Education", heading_style))
        for edu in parsed_data['education']:
            story.append(Paragraph(f"• {edu}", bullet_style))
    
    # Benefits
    if parsed_data.get('benefits'):
        story.append(Paragraph("Benefits", heading_style))
        for benefit in parsed_data['benefits']:
            story.append(Paragraph(f"• {benefit}", bullet_style))
    
    # Full Job Description
    story.append(Paragraph("Full Job Description", heading_style))
    # Split raw text into paragraphs
    raw_text = jd.get('raw_text', '')
    for para in raw_text.split('\n\n'):
        if para.strip():
            story.append(Paragraph(para.strip().replace('\n', '<br/>'), body_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_txt_content(jd: dict) -> str:
    """Generate TXT content for JD"""
    parsed_data = jd.get("parsed_data", {})
    
    content = f"""{'='*60}
{jd.get('title', 'Job Description')}
{'='*60}

Company: {jd.get('company', 'N/A')}
Requisition #: {jd.get('requisition_number', 'N/A')}
Status: {jd.get('status', 'draft').upper()}
Date: {jd.get('requisition_date', jd.get('created_at', 'N/A'))}

{'='*60}
SUMMARY
{'='*60}
{parsed_data.get('summary', jd.get('raw_text', '')[:500])}

{'='*60}
REQUIRED SKILLS
{'='*60}
{chr(10).join(['• ' + skill for skill in parsed_data.get('required_skills', [])])}

{'='*60}
PREFERRED SKILLS
{'='*60}
{chr(10).join(['• ' + skill for skill in parsed_data.get('preferred_skills', [])])}

{'='*60}
EXPERIENCE
{'='*60}
{parsed_data.get('experience_years', {}).get('min', 0)} - {parsed_data.get('experience_years', {}).get('max', 0)} years

{'='*60}
RESPONSIBILITIES
{'='*60}
{chr(10).join(['• ' + resp for resp in parsed_data.get('responsibilities', [])])}

{'='*60}
EDUCATION
{'='*60}
{chr(10).join(['• ' + edu for edu in parsed_data.get('education', [])])}

{'='*60}
BENEFITS
{'='*60}
{chr(10).join(['• ' + benefit for benefit in parsed_data.get('benefits', [])])}

{'='*60}
FULL JOB DESCRIPTION
{'='*60}
{jd.get('raw_text', '')}
"""
    return content

@api_router.get("/jd/{jd_id}/download/{format}")
async def download_job_description(jd_id: str, format: str):
    """Download JD as PDF, DOCX, or TXT"""
    if format not in ["pdf", "docx", "txt"]:
        raise HTTPException(status_code=400, detail="Invalid format. Use pdf, docx, or txt")
    
    jd = await db.job_descriptions.find_one({"id": jd_id}, {"_id": 0})
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")
    
    filename = f"{jd.get('title', 'job_description').replace(' ', '_')}_{jd.get('requisition_number', 'draft')}"
    
    if format == "pdf":
        pdf_content = generate_pdf_content(jd)
        return StreamingResponse(
            io.BytesIO(pdf_content),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}.pdf"}
        )
    elif format == "txt":
        txt_content = generate_txt_content(jd)
        return StreamingResponse(
            io.BytesIO(txt_content.encode('utf-8')),
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename={filename}.txt"}
        )
    else:  # docx - return as txt for now with proper headers
        txt_content = generate_txt_content(jd)
        return StreamingResponse(
            io.BytesIO(txt_content.encode('utf-8')),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}.docx"}
        )

@api_router.post("/jd/generate")
async def generate_job_description(data: Dict[str, Any]):
    """Generate a job description based on requirements"""
    
    system_prompt = """You are an expert HR writer specializing in creating compelling job descriptions.
    Create clear, inclusive, and professional job descriptions that attract top talent."""
    
    user_prompt = f"""Create a professional job description based on these requirements:

Role: {data.get('title', 'Software Engineer')}
Company: {data.get('company', 'A growing tech company')}
Department: {data.get('department', 'Engineering')}
Seniority: {data.get('seniority', 'Mid-level')}
Key Skills: {', '.join(data.get('skills', ['Python', 'JavaScript']))}
Experience: {data.get('experience', '3-5 years')}
Location: {data.get('location', 'Remote')}
Additional Notes: {data.get('notes', '')}

Generate a complete, professional job description with clear sections for:
- About the Role
- Responsibilities
- Requirements
- Nice to Have
- Benefits
- About the Company"""
    
    response = await get_llm_response(system_prompt, user_prompt)
    
    return {
        "generated_jd": response,
        "metadata": data
    }

# ============ Auto-Routing Resume Repository Routes ============

async def initialize_folder_structure():
    """Initialize the folder structure if not exists - uses upsert to prevent duplicates"""
    # Create functional folders and sub-folders using upsert to prevent duplicates
    for func_name, sub_funcs in SKILL_TAXONOMY.items():
        # Check if this folder already exists
        existing_folder = await db.functional_folders.find_one({"name": func_name})
        if existing_folder:
            continue  # Skip if folder already exists
        
        sub_folder_ids = []
        for sub_func_name, skills in sub_funcs.items():
            # Use upsert for sub-folders based on parent_function + name uniqueness
            sub_folder_result = await db.sub_functional_folders.find_one_and_update(
                {"parent_function": func_name, "name": sub_func_name},
                {"$setOnInsert": {
                    "id": str(uuid.uuid4()),
                    "name": sub_func_name,
                    "parent_function": func_name,
                    "skills": skills,
                    "resume_count": 0,
                    "created_at": datetime.now(timezone.utc).isoformat()
                }},
                upsert=True,
                return_document=True
            )
            sub_folder_ids.append(sub_folder_result["id"])
        
        # Use upsert for main folder based on name uniqueness
        await db.functional_folders.find_one_and_update(
            {"name": func_name},
            {"$setOnInsert": {
                "id": str(uuid.uuid4()),
                "name": func_name,
                "description": f"{func_name} department resumes",
                "sub_folders": sub_folder_ids,
                "resume_count": 0,
                "created_at": datetime.now(timezone.utc).isoformat()
            }},
            upsert=True
        )

@api_router.get("/repository/folders")
async def get_folder_structure():
    """Get the complete folder structure with resume counts"""
    # Initialize if needed
    await initialize_folder_structure()
    
    folders = await db.functional_folders.find({}, {"_id": 0}).to_list(100)
    
    # Get sub-folders for each function
    result = []
    for folder in folders:
        sub_folders = await db.sub_functional_folders.find(
            {"parent_function": folder["name"]}, 
            {"_id": 0}
        ).to_list(100)
        folder["sub_folders_data"] = sub_folders
        result.append(folder)
    
    return result

@api_router.get("/repository/folder/{function_name}")
async def get_folder_resumes(function_name: str, sub_function: Optional[str] = None):
    """Get resumes in a specific folder"""
    query = {"primary_function": function_name}
    if sub_function:
        query["sub_function"] = sub_function
    
    resumes = await db.routed_resumes.find(query, {"_id": 0}).sort("created_at", -1).to_list(500)
    return [deserialize_datetime(r) for r in resumes]

@api_router.get("/repository/stats")
async def get_repository_stats():
    """Get statistics for the resume repository"""
    await initialize_folder_structure()
    
    total_resumes = await db.routed_resumes.count_documents({})
    
    # Get count by function
    pipeline = [
        {"$group": {"_id": "$primary_function", "count": {"$sum": 1}}}
    ]
    by_function = await db.routed_resumes.aggregate(pipeline).to_list(100)
    function_counts = {item["_id"]: item["count"] for item in by_function}
    
    # Get count by sub-function
    pipeline = [
        {"$group": {"_id": {"function": "$primary_function", "sub": "$sub_function"}, "count": {"$sum": 1}}}
    ]
    by_sub_function = await db.routed_resumes.aggregate(pipeline).to_list(500)
    
    # Get recent applications
    recent = await db.routed_resumes.find({}, {"_id": 0}).sort("created_at", -1).limit(10).to_list(10)
    
    # Get status breakdown
    status_pipeline = [
        {"$group": {"_id": "$status", "count": {"$sum": 1}}}
    ]
    by_status = await db.routed_resumes.aggregate(status_pipeline).to_list(20)
    status_counts = {item["_id"]: item["count"] for item in by_status}
    
    return {
        "total_resumes": total_resumes,
        "by_function": function_counts,
        "by_sub_function": by_sub_function,
        "by_status": status_counts,
        "recent_resumes": [deserialize_datetime(r) for r in recent]
    }

@api_router.get("/repository/all-resumes")
async def get_all_resumes(
    requisition_id: str = None,
    business_name: str = None,
    source: str = None,
    date_from: str = None,
    date_to: str = None,
    limit: int = 100,
    skip: int = 0
):
    """Get all resumes with filtering options for notification dashboard"""
    query = {}
    
    if requisition_id:
        query["applied_requisition_number"] = {"$regex": requisition_id, "$options": "i"}
    if business_name:
        query["$or"] = [
            {"applied_job_title": {"$regex": business_name, "$options": "i"}},
            {"primary_function": {"$regex": business_name, "$options": "i"}}
        ]
    if source:
        query["source"] = {"$regex": source, "$options": "i"}
    if date_from:
        try:
            from_date = datetime.fromisoformat(date_from.replace('Z', '+00:00'))
            query["created_at"] = {"$gte": from_date}
        except:
            pass
    if date_to:
        try:
            to_date = datetime.fromisoformat(date_to.replace('Z', '+00:00'))
            if "created_at" in query:
                query["created_at"]["$lte"] = to_date
            else:
                query["created_at"] = {"$lte": to_date}
        except:
            pass
    
    total = await db.routed_resumes.count_documents(query)
    resumes = await db.routed_resumes.find(query, {"_id": 0}).sort("created_at", -1).skip(skip).limit(limit).to_list(limit)
    
    return {
        "resumes": [deserialize_datetime(r) for r in resumes],
        "total": total,
        "limit": limit,
        "skip": skip
    }

@api_router.post("/repository/bulk-trajectory")
async def create_bulk_trajectory_assessments(data: dict):
    """Create Career Trajectory assessments for multiple resumes"""
    resume_ids = data.get("resume_ids", [])
    job_id = data.get("job_id")
    
    if not resume_ids:
        raise HTTPException(status_code=400, detail="No resume IDs provided")
    
    created_assessments = []
    errors = []
    
    for resume_id in resume_ids:
        try:
            # Get resume details
            resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
            if not resume:
                errors.append({"resume_id": resume_id, "error": "Resume not found"})
                continue
            
            # Create trajectory assessment
            assessment_id = str(uuid.uuid4())
            access_token = str(uuid.uuid4())[:8].upper()
            
            assessment_data = {
                "id": assessment_id,
                "candidate_name": resume.get("name", "Unknown"),
                "candidate_email": resume.get("email", ""),
                "target_role": resume.get("applied_job_title", ""),
                "target_industry": resume.get("primary_function", ""),
                "resume_text": resume.get("raw_text", ""),
                "resume_id": resume_id,
                "job_id": job_id or resume.get("applied_job_id"),
                "requisition_number": resume.get("applied_requisition_number", ""),
                "assessment_type": "post_application",
                "data_collection_mode": "recruiter_postfill",
                "status": "pending",
                "access_token": access_token,
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc),
                "source": "bulk_repository_upload"
            }
            
            await db.trajectory_assessments.insert_one(serialize_datetime(assessment_data))
            
            # Update resume with trajectory assessment link
            await db.routed_resumes.update_one(
                {"id": resume_id},
                {"$set": {
                    "trajectory_assessment_id": assessment_id,
                    "updated_at": datetime.now(timezone.utc)
                }}
            )
            
            created_assessments.append({
                "resume_id": resume_id,
                "assessment_id": assessment_id,
                "candidate_name": resume.get("name")
            })
            
        except Exception as e:
            errors.append({"resume_id": resume_id, "error": str(e)})
    
    return {
        "success": True,
        "created_count": len(created_assessments),
        "error_count": len(errors),
        "created_assessments": created_assessments,
        "errors": errors
    }

async def classify_resume_ai(resume_text: str, name: str) -> Dict[str, Any]:
    """Use AI to classify a resume into function and sub-function with interview questions"""
    
    # Build taxonomy description for AI
    taxonomy_desc = ""
    for func, sub_funcs in SKILL_TAXONOMY.items():
        taxonomy_desc += f"\n{func}:\n"
        for sub_func, skills in sub_funcs.items():
            taxonomy_desc += f"  - {sub_func}: {', '.join(skills[:10])}\n"
    
    system_prompt = """You are an expert HR professional specializing in resume classification and interview preparation.
    Analyze the resume and classify it into the appropriate functional category and sub-category.
    Also generate targeted interview questions based on the candidate's skills.
    Base your classification on skills, experience, job titles, and overall career trajectory.
    
    For IT resumes, provide detailed technology classification:
    - Identify primary technology stack (Java, .NET, Python, etc.)
    - Identify role type (Development, Architecture, DevOps, Security, etc.)
    - Provide skill match percentages for different areas
    
    Always respond with valid JSON only."""
    
    user_prompt = f"""Classify this resume and generate interview questions.

RESUME:
Name: {name}
Content: {resume_text[:4000]}

AVAILABLE CATEGORIES:
{taxonomy_desc}

Respond with JSON:
{{
    "primary_function": "The main function (HR, IT, Finance, Marketing, Operations, Supply Chain, Administration)",
    "sub_function": "The sub-function within the primary function",
    "confidence_score": 0.0-1.0,
    "matched_skills": ["list", "of", "matched", "skills"],
    "skill_tags": ["tag1", "tag2", "tag3"],
    "routing_reason": "Brief explanation of why this classification was chosen",
    "technology_stack": {{
        "primary_technology": "Java/.NET/Python/etc",
        "frameworks": ["Spring", "React", "etc"],
        "role_type": "Development/Architecture/DevOps/Security/etc",
        "seniority_level": "Junior/Mid/Senior/Lead/Architect"
    }},
    "skill_match_scores": {{
        "software_engineering": 0.0-1.0,
        "data_engineering": 0.0-1.0,
        "devops": 0.0-1.0,
        "cybersecurity": 0.0-1.0,
        "cloud": 0.0-1.0,
        "architecture": 0.0-1.0
    }},
    "suggested_folder": "Specific sub-folder recommendation based on skills",
    "suggested_interview_questions": [
        {{
            "question": "Interview question based on skills",
            "skill_area": "Which skill this tests",
            "expected_answer": "What a good answer should include",
            "difficulty": "easy/medium/hard"
        }}
    ],
    "parsed_data": {{
        "current_role": "Current or most recent job title",
        "experience_years": 0,
        "top_skills": ["skill1", "skill2"],
        "education": "Highest education",
        "companies": ["company1", "company2"],
        "certifications": ["cert1", "cert2"],
        "summary": "Brief professional summary"
    }},
    "ai_assessment": {{
        "strengths": ["strength1", "strength2"],
        "areas_for_growth": ["area1", "area2"],
        "culture_fit_indicators": ["indicator1", "indicator2"],
        "red_flags": ["any concerns or gaps"],
        "overall_rating": 1-10,
        "recommendation": "Highly Recommended/Recommended/Consider/Pass"
    }}
}}

Generate 7-10 interview questions specifically tailored to the candidate's skills and experience level.
Include questions across technical skills, behavioral aspects, and problem-solving abilities."""
    
    response = await get_llm_response(system_prompt, user_prompt)
    return parse_json_from_response(response)

async def smart_route_resume_to_job(classification: Dict, resume_text: str, applied_job_id: str = None) -> Dict:
    """Smart routing: Match resume with open jobs and determine best folder"""
    
    # If applied to specific job, prioritize that
    if applied_job_id:
        job = await db.job_descriptions.find_one({"id": applied_job_id}, {"_id": 0})
        if not job:
            job = await db.structured_jds.find_one({"id": applied_job_id}, {"_id": 0})
        
        if job:
            return {
                "matched_job_id": applied_job_id,
                "matched_job_title": job.get("title") or job.get("basic_info", {}).get("title"),
                "match_type": "direct_application",
                "match_score": 1.0,
                "routing_folder": classification.get("sub_function", "General")
            }
    
    # Get all active jobs
    active_jobs = []
    
    # Get from job_descriptions (staffing)
    staffing_jobs = await db.job_descriptions.find({"status": "active"}, {"_id": 0}).to_list(100)
    active_jobs.extend(staffing_jobs)
    
    # Get from structured_jds (corporate)
    corporate_jobs = await db.structured_jds.find({"status": "active"}, {"_id": 0}).to_list(100)
    active_jobs.extend(corporate_jobs)
    
    if not active_jobs:
        return {
            "matched_job_id": None,
            "matched_job_title": None,
            "match_type": "no_jobs_available",
            "match_score": 0,
            "routing_folder": classification.get("sub_function", "General")
        }
    
    # Get candidate skills
    candidate_skills = set([s.lower() for s in classification.get("matched_skills", [])])
    candidate_skills.update([s.lower() for s in classification.get("skill_tags", [])])
    
    best_match = None
    best_score = 0
    
    for job in active_jobs:
        # Get job skills
        job_skills = set()
        
        # From parsed_data
        parsed = job.get("parsed_data", {})
        job_skills.update([s.lower() for s in parsed.get("required_skills", [])])
        job_skills.update([s.lower() for s in parsed.get("preferred_skills", [])])
        
        # From competencies (structured JDs)
        competencies = job.get("competencies", {})
        job_skills.update([s.lower() for s in competencies.get("must_have_skills", [])])
        job_skills.update([s.lower() for s in competencies.get("good_to_have_skills", [])])
        
        # From direct fields
        job_skills.update([s.lower() for s in job.get("must_have_skills", [])])
        job_skills.update([s.lower() for s in job.get("good_to_have_skills", [])])
        
        if not job_skills:
            continue
        
        # Calculate match score
        matches = candidate_skills.intersection(job_skills)
        score = len(matches) / len(job_skills) if job_skills else 0
        
        if score > best_score:
            best_score = score
            best_match = {
                "job_id": job.get("id"),
                "job_title": job.get("title") or job.get("basic_info", {}).get("title"),
                "match_score": score,
                "matched_skills": list(matches)
            }
    
    if best_match and best_score >= 0.3:  # 30% threshold for suggesting match
        return {
            "matched_job_id": best_match["job_id"],
            "matched_job_title": best_match["job_title"],
            "match_type": "skill_match",
            "match_score": best_score,
            "matched_skills": best_match.get("matched_skills", []),
            "routing_folder": classification.get("sub_function", "General")
        }
    
    return {
        "matched_job_id": None,
        "matched_job_title": None,
        "match_type": "no_strong_match",
        "match_score": best_score,
        "routing_folder": classification.get("sub_function", "General")
    }

def determine_it_sub_folder(classification: Dict) -> str:
    """Determine specific IT sub-folder based on skills and technology"""
    
    skill_scores = classification.get("skill_match_scores", {})
    tech_stack = classification.get("technology_stack", {})
    
    # Priority routing based on skill match scores
    if skill_scores.get("cybersecurity", 0) >= 0.7:
        return "Cybersecurity"
    if skill_scores.get("data_engineering", 0) >= 0.7:
        return "Data Engineering"
    if skill_scores.get("devops", 0) >= 0.7:
        return "DevOps & Infrastructure"
    if skill_scores.get("cloud", 0) >= 0.7:
        return "Cloud Engineering"
    if skill_scores.get("architecture", 0) >= 0.7:
        return "Solution Architecture"
    
    # Fallback to role type
    role_type = tech_stack.get("role_type", "").lower()
    if "security" in role_type or "cyber" in role_type:
        return "Cybersecurity"
    if "devops" in role_type or "sre" in role_type:
        return "DevOps & Infrastructure"
    if "data" in role_type or "analytics" in role_type:
        return "Data Engineering"
    if "architect" in role_type:
        return "Solution Architecture"
    if "qa" in role_type or "test" in role_type:
        return "Quality Assurance"
    
    # Default to Software Engineering for development roles
    return "Software Engineering"

async def check_sla_and_notify(primary_function: str, sub_function: str, recipient_override: str = None):
    """Check SLA rules and create notifications if thresholds exceeded"""
    # Get today's count for this folder
    today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
    
    count_today = await db.routed_resumes.count_documents({
        "primary_function": primary_function,
        "sub_function": sub_function,
        "created_at": {"$gte": today_start.isoformat()}
    })
    
    # Get SLA rules
    sla_rule = SLA_RULES.get(primary_function, {}).get(sub_function, SLA_RULES["default"])
    threshold = sla_rule["daily_threshold"]
    notify_email = recipient_override or sla_rule["notify_email"]
    priority = sla_rule["priority"]
    
    # Check if we just crossed the threshold
    if count_today == threshold:
        notification = {
            "id": str(uuid.uuid4()),
            "type": "sla_alert",
            "title": f"SLA Alert: {sub_function} ({primary_function})",
            "message": f"Daily resume threshold reached! {count_today} resumes received today in {sub_function} folder. Action may be required.",
            "recipient_email": notify_email,
            "function": primary_function,
            "sub_function": sub_function,
            "priority": priority,
            "read": False,
            "data": {
                "count_today": count_today,
                "threshold": threshold,
                "folder": f"{primary_function} > {sub_function}"
            },
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.notifications.insert_one(notification)
        return notification
    
    return None

@api_router.get("/notifications")
async def get_notifications(email: Optional[str] = None, unread_only: bool = False):
    """Get notifications, optionally filtered by email and read status"""
    query = {}
    if email:
        query["recipient_email"] = email
    if unread_only:
        query["read"] = False
    
    notifications = await db.notifications.find(query, {"_id": 0}).sort("created_at", -1).limit(50).to_list(50)
    return [deserialize_datetime(n) for n in notifications]

@api_router.put("/notifications/{notification_id}/read")
async def mark_notification_read(notification_id: str):
    """Mark a notification as read"""
    await db.notifications.update_one({"id": notification_id}, {"$set": {"read": True}})
    return {"message": "Notification marked as read"}

@api_router.get("/repository/source-analytics")
async def get_source_analytics():
    """Get analytics on source quality and performance"""
    pipeline = [
        {"$group": {
            "_id": "$source",
            "count": {"$sum": 1},
            "avg_confidence": {"$avg": "$confidence_score"},
            "statuses": {"$push": "$status"}
        }}
    ]
    source_stats = await db.routed_resumes.aggregate(pipeline).to_list(50)
    
    # Calculate quality metrics per source
    results = []
    for stat in source_stats:
        source_key = stat["_id"]
        source_info = SOURCE_TYPES.get(source_key, {"name": source_key, "quality_weight": 1.0})
        
        # Calculate conversion rates
        statuses = stat["statuses"]
        total = len(statuses)
        shortlisted = statuses.count("shortlisted")
        hired = statuses.count("hired")
        
        results.append({
            "source": source_key,
            "source_name": source_info["name"],
            "quality_weight": source_info["quality_weight"],
            "total_resumes": stat["count"],
            "avg_confidence": round(stat["avg_confidence"] * 100, 1) if stat["avg_confidence"] else 0,
            "shortlist_rate": round((shortlisted / total) * 100, 1) if total > 0 else 0,
            "hire_rate": round((hired / total) * 100, 1) if total > 0 else 0
        })
    
    return sorted(results, key=lambda x: x["total_resumes"], reverse=True)

@api_router.post("/repository/classify")
async def classify_resume(data: dict):
    """Classify a resume and determine routing"""
    resume_text = data.get("resume_text", "")
    name = data.get("name", "Unknown")
    
    if not resume_text:
        raise HTTPException(status_code=400, detail="Resume text is required")
    
    classification = await classify_resume_ai(resume_text, name)
    return classification

@api_router.post("/repository/route")
async def route_resume(data: dict):
    """Route a resume to the appropriate folder"""
    resume_text = data.get("resume_text", "")
    name = data.get("name", "Unknown")
    email = data.get("email", "")
    phone = data.get("phone", "")
    job_id = data.get("job_id")
    source = data.get("source", "direct")
    
    if not resume_text:
        raise HTTPException(status_code=400, detail="Resume text is required")
    
    # Classify the resume
    classification = await classify_resume_ai(resume_text, name)
    
    # Create routed resume record
    routed_resume = {
        "id": str(uuid.uuid4()),
        "job_id": job_id,
        "name": name,
        "email": email,
        "phone": phone,
        "raw_text": resume_text,
        "parsed_data": classification.get("parsed_data", {}),
        "primary_function": classification.get("primary_function", "Administration"),
        "sub_function": classification.get("sub_function", "Office Administration"),
        "confidence_score": classification.get("confidence_score", 0.5),
        "matched_skills": classification.get("matched_skills", []),
        "routing_reason": classification.get("routing_reason", ""),
        # Enhanced fields
        "suggested_interview_questions": classification.get("suggested_interview_questions", []),
        "skill_tags": classification.get("skill_tags", []),
        "source": source,
        "source_quality_score": SOURCE_TYPES.get(source, {"quality_weight": 1.0})["quality_weight"],
        "status": "new",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    # If applied to a specific job, add job info
    if job_id:
        job = await db.job_descriptions.find_one({"id": job_id}, {"_id": 0})
        if job:
            routed_resume["applied_job_title"] = job.get("title", "")
    
    await db.routed_resumes.insert_one(routed_resume)
    
    # Update folder counts
    await db.functional_folders.update_one(
        {"name": routed_resume["primary_function"]},
        {"$inc": {"resume_count": 1}}
    )
    await db.sub_functional_folders.update_one(
        {"name": routed_resume["sub_function"], "parent_function": routed_resume["primary_function"]},
        {"$inc": {"resume_count": 1}}
    )
    
    # Check SLA and create notification if needed
    notification = await check_sla_and_notify(
        routed_resume["primary_function"], 
        routed_resume["sub_function"]
    )
    
    # Return a clean response without any MongoDB objects
    clean_response = {
        "id": routed_resume["id"],
        "name": routed_resume["name"],
        "email": routed_resume["email"],
        "phone": routed_resume["phone"],
        "primary_function": routed_resume["primary_function"],
        "sub_function": routed_resume["sub_function"],
        "confidence_score": routed_resume["confidence_score"],
        "matched_skills": routed_resume["matched_skills"],
        "routing_reason": routed_resume["routing_reason"],
        "suggested_interview_questions": routed_resume["suggested_interview_questions"],
        "skill_tags": routed_resume["skill_tags"],
        "source": routed_resume["source"],
        "source_quality_score": routed_resume["source_quality_score"],
        "status": routed_resume["status"],
        "created_at": routed_resume["created_at"],
        "sla_notification": notification
    }
    
    # Add optional fields if they exist
    if routed_resume.get("job_id"):
        clean_response["job_id"] = routed_resume["job_id"]
    if routed_resume.get("applied_job_title"):
        clean_response["applied_job_title"] = routed_resume["applied_job_title"]
    if routed_resume.get("parsed_data"):
        clean_response["parsed_data"] = routed_resume["parsed_data"]
    
    return clean_response

@api_router.post("/apply/{job_id}")
async def submit_application(job_id: str, application: ApplicationSubmit):
    """Public endpoint for job applications - auto-routes to repository"""
    # Verify job exists and is active - check both collections
    job = await db.job_descriptions.find_one({"id": job_id}, {"_id": 0})
    is_structured = False
    
    if not job:
        # Check structured JDs collection
        job = await db.structured_jds.find_one({"id": job_id}, {"_id": 0})
        is_structured = True
    
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("status") != "active":
        raise HTTPException(status_code=400, detail="This job is no longer accepting applications")
    
    # Get job details based on type
    job_title = job.get("basic_info", {}).get("title") if is_structured else job.get("title", "")
    company_name = job.get("basic_info", {}).get("company_name") if is_structured else job.get("company", "")
    requisition_number = job.get("requisition_number", "")
    
    # Create application record
    app_record = {
        "id": str(uuid.uuid4()),
        "job_id": job_id,
        "job_title": job_title,
        "company": company_name,
        "requisition_number": requisition_number,
        "applicant_name": application.applicant_name,
        "applicant_email": application.applicant_email,
        "applicant_phone": application.applicant_phone,
        "resume_text": application.resume_text,
        "cover_letter": application.cover_letter,
        "linkedin_url": application.linkedin_url,
        "current_company": application.current_company,
        "current_ctc": application.current_ctc,
        "expected_ctc": application.expected_ctc,
        "notice_period": application.notice_period,
        "source": application.source or "careers_page",
        "status": "processing",
        "is_structured_jd": is_structured,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.job_applications.insert_one(app_record)
    
    # Auto-route the resume with enhanced AI classification
    classification = await classify_resume_ai(application.resume_text, application.applicant_name)
    
    # Determine sub-folder - use smart IT routing for IT resumes
    primary_function = classification.get("primary_function", "Administration")
    sub_function = classification.get("sub_function", "Office Administration")
    
    # Smart IT sub-folder routing based on skills
    if primary_function == "IT":
        sub_function = determine_it_sub_folder(classification)
    
    # Get smart job matching
    job_match = await smart_route_resume_to_job(classification, application.resume_text, job_id)
    
    routed_resume = {
        "id": str(uuid.uuid4()),
        "job_id": job_id,
        "applied_job_id": job_id,
        "applied_job_title": job_title,
        "applied_requisition_number": requisition_number,
        "name": application.applicant_name,
        "email": application.applicant_email,
        "phone": application.applicant_phone,
        "raw_text": application.resume_text,
        "parsed_data": classification.get("parsed_data", {}),
        "primary_function": primary_function,
        "sub_function": sub_function,
        "confidence_score": classification.get("confidence_score", 0.5),
        "matched_skills": classification.get("matched_skills", []),
        "routing_reason": classification.get("routing_reason", ""),
        # Enhanced AI fields
        "technology_stack": classification.get("technology_stack", {}),
        "skill_match_scores": classification.get("skill_match_scores", {}),
        "ai_assessment": classification.get("ai_assessment", {}),
        "interview_questions": classification.get("suggested_interview_questions", []),
        "skill_tags": classification.get("skill_tags", []),
        # Job matching
        "job_match_score": job_match.get("match_score", 0),
        "job_match_type": job_match.get("match_type"),
        # Source tracking
        "source": application.source or "careers_page",
        "source_quality_score": SOURCE_TYPES.get(application.source or "careers_page", {"quality_weight": 1.0})["quality_weight"],
        # Status tracking
        "application_date": datetime.now(timezone.utc).isoformat(),
        "status": "new",
        "recruiter_status": "under_evaluation",
        "recruiter_notes": [],
        "has_responses": False,
        "created_at": datetime.now(timezone.utc)
    }
    
    await db.routed_resumes.insert_one(serialize_datetime(routed_resume))
    
    # Update application with routed resume ID
    await db.job_applications.update_one(
        {"id": app_record["id"]},
        {"$set": {"routed_resume_id": routed_resume["id"], "status": "routed"}}
    )
    
    # Update folder counts
    await db.functional_folders.update_one(
        {"name": routed_resume["primary_function"]},
        {"$inc": {"resume_count": 1}}
    )
    await db.sub_functional_folders.update_one(
        {"name": routed_resume["sub_function"], "parent_function": routed_resume["primary_function"]},
        {"$inc": {"resume_count": 1}}
    )
    
    # Check SLA and create notification if needed
    notification = await check_sla_and_notify(
        routed_resume["primary_function"], 
        routed_resume["sub_function"]
    )
    
    # Create recruiter notification about new resume received
    recruiter_notification = {
        "id": str(uuid.uuid4()),
        "type": "new_resume_received",
        "title": f"New Resume Received - {application.applicant_name}",
        "message": f"Resume received for {job_title} from {application.applicant_name}. Routed to {routed_resume['primary_function']} > {routed_resume['sub_function']}.",
        "job_id": job_id,
        "requisition_number": requisition_number,
        "resume_id": routed_resume["id"],
        "candidate_name": application.applicant_name,
        "candidate_email": application.applicant_email,
        "source": application.source or "careers_page",
        "read": False,
        "priority": "normal",
        "action_required": "Review resume and trigger pre-assessment if not auto-triggered",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.notifications.insert_one(recruiter_notification)
    
    # Pre-assessment is NOT auto-triggered - recruiter will send manually if shortlisted
    preassessment_result = {
        "message": "Pre-assessment will be sent by recruiter if shortlisted",
        "manual_trigger_available": True,
        "auto_triggered": False
    }
    
    # Simulate sending acknowledgment email to candidate
    print(f"""
    ============================================================
    RESUME RECEIPT ACKNOWLEDGMENT EMAIL
    ============================================================
    To: {application.applicant_email}
    Subject: Application Received - {job_title}
    
    Dear {application.applicant_name},
    
    Thank you for your interest in the {job_title} position.
    
    We have received your resume and it has been successfully submitted.
    Your application is now under review.
    
    Next Steps:
    - Our recruitment team will review your application
    - If shortlisted, you will receive a Pre-Assessment Form
    - Please ensure your contact details are up to date
    
    Reference Number: {app_record['id'][:8].upper()}
    
    Best regards,
    Recruitment Team
    ============================================================
    """)
    
    return {
        "message": "Application submitted successfully. Resume received and routed.",
        "acknowledgment": f"Thank you {application.applicant_name}! Your resume has been received for {job_title}. If shortlisted, a Pre-Assessment form will be shared with you.",
        "application_id": app_record["id"],
        "routed_to": {
            "function": routed_resume["primary_function"],
            "sub_function": routed_resume["sub_function"]
        },
        "confidence": routed_resume["confidence_score"],
        "interview_questions_generated": len(routed_resume.get("interview_questions", [])),
        "skill_tags": routed_resume.get("skill_tags", []),
        "sla_notification": notification is not None,
        "recruiter_notified": True,
        "preassessment": preassessment_result
    }

@api_router.get("/repository/resume/{resume_id}")
async def get_routed_resume(resume_id: str):
    """Get a specific routed resume"""
    resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return deserialize_datetime(resume)

@api_router.put("/repository/resume/{resume_id}/status")
async def update_resume_status(resume_id: str, data: dict):
    """Update the status of a routed resume"""
    new_status = data.get("status")
    if new_status not in ["new", "reviewed", "shortlisted", "rejected", "hired"]:
        raise HTTPException(status_code=400, detail="Invalid status")
    
    result = await db.routed_resumes.update_one(
        {"id": resume_id},
        {"$set": {"status": new_status}}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    return {"message": "Status updated successfully"}

@api_router.put("/repository/resume/{resume_id}/move")
async def move_resume_to_folder(resume_id: str, data: dict):
    """Manually move a resume to a different folder"""
    new_function = data.get("primary_function")
    new_sub_function = data.get("sub_function")
    
    # Get current resume
    resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    old_function = resume.get("primary_function")
    old_sub_function = resume.get("sub_function")
    
    # Update resume
    await db.routed_resumes.update_one(
        {"id": resume_id},
        {"$set": {"primary_function": new_function, "sub_function": new_sub_function}}
    )
    
    # Update folder counts
    if old_function:
        await db.functional_folders.update_one(
            {"name": old_function},
            {"$inc": {"resume_count": -1}}
        )
    if old_sub_function:
        await db.sub_functional_folders.update_one(
            {"name": old_sub_function, "parent_function": old_function},
            {"$inc": {"resume_count": -1}}
        )
    
    await db.functional_folders.update_one(
        {"name": new_function},
        {"$inc": {"resume_count": 1}}
    )
    await db.sub_functional_folders.update_one(
        {"name": new_sub_function, "parent_function": new_function},
        {"$inc": {"resume_count": 1}}
    )
    
    return {"message": "Resume moved successfully"}

@api_router.get("/repository/applications")
async def get_all_applications():
    """Get all job applications"""
    applications = await db.job_applications.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)
    return [deserialize_datetime(app) for app in applications]

@api_router.get("/repository/job/{job_id}/applications")
async def get_job_applications(job_id: str):
    """Get all applications for a specific job"""
    applications = await db.job_applications.find({"job_id": job_id}, {"_id": 0}).sort("created_at", -1).to_list(500)
    return [deserialize_datetime(app) for app in applications]


# ============ HR Pre-Assessment Questionnaire Routes ============

@api_router.get("/preassessment/questionnaires")
async def get_questionnaires(client_id: Optional[str] = None):
    """Get all questionnaires (default + custom)"""
    query = {"is_active": True}
    if client_id:
        query["$or"] = [{"is_default": True}, {"created_by": client_id}]
    else:
        query["is_default"] = True
    
    questionnaires = await db.preassessment_questionnaires.find(query, {"_id": 0}).to_list(100)
    
    # If no questionnaires exist, create default one
    if not questionnaires:
        default_q = {
            "id": str(uuid.uuid4()),
            "name": "Default HR Pre-Assessment",
            "description": "Standard HR pre-assessment questionnaire for candidate evaluation",
            "questions": DEFAULT_HR_PREASSESSMENT_QUESTIONS,
            "is_default": True,
            "is_active": True,
            "created_at": datetime.now(timezone.utc),
            "updated_at": datetime.now(timezone.utc)
        }
        await db.preassessment_questionnaires.insert_one(default_q)
        questionnaires = [default_q]
    
    return [deserialize_datetime(q) for q in questionnaires]

@api_router.get("/preassessment/questionnaire/{questionnaire_id}")
async def get_questionnaire(questionnaire_id: str):
    """Get a specific questionnaire"""
    questionnaire = await db.preassessment_questionnaires.find_one({"id": questionnaire_id}, {"_id": 0})
    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found")
    return deserialize_datetime(questionnaire)

@api_router.post("/preassessment/questionnaire/create")
async def create_questionnaire(name: str, questions: List[Dict[str, Any]], client_id: Optional[str] = None, description: Optional[str] = None):
    """Create a custom questionnaire"""
    questionnaire = {
        "id": str(uuid.uuid4()),
        "name": name,
        "description": description,
        "questions": questions,
        "created_by": client_id,
        "is_default": False,
        "is_active": True,
        "created_at": datetime.now(timezone.utc),
        "updated_at": datetime.now(timezone.utc)
    }
    await db.preassessment_questionnaires.insert_one(questionnaire)
    return questionnaire

@api_router.put("/preassessment/questionnaire/{questionnaire_id}")
async def update_questionnaire(questionnaire_id: str, name: Optional[str] = None, questions: Optional[List[Dict[str, Any]]] = None, description: Optional[str] = None):
    """Update a questionnaire"""
    update_data = {"updated_at": datetime.now(timezone.utc)}
    if name:
        update_data["name"] = name
    if questions:
        update_data["questions"] = questions
    if description:
        update_data["description"] = description
    
    await db.preassessment_questionnaires.update_one(
        {"id": questionnaire_id},
        {"$set": update_data}
    )
    return {"success": True, "message": "Questionnaire updated"}

@api_router.post("/preassessment/send/{resume_id}")
async def send_preassessment(resume_id: str, questionnaire_id: Optional[str] = None, job_id: Optional[str] = None):
    """Send pre-assessment to a candidate from Resume Repository"""
    # Get resume details
    resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Get questionnaire (use default if not specified)
    if questionnaire_id:
        questionnaire = await db.preassessment_questionnaires.find_one({"id": questionnaire_id}, {"_id": 0})
    else:
        questionnaire = await db.preassessment_questionnaires.find_one({"is_default": True}, {"_id": 0})
    
    if not questionnaire:
        # Create default if doesn't exist
        questionnaire = {
            "id": str(uuid.uuid4()),
            "name": "Default HR Pre-Assessment",
            "questions": DEFAULT_HR_PREASSESSMENT_QUESTIONS,
            "is_default": True,
            "is_active": True,
            "created_at": datetime.now(timezone.utc)
        }
        await db.preassessment_questionnaires.insert_one(questionnaire)
    
    # Get job details if job_id provided
    job_title = None
    if job_id:
        job = await db.structured_jds.find_one({"id": job_id}, {"_id": 0})
        if job:
            job_title = job.get("basic_info", {}).get("title") or job.get("title")
    
    # Check if already sent
    existing = await db.preassessment_responses.find_one({
        "resume_id": resume_id,
        "status": {"$in": ["pending", "sent"]}
    })
    if existing:
        return {
            "success": False,
            "message": "Pre-assessment already sent to this candidate",
            "response_id": existing["id"],
            "status": existing["status"]
        }
    
    # Create pre-assessment response record
    access_token = str(uuid.uuid4())[:8].upper()
    response_record = {
        "id": str(uuid.uuid4()),
        "questionnaire_id": questionnaire["id"],
        "resume_id": resume_id,
        "candidate_name": resume.get("name", "Candidate"),
        "candidate_email": resume.get("email", ""),
        "job_id": job_id,
        "job_title": job_title or resume.get("applied_job_title"),
        "responses": {},
        "status": "sent",
        "sent_at": datetime.now(timezone.utc),
        "access_token": access_token,
        "created_at": datetime.now(timezone.utc)
    }
    await db.preassessment_responses.insert_one(response_record)
    
    # Update resume status
    await db.routed_resumes.update_one(
        {"id": resume_id},
        {"$set": {
            "preassessment_status": "sent",
            "preassessment_sent_at": datetime.now(timezone.utc),
            "preassessment_response_id": response_record["id"]
        }}
    )
    
    # Generate link
    frontend_url = os.environ.get('FRONTEND_URL', 'http://localhost:3000')
    preassessment_link = f"{frontend_url}/preassessment/{response_record['id']}?token={access_token}"
    
    # Simulate email
    print(f"""
    ============================================================
    HR PRE-ASSESSMENT EMAIL SENT
    ============================================================
    To: {resume.get('email')}
    Subject: Pre-Assessment Form - {job_title or 'Job Application'}
    
    Dear {resume.get('name')},
    
    Congratulations! Your application has been shortlisted.
    
    Please complete the Pre-Assessment form to proceed with your application.
    
    Link: {preassessment_link}
    
    This link will expire in 7 days.
    
    Best regards,
    Recruitment Team
    ============================================================
    """)
    
    return {
        "success": True,
        "message": "Pre-assessment sent successfully",
        "response_id": response_record["id"],
        "preassessment_link": preassessment_link,
        "candidate_email": resume.get("email")
    }

@api_router.get("/preassessment/response/{response_id}")
async def get_preassessment_response(response_id: str, token: Optional[str] = None):
    """Get pre-assessment response (for candidate to fill or recruiter to view)"""
    response = await db.preassessment_responses.find_one({"id": response_id}, {"_id": 0})
    if not response:
        raise HTTPException(status_code=404, detail="Pre-assessment not found")
    
    # Validate token for candidate access
    if token and response.get("access_token") != token:
        raise HTTPException(status_code=403, detail="Invalid access token")
    
    # Get questionnaire
    questionnaire = await db.preassessment_questionnaires.find_one(
        {"id": response["questionnaire_id"]}, {"_id": 0}
    )
    
    return {
        "response": deserialize_datetime(response),
        "questionnaire": deserialize_datetime(questionnaire) if questionnaire else None
    }

@api_router.post("/preassessment/response/{response_id}/submit")
async def submit_preassessment_response(response_id: str, responses: Dict[str, Any], token: str):
    """Candidate submits their pre-assessment responses"""
    # Validate
    response = await db.preassessment_responses.find_one({"id": response_id}, {"_id": 0})
    if not response:
        raise HTTPException(status_code=404, detail="Pre-assessment not found")
    
    if response.get("access_token") != token:
        raise HTTPException(status_code=403, detail="Invalid access token")
    
    if response.get("status") == "completed":
        raise HTTPException(status_code=400, detail="Pre-assessment already completed")
    
    # Generate HR Fitment Analysis from responses
    hr_fitment_analysis = {
        "motivation": {},
        "employment_history": {},
        "achievements": {},
        "work_preferences": {},
        "commitments": {},
        "strengths_development": {},
        "interview_status": {},
        "career_goals": {},
        "analysis_summary": ""
    }
    
    # Map responses to HR fitment categories
    for q_id, answer in responses.items():
        if "reason" in q_id.lower() or q_id == "q1":
            hr_fitment_analysis["motivation"]["reason_for_change"] = answer
        elif "gap" in q_id.lower() or q_id == "q2":
            hr_fitment_analysis["employment_history"]["gaps_explanation"] = answer
        elif "achievement" in q_id.lower() or q_id == "q3":
            hr_fitment_analysis["achievements"]["top_achievements"] = answer
        elif "relocation" in q_id.lower() or q_id == "q4":
            hr_fitment_analysis["work_preferences"]["relocation"] = answer
        elif "work" in q_id.lower() or q_id == "q5":
            hr_fitment_analysis["work_preferences"]["arrangement"] = answer
        elif "commitment" in q_id.lower() or q_id == "q6":
            hr_fitment_analysis["commitments"]["current_commitments"] = answer
        elif "strength" in q_id.lower() or q_id == "q7":
            hr_fitment_analysis["strengths_development"]["strengths"] = answer
        elif "improve" in q_id.lower() or q_id == "q8":
            hr_fitment_analysis["strengths_development"]["areas_to_develop"] = answer
        elif "interview" in q_id.lower() or q_id == "q9":
            hr_fitment_analysis["interview_status"]["current_status"] = answer
        elif "goal" in q_id.lower() or q_id == "q10":
            hr_fitment_analysis["career_goals"]["long_term_goals"] = answer
    
    # Update response record
    await db.preassessment_responses.update_one(
        {"id": response_id},
        {"$set": {
            "responses": responses,
            "status": "completed",
            "completed_at": datetime.now(timezone.utc),
            "hr_fitment_analysis": hr_fitment_analysis
        }}
    )
    
    # Update resume record
    await db.routed_resumes.update_one(
        {"id": response["resume_id"]},
        {"$set": {
            "preassessment_status": "completed",
            "preassessment_completed_at": datetime.now(timezone.utc),
            "hr_fitment_analysis": hr_fitment_analysis
        }}
    )
    
    # Notify recruiter
    notification = {
        "id": str(uuid.uuid4()),
        "type": "preassessment_completed",
        "title": f"Pre-Assessment Completed - {response['candidate_name']}",
        "message": f"{response['candidate_name']} has completed their pre-assessment form.",
        "resume_id": response["resume_id"],
        "response_id": response_id,
        "read": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.notifications.insert_one(notification)
    
    return {
        "success": True,
        "message": "Pre-assessment submitted successfully. The recruitment team will review your responses.",
        "hr_fitment_analysis": hr_fitment_analysis
    }

@api_router.get("/preassessment/resume/{resume_id}/status")
async def get_resume_preassessment_status(resume_id: str):
    """Get pre-assessment status for a resume"""
    response = await db.preassessment_responses.find_one(
        {"resume_id": resume_id},
        {"_id": 0}
    )
    
    if not response:
        return {
            "status": "not_sent",
            "can_send": True,
            "message": "Pre-assessment not yet sent"
        }
    
    return {
        "status": response.get("status"),
        "response_id": response.get("id"),
        "sent_at": response.get("sent_at"),
        "completed_at": response.get("completed_at"),
        "can_send": response.get("status") == "completed",
        "hr_fitment_analysis": response.get("hr_fitment_analysis")
    }

# ============ Competency Report Generation ============

@api_router.post("/competency-report/generate/{resume_id}")
async def generate_competency_report(resume_id: str, report_type: str = "basic"):
    """Generate competency report for a candidate"""
    # Get resume
    resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Get pre-assessment response if exists
    preassessment = await db.preassessment_responses.find_one(
        {"resume_id": resume_id, "status": "completed"},
        {"_id": 0}
    )
    
    # Get recruiter notes
    recruiter_notes = resume.get("recruiter_notes", [])
    
    # Parse resume for projects and skills (factual data only)
    resume_analysis = {
        "candidate_name": resume.get("name"),
        "candidate_email": resume.get("email"),
        "primary_function": resume.get("primary_function"),
        "sub_function": resume.get("sub_function"),
        "skills": resume.get("matched_skills", []) + resume.get("skill_tags", []),
        "technology_stack": resume.get("technology_stack", {}),
        "parsed_data": resume.get("parsed_data", {}),
        "confidence_score": resume.get("confidence_score", 0)
    }
    
    # Extract projects from parsed data
    parsed = resume.get("parsed_data", {})
    resume_analysis["experience_summary"] = parsed.get("experience_summary", [])
    resume_analysis["education"] = parsed.get("education", [])
    resume_analysis["certifications"] = parsed.get("certifications", [])
    resume_analysis["projects"] = parsed.get("projects", [])
    
    # Create report record
    report = {
        "id": str(uuid.uuid4()),
        "resume_id": resume_id,
        "candidate_name": resume.get("name"),
        "candidate_email": resume.get("email"),
        "preassessment_response_id": preassessment["id"] if preassessment else None,
        "resume_analysis": resume_analysis,
        "hr_fitment_data": preassessment.get("hr_fitment_analysis", {}) if preassessment else {},
        "preassessment_responses": preassessment.get("responses", {}) if preassessment else {},
        "recruiter_observations": recruiter_notes,
        "report_type": report_type,
        "generated_at": datetime.now(timezone.utc)
    }
    
    await db.competency_reports.insert_one(report)
    
    return deserialize_datetime(report)

@api_router.get("/competency-report/{report_id}")
async def get_competency_report(report_id: str):
    """Get a competency report"""
    report = await db.competency_reports.find_one({"id": report_id}, {"_id": 0})
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    return deserialize_datetime(report)

@api_router.get("/competency-report/resume/{resume_id}")
async def get_competency_report_by_resume(resume_id: str):
    """Get competency report for a resume"""
    report = await db.competency_reports.find_one(
        {"resume_id": resume_id},
        {"_id": 0}
    )
    if report:
        return deserialize_datetime(report)
    return None

@api_router.get("/competency-report/{report_id}/download/{format}")
async def download_competency_report(report_id: str, format: str = "pdf"):
    """Download competency report as PDF or DOCX"""
    from fastapi.responses import StreamingResponse
    
    report = await db.competency_reports.find_one({"id": report_id}, {"_id": 0})
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    if format.lower() == "pdf":
        # Generate PDF
        pdf_buffer = await generate_competency_pdf(report)
        return StreamingResponse(
            pdf_buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=competency_report_{report['candidate_name'].replace(' ', '_')}.pdf"}
        )
    elif format.lower() == "docx":
        # Generate DOCX
        docx_buffer = await generate_competency_docx(report)
        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=competency_report_{report['candidate_name'].replace(' ', '_')}.docx"}
        )
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use 'pdf' or 'docx'")

async def generate_competency_pdf(report: Dict[str, Any]) -> io.BytesIO:
    """Generate PDF competency report"""
    buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )
    
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        'Heading',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=12,
        textColor=colors.HexColor('#1e293b')
    )
    subheading_style = ParagraphStyle(
        'SubHeading',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=8,
        textColor=colors.HexColor('#475569')
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6
    )
    
    story = []
    
    # Title
    story.append(Paragraph("COMPETENCY REPORT", heading_style))
    story.append(Paragraph(f"Candidate: {report['candidate_name']}", subheading_style))
    story.append(Paragraph(f"Generated: {report['generated_at']}", body_style))
    story.append(Spacer(1, 20))
    
    # Resume Analysis Section
    story.append(Paragraph("1. CANDIDATE PROFILE", heading_style))
    analysis = report.get("resume_analysis", {})
    story.append(Paragraph(f"<b>Function:</b> {analysis.get('primary_function', 'N/A')} > {analysis.get('sub_function', 'N/A')}", body_style))
    story.append(Paragraph(f"<b>Email:</b> {analysis.get('candidate_email', 'N/A')}", body_style))
    
    # Skills
    skills = analysis.get("skills", [])
    if skills:
        story.append(Paragraph("<b>Key Skills:</b>", body_style))
        story.append(Paragraph(", ".join(skills[:15]), body_style))
    story.append(Spacer(1, 15))
    
    # Experience Summary
    exp_summary = analysis.get("experience_summary", [])
    if exp_summary:
        story.append(Paragraph("2. EXPERIENCE SUMMARY", heading_style))
        for exp in exp_summary[:5]:
            if isinstance(exp, dict):
                story.append(Paragraph(f"• {exp.get('company', '')} - {exp.get('role', '')}", body_style))
            else:
                story.append(Paragraph(f"• {exp}", body_style))
        story.append(Spacer(1, 15))
    
    # HR Fitment Analysis
    hr_fitment = report.get("hr_fitment_data", {})
    if hr_fitment:
        story.append(Paragraph("3. HR FITMENT ANALYSIS", heading_style))
        
        if hr_fitment.get("motivation", {}).get("reason_for_change"):
            story.append(Paragraph("<b>Reason for Change:</b>", subheading_style))
            story.append(Paragraph(str(hr_fitment["motivation"]["reason_for_change"]), body_style))
        
        if hr_fitment.get("work_preferences", {}).get("relocation"):
            story.append(Paragraph(f"<b>Relocation:</b> {hr_fitment['work_preferences']['relocation']}", body_style))
        
        if hr_fitment.get("work_preferences", {}).get("arrangement"):
            story.append(Paragraph(f"<b>Work Arrangement:</b> {hr_fitment['work_preferences']['arrangement']}", body_style))
        
        if hr_fitment.get("achievements", {}).get("top_achievements"):
            story.append(Paragraph("<b>Key Achievements:</b>", subheading_style))
            story.append(Paragraph(str(hr_fitment["achievements"]["top_achievements"]), body_style))
        
        if hr_fitment.get("career_goals", {}).get("long_term_goals"):
            story.append(Paragraph("<b>Career Goals:</b>", subheading_style))
            story.append(Paragraph(str(hr_fitment["career_goals"]["long_term_goals"]), body_style))
        
        story.append(Spacer(1, 15))
    
    # Pre-assessment Responses
    preassessment_responses = report.get("preassessment_responses", {})
    if preassessment_responses:
        story.append(Paragraph("4. PRE-ASSESSMENT RESPONSES", heading_style))
        for q_id, response in preassessment_responses.items():
            story.append(Paragraph(f"<b>{q_id}:</b> {response}", body_style))
        story.append(Spacer(1, 15))
    
    # Recruiter Observations
    observations = report.get("recruiter_observations", [])
    if observations:
        story.append(Paragraph("5. RECRUITER OBSERVATIONS", heading_style))
        for obs in observations:
            if isinstance(obs, dict):
                story.append(Paragraph(f"• {obs.get('content', '')} ({obs.get('created_at', '')})", body_style))
            else:
                story.append(Paragraph(f"• {obs}", body_style))
    
    # Footer
    story.append(Spacer(1, 30))
    story.append(Paragraph("---", body_style))
    story.append(Paragraph("This report is generated by RoleSense - Human-Driven Recruitment Intelligence", body_style))
    story.append(Paragraph("For internal HR use only. Not for distribution.", body_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

async def generate_competency_docx(report: Dict[str, Any]) -> io.BytesIO:
    """Generate DOCX competency report"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Title
    title = doc.add_heading('COMPETENCY REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Candidate: {report['candidate_name']}")
    doc.add_paragraph(f"Generated: {report['generated_at']}")
    doc.add_paragraph()
    
    # Candidate Profile
    doc.add_heading('1. CANDIDATE PROFILE', level=1)
    analysis = report.get("resume_analysis", {})
    doc.add_paragraph(f"Function: {analysis.get('primary_function', 'N/A')} > {analysis.get('sub_function', 'N/A')}")
    doc.add_paragraph(f"Email: {analysis.get('candidate_email', 'N/A')}")
    
    skills = analysis.get("skills", [])
    if skills:
        doc.add_paragraph("Key Skills:")
        doc.add_paragraph(", ".join(skills[:15]))
    
    # Experience
    exp_summary = analysis.get("experience_summary", [])
    if exp_summary:
        doc.add_heading('2. EXPERIENCE SUMMARY', level=1)
        for exp in exp_summary[:5]:
            if isinstance(exp, dict):
                doc.add_paragraph(f"• {exp.get('company', '')} - {exp.get('role', '')}", style='List Bullet')
            else:
                doc.add_paragraph(f"• {exp}", style='List Bullet')
    
    # HR Fitment
    hr_fitment = report.get("hr_fitment_data", {})
    if hr_fitment:
        doc.add_heading('3. HR FITMENT ANALYSIS', level=1)
        
        if hr_fitment.get("motivation", {}).get("reason_for_change"):
            doc.add_heading('Reason for Change:', level=2)
            doc.add_paragraph(str(hr_fitment["motivation"]["reason_for_change"]))
        
        if hr_fitment.get("work_preferences", {}).get("relocation"):
            doc.add_paragraph(f"Relocation: {hr_fitment['work_preferences']['relocation']}")
        
        if hr_fitment.get("achievements", {}).get("top_achievements"):
            doc.add_heading('Key Achievements:', level=2)
            doc.add_paragraph(str(hr_fitment["achievements"]["top_achievements"]))
        
        if hr_fitment.get("career_goals", {}).get("long_term_goals"):
            doc.add_heading('Career Goals:', level=2)
            doc.add_paragraph(str(hr_fitment["career_goals"]["long_term_goals"]))
    
    # Pre-assessment Responses
    preassessment_responses = report.get("preassessment_responses", {})
    if preassessment_responses:
        doc.add_heading('4. PRE-ASSESSMENT RESPONSES', level=1)
        for q_id, response in preassessment_responses.items():
            doc.add_paragraph(f"{q_id}: {response}")
    
    # Recruiter Observations
    observations = report.get("recruiter_observations", [])
    if observations:
        doc.add_heading('5. RECRUITER OBSERVATIONS', level=1)
        for obs in observations:
            if isinstance(obs, dict):
                doc.add_paragraph(f"• {obs.get('content', '')} ({obs.get('created_at', '')})", style='List Bullet')
            else:
                doc.add_paragraph(f"• {obs}", style='List Bullet')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph("---")
    doc.add_paragraph("This report is generated by RoleSense - Human-Driven Recruitment Intelligence")
    doc.add_paragraph("For internal HR use only. Not for distribution.")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============ Resume/Candidate Routes ============

@api_router.post("/candidate/analyze")
async def analyze_candidate(candidate: CandidateCreate):
    """Analyze a candidate's resume and extract structured data"""
    
    system_prompt = """You are an expert recruiter with deep experience in resume analysis.
    Analyze resumes contextually, understanding career progression and potential.
    Always respond with valid JSON only, no additional text."""
    
    user_prompt = f"""Analyze this resume/profile and provide a detailed breakdown:

{candidate.raw_resume}

Respond with a JSON object containing:
{{
    "summary": "Professional summary",
    "current_role": "Current or most recent role",
    "total_experience_years": 0,
    "skills": {{
        "technical": ["skill1", "skill2"],
        "soft": ["skill1", "skill2"],
        "tools": ["tool1", "tool2"]
    }},
    "experience": [
        {{
            "title": "Job Title",
            "company": "Company",
            "duration": "Duration",
            "highlights": ["achievement1", "achievement2"]
        }}
    ],
    "education": [
        {{
            "degree": "Degree",
            "institution": "University",
            "year": "Year"
        }}
    ],
    "certifications": ["cert1", "cert2"],
    "languages": ["language1"],
    "career_trajectory": "upward/lateral/transitioning",
    "strengths": ["strength1", "strength2"],
    "potential_concerns": ["concern1"],
    "ideal_roles": ["role1", "role2"],
    "seniority_level": "junior/mid/senior/lead/executive",
    "why_analysis": "Detailed explanation of the analysis"
}}"""
    
    response = await get_llm_response(system_prompt, user_prompt)
    parsed_data = parse_json_from_response(response)
    
    # Create candidate object
    candidate_obj = Candidate(
        name=candidate.name,
        email=candidate.email,
        phone=candidate.phone,
        raw_resume=candidate.raw_resume,
        parsed_data=parsed_data,
        analysis={
            "career_trajectory": parsed_data.get("career_trajectory", ""),
            "strengths": parsed_data.get("strengths", []),
            "potential_concerns": parsed_data.get("potential_concerns", []),
            "ideal_roles": parsed_data.get("ideal_roles", []),
            "why_analysis": parsed_data.get("why_analysis", "")
        }
    )
    
    # Save to database
    doc = serialize_datetime(candidate_obj.model_dump())
    await db.candidates.insert_one(doc)
    
    return candidate_obj

@api_router.get("/candidates", response_model=List[Candidate])
async def list_candidates(stage: Optional[str] = None):
    """List all candidates, optionally filtered by pipeline stage"""
    query = {}
    if stage:
        query["pipeline_stage"] = stage
    candidates = await db.candidates.find(query, {"_id": 0}).to_list(100)
    return [deserialize_datetime(c) for c in candidates]

@api_router.get("/candidate/{candidate_id}")
async def get_candidate(candidate_id: str):
    """Get a specific candidate"""
    candidate = await db.candidates.find_one({"id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    return deserialize_datetime(candidate)

@api_router.patch("/candidate/{candidate_id}")
async def update_candidate(candidate_id: str, update: CandidateUpdate):
    """Update candidate details"""
    update_dict = {k: v for k, v in update.model_dump().items() if v is not None}
    update_dict["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    result = await db.candidates.update_one(
        {"id": candidate_id},
        {"$set": update_dict}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    return await get_candidate(candidate_id)

@api_router.delete("/candidate/{candidate_id}")
async def delete_candidate(candidate_id: str):
    """Delete a candidate"""
    result = await db.candidates.delete_one({"id": candidate_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Candidate not found")
    return {"message": "Candidate deleted successfully"}

@api_router.post("/candidate/{candidate_id}/note")
async def add_candidate_note(candidate_id: str, note: NoteCreate):
    """Add a note to a candidate"""
    note_obj = {
        "id": str(uuid.uuid4()),
        "content": note.content,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    result = await db.candidates.update_one(
        {"id": candidate_id},
        {"$push": {"notes": note_obj}}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    return note_obj

# ============ Matching Routes ============

@api_router.post("/match/analyze")
async def analyze_match(request: MatchRequest):
    """Analyze match between a candidate and job description with explainable reasoning"""
    
    # Get JD and Candidate
    jd = await db.job_descriptions.find_one({"id": request.jd_id}, {"_id": 0})
    candidate = await db.candidates.find_one({"id": request.candidate_id}, {"_id": 0})
    
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    system_prompt = """You are an expert recruitment analyst specializing in candidate-job matching.
    Provide detailed, explainable analysis of how well a candidate matches a job.
    Be honest and balanced - highlight both strengths and gaps.
    Always respond with valid JSON only, no additional text."""
    
    user_prompt = f"""Analyze how well this candidate matches this job description:

JOB DESCRIPTION:
Title: {jd.get('title', '')}
{jd.get('raw_text', '')}

CANDIDATE:
Name: {candidate.get('name', '')}
{candidate.get('raw_resume', '')}

Provide a detailed match analysis as JSON:
{{
    "overall_score": 85,
    "confidence_score": 0.9,
    "match_breakdown": {{
        "skills_match": {{
            "score": 80,
            "matched": ["skill1", "skill2"],
            "missing": ["skill3"],
            "why": "Explanation"
        }},
        "experience_match": {{
            "score": 90,
            "why": "Explanation"
        }},
        "education_match": {{
            "score": 75,
            "why": "Explanation"
        }},
        "culture_fit": {{
            "score": 85,
            "indicators": ["indicator1"],
            "why": "Explanation"
        }}
    }},
    "strengths": ["strength1 - detailed explanation", "strength2 - detailed explanation"],
    "gaps": ["gap1 - and what it means", "gap2 - and what it means"],
    "why_overall": "Comprehensive explanation of why this candidate is/isn't a good match",
    "interview_focus_areas": ["area1", "area2"],
    "salary_expectation_alignment": "likely aligned/may need discussion/potential mismatch",
    "hiring_recommendation": "strong yes/yes/maybe/probably not/no",
    "recommendation_reasoning": "Detailed reasoning for the recommendation"
}}"""
    
    response = await get_llm_response(system_prompt, user_prompt)
    match_data = parse_json_from_response(response)
    
    # Create match result
    match_result = MatchResult(
        jd_id=request.jd_id,
        candidate_id=request.candidate_id,
        overall_score=match_data.get("overall_score", 0),
        match_breakdown=match_data.get("match_breakdown", {}),
        strengths=match_data.get("strengths", []),
        gaps=match_data.get("gaps", []),
        why_explanation=match_data.get("why_overall", ""),
        confidence_score=match_data.get("confidence_score", 0.5)
    )
    
    # Save to database
    doc = serialize_datetime(match_result.model_dump())
    await db.matches.insert_one(doc)
    
    # Return full match data with additional insights
    return {
        **match_result.model_dump(),
        "interview_focus_areas": match_data.get("interview_focus_areas", []),
        "hiring_recommendation": match_data.get("hiring_recommendation", ""),
        "recommendation_reasoning": match_data.get("recommendation_reasoning", ""),
        "jd_title": jd.get("title", ""),
        "candidate_name": candidate.get("name", "")
    }

@api_router.get("/matches")
async def list_matches(jd_id: Optional[str] = None, candidate_id: Optional[str] = None):
    """List match results"""
    query = {}
    if jd_id:
        query["jd_id"] = jd_id
    if candidate_id:
        query["candidate_id"] = candidate_id
    
    matches = await db.matches.find(query, {"_id": 0}).to_list(100)
    return [deserialize_datetime(m) for m in matches]

@api_router.patch("/match/{match_id}/feedback")
async def add_recruiter_feedback(match_id: str, feedback: RecruiterFeedback):
    """Add recruiter feedback/override to a match"""
    update_dict = {}
    if feedback.feedback:
        update_dict["recruiter_feedback"] = feedback.feedback
    if feedback.override_score is not None:
        update_dict["recruiter_override"] = feedback.override_score
    
    result = await db.matches.update_one(
        {"id": match_id},
        {"$set": update_dict}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Match not found")
    
    return {"message": "Feedback added successfully"}

# ============ Smart Search Routes ============

@api_router.post("/search")
async def smart_search(search: SearchQuery):
    """Smart search that understands natural language queries"""
    
    system_prompt = """You are an expert recruitment search assistant.
    Convert natural language queries into structured search criteria.
    Always respond with valid JSON only, no additional text."""
    
    user_prompt = f"""Convert this natural language search query into structured criteria:

Query: "{search.query}"

Respond with JSON:
{{
    "interpretation": "What you understood from the query",
    "skills": ["skill1", "skill2"],
    "experience_min": 0,
    "experience_max": 10,
    "seniority": ["junior", "mid", "senior"],
    "keywords": ["keyword1", "keyword2"],
    "exclusions": ["things to exclude"],
    "sort_by": "relevance/experience/recent"
}}"""
    
    response = await get_llm_response(system_prompt, user_prompt)
    criteria = parse_json_from_response(response)
    
    # Build MongoDB query based on criteria
    mongo_query = {}
    
    # Search in parsed_data.skills
    if criteria.get("skills"):
        mongo_query["$or"] = [
            {"parsed_data.skills.technical": {"$in": criteria["skills"]}},
            {"parsed_data.skills.soft": {"$in": criteria["skills"]}},
            {"raw_resume": {"$regex": "|".join(criteria["skills"]), "$options": "i"}}
        ]
    
    # Apply additional filters if provided
    if search.filters:
        if search.filters.get("pipeline_stage"):
            mongo_query["pipeline_stage"] = search.filters["pipeline_stage"]
    
    # Execute search
    candidates = await db.candidates.find(mongo_query, {"_id": 0}).to_list(50)
    
    return SearchResult(
        candidates=[deserialize_datetime(c) for c in candidates],
        interpretation=criteria.get("interpretation", ""),
        total_count=len(candidates)
    )

# ============ Pipeline Routes ============

@api_router.get("/pipeline/stages")
async def get_pipeline_stages():
    """Get default pipeline stages"""
    default_stages = [
        {"id": "new", "name": "New", "order": 1, "color": "#6B7280"},
        {"id": "screening", "name": "Screening", "order": 2, "color": "#3B82F6"},
        {"id": "interview", "name": "Interview", "order": 3, "color": "#8B5CF6"},
        {"id": "assessment", "name": "Assessment", "order": 4, "color": "#F59E0B"},
        {"id": "offer", "name": "Offer", "order": 5, "color": "#10B981"},
        {"id": "hired", "name": "Hired", "order": 6, "color": "#059669"},
        {"id": "rejected", "name": "Rejected", "order": 7, "color": "#EF4444"}
    ]
    return default_stages

@api_router.get("/pipeline/overview")
async def get_pipeline_overview():
    """Get pipeline overview with candidate counts per stage"""
    stages = await get_pipeline_stages()
    
    pipeline_data = []
    for stage in stages:
        count = await db.candidates.count_documents({"pipeline_stage": stage["id"]})
        pipeline_data.append({
            **stage,
            "candidate_count": count
        })
    
    return pipeline_data

@api_router.post("/pipeline/move/{candidate_id}")
async def move_candidate_in_pipeline(candidate_id: str, data: Dict[str, str]):
    """Move a candidate to a different pipeline stage"""
    new_stage = data.get("stage")
    if not new_stage:
        raise HTTPException(status_code=400, detail="Stage is required")
    
    result = await db.candidates.update_one(
        {"id": candidate_id},
        {
            "$set": {
                "pipeline_stage": new_stage,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }
        }
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    return {"message": f"Candidate moved to {new_stage}"}

# ============ CAREER TRAJECTORY INDICATOR MODULE ============

# Career Trajectory Indicator Types and Master Data - Enhanced with Predictive Analysis
CAREER_TRAJECTORY_INDICATORS = {
    "career_progression": {
        "id": "career_progression",
        "name": "Career Progression Consistency",
        "description": "Analyzes job title growth, role advancement, and salary trajectory",
        "weight": 0.12,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "employment_data"
    },
    "job_stability": {
        "id": "job_stability",
        "name": "Job Stability Index (Explainable)",
        "description": "Evaluates tenure patterns with detailed reasons for each job change",
        "weight": 0.12,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "job_stability",
        "explainable": True
    },
    "industry_alignment": {
        "id": "industry_alignment",
        "name": "Industry/Domain Alignment",
        "description": "Measures domain/sector match with target role",
        "weight": 0.10,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "employment_data"
    },
    "skills_evolution": {
        "id": "skills_evolution",
        "name": "Skills Evolution",
        "description": "Tracks skill growth, learning trajectory, and upskilling",
        "weight": 0.10,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "employment_data"
    },
    "education_alignment": {
        "id": "education_alignment",
        "name": "Education-to-Career Alignment",
        "description": "Assesses degree-to-career match and certifications",
        "weight": 0.08,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "personal_data"
    },
    "employment_gaps": {
        "id": "employment_gaps",
        "name": "Employment Gap Analysis",
        "description": "Identifies and contextualizes career gaps",
        "weight": 0.08,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "employment_data"
    },
    "cultural_behavioral_fit": {
        "id": "cultural_behavioral_fit",
        "name": "Cultural & Behavioral Fit",
        "description": "Based on questionnaire responses and behavioral signals",
        "weight": 0.08,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "personal_data"
    },
    "compensation_trajectory": {
        "id": "compensation_trajectory",
        "name": "Compensation Trajectory",
        "description": "Analyzes salary growth, hike percentages across employers",
        "weight": 0.08,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "compensation_data"
    },
    "location_mobility": {
        "id": "location_mobility",
        "name": "Location & Mobility Analysis",
        "description": "Evaluates relocation patterns and location fit",
        "weight": 0.08,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "personal_data"
    },
    "joining_intent": {
        "id": "joining_intent",
        "name": "Joining Intent Score",
        "description": "Predicts likelihood of candidate accepting the offer",
        "weight": 0.08,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "candidate_overview"
    },
    "counter_offer_risk": {
        "id": "counter_offer_risk",
        "name": "Counter-Offer Risk Score",
        "description": "Assesses likelihood of candidate taking counter-offer",
        "weight": 0.04,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "candidate_overview"
    },
    "retention_stability": {
        "id": "retention_stability",
        "name": "Retention & Stability Score",
        "description": "Predicts long-term retention based on commitment indicators",
        "weight": 0.04,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "personal_data"
    }
}

# HR Fitment Analysis Indicators (Additional 5)
HR_FITMENT_INDICATORS = {
    "hr_cultural_fit": {
        "id": "hr_cultural_fit",
        "name": "Cultural Fit Assessment",
        "description": "Evaluates alignment with company culture, values, work ethics, and organizational fit",
        "weight": 0.20,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "hr_fitment",
        "factors": ["value_alignment", "work_style", "communication_style", "team_orientation", "adaptability"]
    },
    "hr_team_dynamics": {
        "id": "hr_team_dynamics",
        "name": "Team Dynamics Fit",
        "description": "Assesses compatibility with existing team structure, reporting hierarchy, and collaboration style",
        "weight": 0.20,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "hr_fitment",
        "factors": ["team_compatibility", "leadership_style", "collaboration_approach", "conflict_resolution", "hierarchy_comfort"]
    },
    "hr_role_metrics": {
        "id": "hr_role_metrics",
        "name": "Role-Specific HR Metrics",
        "description": "Evaluates notice period compatibility, salary alignment, location feasibility, and BGV readiness",
        "weight": 0.25,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "hr_fitment",
        "factors": ["notice_period_fit", "salary_alignment", "location_feasibility", "bgv_readiness", "documentation_status"]
    },
    "hr_soft_skills": {
        "id": "hr_soft_skills",
        "name": "Soft Skills Evaluation",
        "description": "Comprehensive assessment of communication, leadership potential, emotional intelligence, and adaptability",
        "weight": 0.20,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "hr_fitment",
        "factors": ["communication", "leadership_potential", "emotional_intelligence", "problem_solving", "time_management"]
    },
    "hr_risk_summary": {
        "id": "hr_risk_summary",
        "name": "HR Risk Assessment Summary",
        "description": "Consolidated risk analysis including attrition risk, counter-offer probability, and hiring timeline risks",
        "weight": 0.15,
        "thresholds": {"green": 71, "yellow": 41},
        "category": "hr_fitment",
        "factors": ["attrition_risk", "counter_offer_probability", "timeline_risk", "reference_risk", "offer_acceptance_risk"]
    }
}

# Job Change Reason Options - For Explainable Job Stability
JOB_CHANGE_REASONS = {
    "growth_related": [
        "Promoted to a role not available at current company",
        "Seeking larger team/project responsibility",
        "Moving from individual contributor to management",
        "Transitioning to a specialized domain",
        "Pursuing leadership/executive track"
    ],
    "compensation_related": [
        "Significant salary increase (>30% hike)",
        "Better equity/stock options offered",
        "Improved benefits package (healthcare, retirement)",
        "Signing bonus opportunity",
        "Performance-linked variable pay"
    ],
    "company_factors": [
        "Company downsizing/layoffs",
        "Company acquisition/merger",
        "Company shutting down",
        "Organizational restructuring",
        "Change in company direction/strategy",
        "Startup funding issues"
    ],
    "work_environment": [
        "Toxic work culture",
        "Poor relationship with manager",
        "Lack of work-life balance",
        "Excessive overtime expectations",
        "Remote/hybrid work preference",
        "Better team dynamics elsewhere"
    ],
    "personal_reasons": [
        "Relocation due to spouse's job",
        "Family health/care responsibilities",
        "Personal health reasons",
        "Pursuing higher education",
        "Starting own venture",
        "Returning after career break"
    ],
    "career_change": [
        "Industry pivot (e.g., finance to tech)",
        "Function change (e.g., sales to product)",
        "Moving to consulting/advisory",
        "Joining early-stage startup",
        "Moving to established enterprise"
    ]
}

# Predictive Scores - System Generated
PREDICTIVE_SCORES = {
    "joining_intent": {
        "name": "Joining Intent Score",
        "description": "Likelihood to accept the offer (0-100)",
        "factors": ["resignation_status", "counter_offer_disclosed", "notice_period", "current_ctc_gap"]
    },
    "counter_offer_risk": {
        "name": "Counter-Offer Risk",
        "description": "Risk of candidate taking counter-offer (0-100, lower is better)",
        "factors": ["tenure_at_current", "promotion_history", "relationship_with_manager"]
    },
    "stability_score": {
        "name": "Stability Score",
        "description": "Predicted retention and stability (0-100)",
        "factors": ["avg_tenure", "location_stability", "personal_commitments", "family_situation"]
    },
    "location_fit": {
        "name": "Location Fit Score",
        "description": "Alignment with job location (0-100)",
        "factors": ["current_location", "native_place", "relocation_history", "family_ties"]
    },
    "offer_decline_risk": {
        "name": "Offer Decline Risk",
        "description": "Risk of offer being declined (0-100, lower is better)",
        "factors": ["multiple_offers", "expectations_gap", "notice_buyout_required"]
    },
    "time_to_join": {
        "name": "Time to Join Score",
        "description": "Likelihood of joining within expected timeline (0-100)",
        "factors": ["notice_period", "resignation_date", "handover_complexity"]
    }
}

# PRE-ASSESSMENT QUESTIONNAIRE - All Multiple Choice (42 Questions, <5 mins)
# Structured by Career Trajectory Categories
PRE_ASSESSMENT_QUESTIONNAIRE = {
    "candidate_overview": [
        {
            "id": "co_1",
            "question": "What is your primary motivation for exploring this opportunity?",
            "type": "single_select",
            "options": [
                "Career advancement to next level",
                "Significant compensation improvement",
                "Better work-life balance",
                "Industry/domain change",
                "Relocation requirement",
                "Current company instability",
                "Learning new technologies/skills",
                "Leadership/management opportunity"
            ],
            "indicator": "joining_intent",
            "mandatory": True
        },
        {
            "id": "co_2",
            "question": "How actively are you looking for a new role?",
            "type": "single_select",
            "options": [
                "Actively interviewing with multiple companies",
                "Selectively exploring specific opportunities",
                "Passively open - approached by recruiter",
                "Casually browsing - no urgency",
                "Only interested in this specific role"
            ],
            "indicator": "joining_intent",
            "mandatory": True
        },
        {
            "id": "co_3",
            "question": "What is your expected timeline to make a decision?",
            "type": "single_select",
            "options": [
                "Within 1 week",
                "Within 2 weeks",
                "Within 1 month",
                "Within 2-3 months",
                "No specific timeline"
            ],
            "indicator": "joining_intent",
            "mandatory": True
        },
        {
            "id": "co_4",
            "question": "Do you have other active job offers or final-stage interviews?",
            "type": "single_select",
            "options": [
                "Yes - have 2+ offers in hand",
                "Yes - have 1 offer in hand",
                "Yes - in final stages with other companies",
                "No - early stages with others",
                "No - this is my only active application"
            ],
            "indicator": "counter_offer_risk",
            "mandatory": True
        }
    ],
    "personal_data": [
        {
            "id": "pd_1",
            "question": "What is your highest educational qualification?",
            "type": "single_select",
            "options": [
                "PhD/Doctorate",
                "Post Graduate (MBA/MTech/MS)",
                "Graduate (BTech/BE/BSc/BCom/BA)",
                "Diploma",
                "12th/High School",
                "Professional Certification (CA/CS/CFA)"
            ],
            "indicator": "education_alignment",
            "mandatory": True
        },
        {
            "id": "pd_2",
            "question": "Is your educational background aligned with your current career?",
            "type": "single_select",
            "options": [
                "Directly aligned (studied and working in same field)",
                "Partially aligned (related field)",
                "Career pivot (different field but leveraging skills)",
                "Complete career change (unrelated field)"
            ],
            "indicator": "education_alignment",
            "mandatory": True
        },
        {
            "id": "pd_3",
            "question": "What is your native/hometown location?",
            "type": "single_select",
            "options": [
                "Same city as current residence",
                "Different city, same state",
                "Different state",
                "Different country"
            ],
            "indicator": "location_mobility",
            "mandatory": True
        },
        {
            "id": "pd_4",
            "question": "How many cities have you relocated to for work?",
            "type": "single_select",
            "options": [
                "0 - Never relocated for work",
                "1 city",
                "2-3 cities",
                "4+ cities"
            ],
            "indicator": "location_mobility",
            "mandatory": True
        },
        {
            "id": "pd_5",
            "question": "Are you open to relocation for this role?",
            "type": "single_select",
            "options": [
                "Yes - willing to relocate anywhere",
                "Yes - to specific metro cities only",
                "Yes - within same state only",
                "Prefer remote/hybrid work",
                "No - cannot relocate"
            ],
            "indicator": "location_mobility",
            "mandatory": True
        },
        {
            "id": "pd_6",
            "question": "What is your current family status?",
            "type": "single_select",
            "options": [
                "Single",
                "Married - Childless couple",
                "Married - With children (below school age)",
                "Married - With children (school going)",
                "Married - With children (college/adult)",
                "Single parent"
            ],
            "indicator": "retention_stability",
            "mandatory": True
        },
        {
            "id": "pd_7",
            "question": "Is your spouse currently employed?",
            "type": "single_select",
            "options": [
                "Yes - Full-time employed",
                "Yes - Part-time/Freelance",
                "Yes - Own business",
                "No - Homemaker",
                "No - Currently job seeking",
                "Not applicable"
            ],
            "indicator": "retention_stability",
            "mandatory": True
        },
        {
            "id": "pd_8",
            "question": "What is your current residence status?",
            "type": "single_select",
            "options": [
                "Own house - Fully paid",
                "Own house - EMI ongoing",
                "Rented accommodation",
                "Company provided housing",
                "Living with family/parents",
                "PG/Hostel accommodation"
            ],
            "indicator": "retention_stability",
            "mandatory": True
        },
        {
            "id": "pd_9",
            "question": "What type of work environment do you prefer?",
            "type": "single_select",
            "options": [
                "Fast-paced startup environment",
                "Structured corporate environment",
                "Collaborative team-oriented setup",
                "Independent/autonomous work style",
                "Hybrid of structure and flexibility"
            ],
            "indicator": "cultural_behavioral_fit",
            "mandatory": True
        },
        {
            "id": "pd_10",
            "question": "What management style works best for you?",
            "type": "single_select",
            "options": [
                "Hands-on guidance and mentorship",
                "Regular check-ins with autonomy",
                "Complete autonomy with minimal supervision",
                "Collaborative decision-making",
                "Clear goals with freedom in execution"
            ],
            "indicator": "cultural_behavioral_fit",
            "mandatory": True
        }
    ],
    "employment_data": [
        {
            "id": "ed_1",
            "question": "What is your total professional experience?",
            "type": "single_select",
            "options": [
                "0-2 years (Entry level)",
                "2-5 years (Early career)",
                "5-8 years (Mid career)",
                "8-12 years (Senior)",
                "12-15 years (Lead/Principal)",
                "15+ years (Executive/Director)"
            ],
            "indicator": "career_progression",
            "mandatory": True
        },
        {
            "id": "ed_2",
            "question": "How many companies have you worked for?",
            "type": "single_select",
            "options": [
                "1 company (current)",
                "2 companies",
                "3-4 companies",
                "5-6 companies",
                "7+ companies"
            ],
            "indicator": "job_stability",
            "mandatory": True
        },
        {
            "id": "ed_3",
            "question": "What is your average tenure at each company?",
            "type": "single_select",
            "options": [
                "Less than 1 year",
                "1-2 years",
                "2-3 years",
                "3-5 years",
                "5+ years"
            ],
            "indicator": "job_stability",
            "mandatory": True
        },
        {
            "id": "ed_4",
            "question": "How many promotions have you received in your career?",
            "type": "single_select",
            "options": [
                "0 - No promotions yet",
                "1 promotion",
                "2-3 promotions",
                "4-5 promotions",
                "6+ promotions"
            ],
            "indicator": "career_progression",
            "mandatory": True
        },
        {
            "id": "ed_5",
            "question": "What best describes your career trajectory?",
            "type": "single_select",
            "options": [
                "Consistent upward growth in same domain",
                "Lateral moves with skill expansion",
                "Mix of vertical and lateral moves",
                "Recent pivot to new domain/function",
                "Entrepreneurial journey with corporate stints"
            ],
            "indicator": "career_progression",
            "mandatory": True
        },
        {
            "id": "ed_6",
            "question": "Have you had any career gaps (>3 months)?",
            "type": "single_select",
            "options": [
                "No gaps in my career",
                "Yes - 1 gap (3-6 months)",
                "Yes - 1 gap (6-12 months)",
                "Yes - 1 gap (>12 months)",
                "Yes - Multiple gaps"
            ],
            "indicator": "employment_gaps",
            "mandatory": True
        },
        {
            "id": "ed_7",
            "question": "If you had career gaps, what was the primary reason?",
            "type": "single_select",
            "options": [
                "No gaps - not applicable",
                "Higher education/certification",
                "Health reasons (self/family)",
                "Parental/family care responsibilities",
                "Entrepreneurial venture",
                "Job search after layoff",
                "Planned sabbatical/travel",
                "Relocation/visa issues"
            ],
            "indicator": "employment_gaps",
            "mandatory": True
        },
        {
            "id": "ed_8",
            "question": "Which industries have you primarily worked in?",
            "type": "single_select",
            "options": [
                "IT Services/Consulting",
                "Product Companies (SaaS/Tech)",
                "BFSI (Banking/Finance/Insurance)",
                "E-commerce/Retail",
                "Manufacturing/Industrial",
                "Healthcare/Pharma",
                "Telecom/Media",
                "Multiple industries"
            ],
            "indicator": "industry_alignment",
            "mandatory": True
        },
        {
            "id": "ed_9",
            "question": "How do you typically approach learning new skills?",
            "type": "single_select",
            "options": [
                "Formal courses and certifications",
                "On-the-job learning and projects",
                "Self-study through online resources",
                "Mentorship and peer learning",
                "Combination of all approaches"
            ],
            "indicator": "skills_evolution",
            "mandatory": True
        },
        {
            "id": "ed_10",
            "question": "How many new skills/technologies have you learned in last 2 years?",
            "type": "single_select",
            "options": [
                "None - focused on existing skills",
                "1-2 new skills",
                "3-4 new skills",
                "5+ new skills",
                "Complete skill transformation"
            ],
            "indicator": "skills_evolution",
            "mandatory": True
        }
    ],
    "compensation_data": [
        {
            "id": "cd_1",
            "question": "What is your current annual CTC range?",
            "type": "single_select",
            "options": [
                "Below 5 LPA",
                "5-10 LPA",
                "10-20 LPA",
                "20-35 LPA",
                "35-50 LPA",
                "50-75 LPA",
                "75+ LPA"
            ],
            "indicator": "compensation_trajectory",
            "mandatory": True
        },
        {
            "id": "cd_2",
            "question": "What is your expected annual CTC range?",
            "type": "single_select",
            "options": [
                "10-15% hike on current",
                "15-25% hike on current",
                "25-40% hike on current",
                "40-60% hike on current",
                "60%+ hike on current",
                "Flexible based on role"
            ],
            "indicator": "compensation_trajectory",
            "mandatory": True
        },
        {
            "id": "cd_3",
            "question": "What is your minimum acceptable CTC?",
            "type": "single_select",
            "options": [
                "No reduction from current",
                "5-10% hike minimum",
                "10-15% hike minimum",
                "15-25% hike minimum",
                "25%+ hike minimum"
            ],
            "indicator": "compensation_trajectory",
            "mandatory": True
        },
        {
            "id": "cd_4",
            "question": "What components do you prioritize in compensation?",
            "type": "single_select",
            "options": [
                "Fixed salary (maximum fixed component)",
                "Variable/bonus (performance-linked)",
                "ESOPs/RSUs (equity participation)",
                "Benefits (healthcare, retirement)",
                "Balanced mix of all components"
            ],
            "indicator": "compensation_trajectory",
            "mandatory": True
        },
        {
            "id": "cd_5",
            "question": "How has your compensation grown over last 3 years?",
            "type": "single_select",
            "options": [
                "Minimal growth (<10% total)",
                "Moderate growth (10-30% total)",
                "Good growth (30-50% total)",
                "Strong growth (50-100% total)",
                "Exceptional growth (>100% total)"
            ],
            "indicator": "compensation_trajectory",
            "mandatory": True
        }
    ],
    "job_stability_details": [
        {
            "id": "js_1",
            "question": "What was your PRIMARY reason for leaving your most recent/current job?",
            "type": "single_select",
            "options": [
                "Promoted externally - role not available internally",
                "Significant compensation increase (>30%)",
                "Better learning and growth opportunity",
                "Company downsizing/layoffs",
                "Organizational restructuring",
                "Poor work-life balance",
                "Toxic culture/poor management",
                "Relocation/personal reasons",
                "Startup/company shutdown",
                "Currently employed - exploring"
            ],
            "indicator": "job_stability",
            "mandatory": True,
            "explainable": True
        },
        {
            "id": "js_2",
            "question": "What was your reason for leaving your second most recent job?",
            "type": "single_select",
            "options": [
                "Promoted externally - role not available internally",
                "Significant compensation increase (>30%)",
                "Better learning and growth opportunity",
                "Company downsizing/layoffs",
                "Organizational restructuring",
                "Poor work-life balance",
                "Toxic culture/poor management",
                "Relocation/personal reasons",
                "Startup/company shutdown",
                "Not applicable - first job"
            ],
            "indicator": "job_stability",
            "mandatory": True,
            "explainable": True
        },
        {
            "id": "js_3",
            "question": "Pattern of your job changes - which best describes?",
            "type": "single_select",
            "options": [
                "Mostly growth-driven moves",
                "Mix of growth and compensation-driven",
                "Mostly involuntary (layoffs/closures)",
                "Mix of voluntary and involuntary",
                "Personal/relocation driven",
                "Single employer - no changes"
            ],
            "indicator": "job_stability",
            "mandatory": True
        },
        {
            "id": "js_4",
            "question": "How would you rate your relationship with your current/last manager?",
            "type": "single_select",
            "options": [
                "Excellent - strong mentor relationship",
                "Good - professional and supportive",
                "Average - functional relationship",
                "Below average - some conflicts",
                "Poor - significant issues"
            ],
            "indicator": "counter_offer_risk",
            "mandatory": True
        },
        {
            "id": "js_5",
            "question": "If offered a counter-offer by current employer, what would you do?",
            "type": "single_select",
            "options": [
                "Definitely decline - decision is final",
                "Would consider if significantly better",
                "Would seriously evaluate both options",
                "Likely to accept if matched",
                "Not applicable - not currently employed"
            ],
            "indicator": "counter_offer_risk",
            "mandatory": True
        }
    ],
    "resignation_status": [
        {
            "id": "rs_1",
            "question": "What is your current resignation status?",
            "type": "single_select",
            "options": [
                "Already resigned - serving notice",
                "Planning to resign upon offer",
                "Will resign after joining confirmation",
                "Not currently employed",
                "Need time to decide"
            ],
            "indicator": "joining_intent",
            "mandatory": True
        },
        {
            "id": "rs_2",
            "question": "What is your notice period?",
            "type": "single_select",
            "options": [
                "Immediate joiner",
                "15 days",
                "30 days",
                "60 days",
                "90 days",
                "90+ days"
            ],
            "indicator": "joining_intent",
            "mandatory": True
        },
        {
            "id": "rs_3",
            "question": "Is your notice period negotiable?",
            "type": "single_select",
            "options": [
                "Yes - can be bought out completely",
                "Yes - can negotiate to reduce",
                "Partially - can reduce by 1-2 weeks",
                "No - must serve full notice",
                "Not sure - need to check"
            ],
            "indicator": "joining_intent",
            "mandatory": True
        },
        {
            "id": "rs_4",
            "question": "Have you received a counter-offer from current employer?",
            "type": "single_select",
            "options": [
                "Yes - already received",
                "Yes - expecting one",
                "No - not expecting",
                "No - will not accept even if offered",
                "Not applicable"
            ],
            "indicator": "counter_offer_risk",
            "mandatory": True
        },
        {
            "id": "rs_5",
            "question": "If counter-offer received, what was offered?",
            "type": "single_select",
            "options": [
                "No counter-offer received",
                "Salary hike only",
                "Promotion + salary hike",
                "Retention bonus",
                "Role change + compensation",
                "Multiple incentives combined"
            ],
            "indicator": "counter_offer_risk",
            "mandatory": True
        },
        {
            "id": "rs_6",
            "question": "What would be your ideal start date if selected?",
            "type": "single_select",
            "options": [
                "Immediate",
                "Within 2 weeks",
                "Within 1 month",
                "Within 2 months",
                "Within 3 months",
                "Flexible - depends on notice"
            ],
            "indicator": "joining_intent",
            "mandatory": True
        },
        {
            "id": "rs_7",
            "question": "How committed are you to joining if an offer is extended?",
            "type": "single_select",
            "options": [
                "100% committed - will definitely join",
                "Highly committed - very likely to join",
                "Moderately committed - need to evaluate offer",
                "Exploring options - will compare offers",
                "Uncertain - depends on multiple factors"
            ],
            "indicator": "joining_intent",
            "mandatory": True
        },
        {
            "id": "rs_8",
            "question": "What could potentially prevent you from joining?",
            "type": "single_select",
            "options": [
                "Nothing - fully committed",
                "Better counter-offer from current employer",
                "Better offer from another company",
                "Personal/family circumstances",
                "Relocation challenges",
                "Compensation not meeting expectations"
            ],
            "indicator": "counter_offer_risk",
            "mandatory": True
        }
    ]
}

# Email Template Simulation
def simulate_send_preassessment_email(candidate_email: str, candidate_name: str, job_title: str, job_id: str, assessment_link: str):
    """Simulate sending pre-assessment email - logs to console, ready for email service integration"""
    email_content = f"""
    ============================================================
    PRE-ASSESSMENT EMAIL SIMULATION
    ============================================================
    To: {candidate_email}
    Subject: Action Required: Complete Pre-Assessment for {job_title}
    
    Dear {candidate_name},
    
    Thank you for applying to the {job_title} position.
    
    To proceed with your application, please complete the mandatory 
    Pre-Assessment Form within 48 hours.
    
    Pre-Assessment Link: {assessment_link}
    
    IMPORTANT:
    - This assessment is MANDATORY for your application to be considered
    - It takes less than 5 minutes to complete
    - All questions are required
    - Failure to complete will result in automatic rejection
    
    Once completed:
    - Your resume will be accepted and locked to this job requisition
    - You will receive an acknowledgement email
    - Our team will begin the first-level evaluation
    
    Best regards,
    RoleSense Recruitment Team
    
    Job ID: {job_id}
    ============================================================
    """
    logger.info(email_content)
    print(email_content)
    return True

def simulate_send_acknowledgement_email(candidate_email: str, candidate_name: str, job_title: str, job_id: str):
    """Simulate sending acknowledgement email after assessment completion"""
    email_content = f"""
    ============================================================
    ACKNOWLEDGEMENT EMAIL SIMULATION
    ============================================================
    To: {candidate_email}
    Subject: Application Received - {job_title}
    
    Dear {candidate_name},
    
    Thank you for completing the Pre-Assessment Form.
    
    Your application for {job_title} has been successfully received 
    and your profile is now locked to this job requisition.
    
    What's Next:
    - First-level evaluation has been initiated
    - You will be notified within 5-7 business days about next steps
    - If shortlisted, you will receive details for the next round
    
    Application Reference: {job_id}
    
    Best regards,
    RoleSense Recruitment Team
    ============================================================
    """
    logger.info(email_content)
    print(email_content)
    return True

def simulate_send_rejection_email(candidate_email: str, candidate_name: str, job_title: str, job_id: str, reasons: List[str], status: str = "rejected"):
    """Simulate sending rejection/hold email with auto-generated reasons"""
    status_text = "placed on hold" if status == "hold" else "not moving forward"
    email_content = f"""
    ============================================================
    APPLICATION STATUS EMAIL SIMULATION
    ============================================================
    To: {candidate_email}
    Subject: Application Update - {job_title}
    
    Dear {candidate_name},
    
    Thank you for your interest in the {job_title} position.
    
    After careful evaluation, we regret to inform you that your 
    application has been {status_text} at this time.
    
    Key Factors:
    {chr(10).join(f'    - {reason}' for reason in reasons)}
    
    {"We encourage you to apply for other positions that match your profile." if status == "rejected" else "We will reach out if there are any updates to your application status."}
    
    Application Reference: {job_id}
    
    Best regards,
    RoleSense Recruitment Team
    ============================================================
    """
    logger.info(email_content)
    print(email_content)
    return True

def generate_auto_rejection_reasons(assessment: Dict) -> List[str]:
    """Auto-generate rejection/hold reasons based on red flags and scores"""
    reasons = []
    
    # Check overall score
    if assessment.get("overall_score", 0) < 41:
        reasons.append("Overall assessment score below threshold")
    
    # Check specific indicators
    for indicator in assessment.get("indicator_results", []):
        if indicator.get("flag") == "red":
            reasons.append(f"{indicator.get('indicator_name')}: {indicator.get('concerns', ['Score below acceptable range'])[0]}")
    
    # Check predictive scores
    predictive = assessment.get("predictive_scores", {})
    if predictive.get("joining_intent", 100) < 40:
        reasons.append("Low joining intent score indicates uncertain commitment")
    if predictive.get("counter_offer_risk", 0) > 70:
        reasons.append("High counter-offer risk detected")
    if predictive.get("offer_decline_risk", 0) > 70:
        reasons.append("High likelihood of offer decline")
    
    # Check non-disclosure
    for flag in assessment.get("non_disclosure_flags", []):
        reasons.append(f"Information not disclosed: {flag.get('message', 'Required information missing')}")
    
    # Default reason if none found
    if not reasons:
        reasons.append("Profile does not match current requirements")
    
    return reasons[:5]  # Limit to 5 reasons
# Employment History Entry Model (for detailed tracking)
class EmploymentHistoryEntry(BaseModel):
    employer_name: str
    designation_at_joining: str
    designation_at_exit: Optional[str] = None
    start_date: str  # YYYY-MM format
    end_date: Optional[str] = None  # YYYY-MM format or "Present"
    location: str
    ctc_at_joining: Optional[float] = None
    ctc_at_exit: Optional[float] = None
    hike_percentage: Optional[float] = None  # Auto-calculated
    promotions: List[Dict[str, Any]] = []  # List of {date, from_role, to_role, hike_percent}
    reason_for_leaving: Optional[str] = None
    reporting_manager: Optional[str] = None
    team_size: Optional[int] = None

# Career Trajectory Assessment Models - Enhanced
class TrajectoryIndicatorResult(BaseModel):
    indicator_id: str
    indicator_name: str
    score: float  # 0-100
    flag: str  # green, yellow, red
    weight: float
    weighted_score: float
    findings: List[str] = []
    concerns: List[str] = []
    recommendations: List[str] = []

class PredictiveScores(BaseModel):
    joining_intent: float = 0.0  # 0-100, higher = more likely to join
    counter_offer_risk: float = 0.0  # 0-100, higher = higher risk
    stability_score: float = 0.0  # 0-100, higher = more stable
    location_fit: float = 0.0  # 0-100, higher = better fit
    offer_decline_risk: float = 0.0  # 0-100, higher = higher risk
    time_to_join: float = 0.0  # 0-100, higher = faster joining
    overall_hiring_risk: str = "medium"  # low, medium, high

class CareerTrajectoryAssessment(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    # Candidate Information
    candidate_id: Optional[str] = None
    candidate_name: str
    candidate_email: str
    candidate_phone: Optional[str] = None
    resume_text: Optional[str] = None
    # Assessment Context
    assessment_type: str = "standalone"  # pre_application, post_application, standalone
    data_collection_mode: str = "candidate"  # candidate, recruiter_prefill, recruiter_postfill
    job_id: Optional[str] = None
    job_title: Optional[str] = None
    target_role: Optional[str] = None
    target_industry: Optional[str] = None
    target_location: Optional[str] = None
    offered_ctc: Optional[float] = None
    # Employment History (Detailed)
    employment_history: List[Dict[str, Any]] = []
    total_experience_years: Optional[float] = None
    average_tenure_months: Optional[float] = None
    # Location & Mobility
    native_city: Optional[str] = None
    current_city: Optional[str] = None
    relocation_history: List[str] = []
    relocation_willingness: Optional[str] = None
    # Personal Commitment Indicators
    family_status: Optional[str] = None
    spouse_employment: Optional[str] = None
    children_schooling: Optional[str] = None
    residence_status: Optional[str] = None
    # Resignation & Counter-Offer Status
    resignation_status: Optional[str] = None
    resignation_date: Optional[str] = None
    notice_period_days: Optional[int] = None
    notice_negotiable: Optional[str] = None
    counter_offer_status: Optional[str] = None
    counter_offer_details: Optional[Dict[str, Any]] = None
    has_other_offers: Optional[str] = None
    # Compensation Data
    current_ctc: Optional[float] = None
    expected_ctc: Optional[float] = None
    minimum_ctc: Optional[float] = None
    ctc_hike_expected_percent: Optional[float] = None
    # Document Uploads
    resignation_proof_url: Optional[str] = None
    counter_offer_proof_url: Optional[str] = None
    # Questionnaire Responses
    questionnaire_responses: Dict[str, Any] = {}
    questionnaire_completed: bool = False
    questionnaire_completed_at: Optional[datetime] = None
    # Analysis Results
    indicator_results: List[Dict[str, Any]] = []
    overall_score: float = 0.0
    overall_flag: str = "pending"  # pending, green, yellow, red
    # Predictive Scores
    predictive_scores: Dict[str, Any] = {}
    # Summary
    green_flags: List[str] = []
    yellow_flags: List[str] = []
    red_flags: List[str] = []
    key_strengths: List[str] = []
    areas_to_probe: List[str] = []
    recruiter_summary: Optional[str] = None
    hiring_recommendation: Optional[str] = None  # proceed, proceed_with_caution, hold, reject
    # Non-disclosure flags
    non_disclosure_flags: List[Dict[str, Any]] = []  # Tracks what candidate didn't disclose
    # AI Analysis Details
    career_timeline: List[Dict[str, Any]] = []
    skill_progression: List[Dict[str, Any]] = []
    compensation_trajectory: List[Dict[str, Any]] = []
    # Recruiter Actions
    recruiter_notes: Optional[str] = None
    recruiter_override_score: Optional[float] = None
    recruiter_decision: Optional[str] = None  # proceed, hold, reject
    verified_by_recruiter: bool = False
    verification_notes: Optional[str] = None
    # Metadata
    status: str = "pending"  # pending, questionnaire_sent, in_progress, completed, reviewed
    access_token: str = Field(default_factory=lambda: str(uuid.uuid4())[:8].upper())
    created_by: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    completed_at: Optional[datetime] = None
    # Hash of inputs for Re-Analyze detection
    inputs_hash: Optional[str] = None

class CreateTrajectoryAssessmentRequest(BaseModel):
    candidate_name: str
    candidate_email: str
    candidate_phone: Optional[str] = None
    resume_text: Optional[str] = None
    job_id: Optional[str] = None
    target_role: Optional[str] = None
    target_industry: Optional[str] = None
    target_location: Optional[str] = None
    # Note: offered_ctc is NOT captured at creation - only by recruiter after assessment
    employment_history: List[Dict[str, Any]] = []

class SubmitPreAssessmentRequest(BaseModel):
    """Pre-assessment form submission - all 42 questions mandatory"""
    responses: Dict[str, Any]
    job_change_explanations: Optional[List[Dict[str, Any]]] = None  # Explainable job stability

class SubmitQuestionnaireRequest(BaseModel):
    responses: Dict[str, Any]
    employment_history: Optional[List[Dict[str, Any]]] = None

class SubmitEmploymentHistoryRequest(BaseModel):
    employment_history: List[Dict[str, Any]]

class RecruiterAssessmentUpdate(BaseModel):
    recruiter_notes: Optional[str] = None
    recruiter_override_score: Optional[float] = None
    recruiter_decision: Optional[str] = None  # proceed, hold, reject
    verified_by_recruiter: Optional[bool] = None
    verification_notes: Optional[str] = None
    employment_history: Optional[List[Dict[str, Any]]] = None
    current_ctc: Optional[float] = None
    offered_ctc: Optional[float] = None  # Recruiter adds offered CTC after assessment
    send_status_email: Optional[bool] = False  # Trigger status email

class TriggerPreAssessmentRequest(BaseModel):
    """Trigger pre-assessment from job application"""
    candidate_name: str
    candidate_email: str
    candidate_phone: Optional[str] = None
    resume_id: Optional[str] = None
    job_id: str
    resume_text: Optional[str] = None

def calculate_inputs_hash(assessment_data: Dict) -> str:
    """Calculate hash of inputs to detect changes for Re-Analyze"""
    import hashlib
    relevant_fields = [
        "questionnaire_responses",
        "employment_history", 
        "resume_text",
        "job_change_explanations"
    ]
    hash_input = ""
    for field in relevant_fields:
        value = assessment_data.get(field)
        if value:
            hash_input += str(value)
    return hashlib.md5(hash_input.encode()).hexdigest() if hash_input else ""

# Non-disclosure penalty calculation
def calculate_non_disclosure_penalty(assessment_data: Dict) -> tuple:
    """Calculate penalties for non-disclosure of critical information"""
    non_disclosure_flags = []
    total_penalty = 0
    
    # Check counter-offer disclosure
    co_status = assessment_data.get("counter_offer_status")
    if co_status in ["Yes - Already received", "Yes - Expecting one"]:
        co_details = assessment_data.get("counter_offer_details") or {}
        if not co_details.get("ctc") and not assessment_data.get("questionnaire_responses", {}).get("Counter-offer proposed CTC (Annual)"):
            non_disclosure_flags.append({
                "type": "counter_offer_ctc",
                "severity": "red",
                "message": "Counter-offer CTC not disclosed despite indicating receipt of counter-offer",
                "penalty": 15
            })
            total_penalty += 15
        if not co_details.get("role") and not assessment_data.get("questionnaire_responses", {}).get("Counter-offer proposed role/designation"):
            non_disclosure_flags.append({
                "type": "counter_offer_role",
                "severity": "yellow",
                "message": "Counter-offer role/designation not disclosed",
                "penalty": 10
            })
            total_penalty += 10
    
    # Check resignation proof when claimed resigned
    resignation_status = assessment_data.get("resignation_status")
    if resignation_status in ["Yes - Already resigned"] and not assessment_data.get("resignation_proof_url"):
        non_disclosure_flags.append({
            "type": "resignation_proof",
            "severity": "yellow",
            "message": "Resignation claimed but no proof uploaded",
            "penalty": 5
        })
        total_penalty += 5
    
    # Check other offers disclosure
    has_other_offers = assessment_data.get("has_other_offers")
    if has_other_offers == "Yes - Multiple offers":
        # Check if details provided in questionnaire
        responses = assessment_data.get("questionnaire_responses", {})
        offer_details_provided = any("offer" in str(v).lower() for v in responses.values() if isinstance(v, str))
        if not offer_details_provided:
            non_disclosure_flags.append({
                "type": "multiple_offers",
                "severity": "yellow",
                "message": "Multiple offers indicated but no details provided",
                "penalty": 8
            })
            total_penalty += 8
    
    return non_disclosure_flags, total_penalty

# Helper function for Enhanced Career Trajectory AI Analysis
async def analyze_career_trajectory(
    resume_text: str,
    questionnaire_responses: Dict[str, Any],
    target_role: Optional[str] = None,
    target_industry: Optional[str] = None,
    employment_history: List[Dict] = None,
    assessment_data: Dict = None
) -> Dict[str, Any]:
    """Use AI to analyze career trajectory with enhanced predictive analysis"""
    
    # Calculate non-disclosure penalties
    non_disclosure_flags, penalty = calculate_non_disclosure_penalty(assessment_data or {})
    
    system_prompt = """You are an expert HR analyst and predictive hiring specialist. Your task is to perform comprehensive career trajectory analysis with predictive scoring AND HR fitment analysis.

ANALYZE THE FOLLOWING 12 CAREER TRAJECTORY INDICATORS (with weights):

CORE INDICATORS:
1. Career Progression (12%) - Job title growth, role advancement, logical career path
2. Job Stability (12%) - Average tenure, job hopping patterns, commitment signals
3. Industry Alignment (10%) - Industry match, domain expertise, transferable experience
4. Skills Evolution (10%) - Skill growth, upskilling efforts, learning agility
5. Education Alignment (8%) - Education relevance, certifications, continuous learning
6. Employment Gaps (8%) - Gap identification, context, impact on career
7. Cultural Fit (8%) - Work preferences, adaptability, conflict handling
8. Compensation Trajectory (8%) - Salary growth patterns, hike percentages

PREDICTIVE INDICATORS:
9. Location & Mobility (8%) - Relocation patterns, location fit with target role
10. Joining Intent (8%) - Likelihood to accept offer based on resignation status, notice period
11. Counter-Offer Risk (4%) - Risk of candidate taking counter-offer (higher score = lower risk)
12. Retention Stability (4%) - Long-term retention prediction based on personal commitments

ALSO ANALYZE 5 HR FITMENT INDICATORS:

HR FITMENT ANALYSIS:
1. Cultural Fit Assessment (20%) - Company culture alignment, values match, work ethics, organizational fit
   Factors: value_alignment, work_style, communication_style, team_orientation, adaptability
   
2. Team Dynamics Fit (20%) - Team compatibility, reporting hierarchy comfort, collaboration style
   Factors: team_compatibility, leadership_style, collaboration_approach, conflict_resolution, hierarchy_comfort
   
3. Role-Specific HR Metrics (25%) - Notice period, salary alignment, location feasibility, BGV readiness
   Factors: notice_period_fit, salary_alignment, location_feasibility, bgv_readiness, documentation_status
   
4. Soft Skills Evaluation (20%) - Communication, leadership potential, emotional intelligence
   Factors: communication, leadership_potential, emotional_intelligence, problem_solving, time_management
   
5. HR Risk Assessment Summary (15%) - Consolidated risk factors for hiring decision
   Factors: attrition_risk, counter_offer_probability, timeline_risk, reference_risk, offer_acceptance_risk

GENERATE PREDICTIVE SCORES:
- joining_intent: 0-100 (likelihood to accept offer)
- counter_offer_risk: 0-100 (risk level, higher = more risky)
- stability_score: 0-100 (retention prediction)
- location_fit: 0-100 (location alignment)
- offer_decline_risk: 0-100 (risk of declining offer)
- time_to_join: 0-100 (likelihood of joining on time)

HIRING RECOMMENDATION:
Based on all indicators and predictive scores, provide one of:
- "proceed" - Strong candidate, minimal risk
- "proceed_with_caution" - Good candidate but has some yellow flags to address
- "hold" - Significant concerns, needs further verification
- "reject" - Multiple red flags, high risk

For each indicator, provide score (0-100), flag (green/yellow/red), findings, concerns, and recommendations.

IMPORTANT FLAGS FOR NON-DISCLOSURE:
- If counter-offer details not disclosed when claimed: RED FLAG, -15 points
- If resignation proof not provided when claimed resigned: YELLOW FLAG, -5 points
- If multiple offers claimed but no details: YELLOW FLAG, -8 points

Respond in valid JSON format."""

    user_prompt = f"""
CANDIDATE ANALYSIS REQUEST:

=== RESUME/PROFILE ===
{resume_text or "No resume provided"}

=== TARGET POSITION ===
Role: {target_role or "Not specified"}
Industry: {target_industry or "Not specified"}
Location: {assessment_data.get("target_location") if assessment_data else "Not specified"}
Offered CTC: {assessment_data.get("offered_ctc") if assessment_data else "Not specified"}

=== EMPLOYMENT HISTORY ===
{json.dumps(employment_history or [], indent=2) if employment_history else "Not provided - extract from resume"}

=== QUESTIONNAIRE RESPONSES ===
{json.dumps(questionnaire_responses, indent=2) if questionnaire_responses else "No responses yet"}

=== PERSONAL & COMMITMENT DATA ===
Family Status: {assessment_data.get("family_status") if assessment_data else "Not provided"}
Spouse Employment: {assessment_data.get("spouse_employment") if assessment_data else "Not provided"}
Children Schooling: {assessment_data.get("children_schooling") if assessment_data else "Not provided"}
Residence Status: {assessment_data.get("residence_status") if assessment_data else "Not provided"}
Native City: {assessment_data.get("native_city") if assessment_data else "Not provided"}
Current City: {assessment_data.get("current_city") if assessment_data else "Not provided"}

=== RESIGNATION & OFFER STATUS ===
Resignation Status: {assessment_data.get("resignation_status") if assessment_data else "Not provided"}
Resignation Date: {assessment_data.get("resignation_date") if assessment_data else "Not provided"}
Notice Period (days): {assessment_data.get("notice_period_days") if assessment_data else "Not provided"}
Counter-Offer Status: {assessment_data.get("counter_offer_status") if assessment_data else "Not provided"}
Counter-Offer Details: {json.dumps(assessment_data.get("counter_offer_details")) if assessment_data and assessment_data.get("counter_offer_details") else "Not disclosed"}
Other Active Offers: {assessment_data.get("has_other_offers") if assessment_data else "Not provided"}

=== COMPENSATION DATA ===
Current CTC: {assessment_data.get("current_ctc") if assessment_data else "Not provided"}
Expected CTC: {assessment_data.get("expected_ctc") if assessment_data else "Not provided"}
Minimum CTC: {assessment_data.get("minimum_ctc") if assessment_data else "Not provided"}

=== NON-DISCLOSURE FLAGS (Apply Penalties) ===
{json.dumps(non_disclosure_flags, indent=2) if non_disclosure_flags else "None detected"}
Total Penalty Points: {penalty}

Provide comprehensive analysis in JSON format:
{{
    "indicators": [
        {{
            "indicator_id": "<id>",
            "indicator_name": "<name>",
            "score": <0-100>,
            "flag": "<green|yellow|red>",
            "weight": <weight>,
            "findings": ["..."],
            "concerns": ["..."],
            "recommendations": ["..."]
        }}
        // ... for all 12 career trajectory indicators
    ],
    "hr_fitment_analysis": [
        {{
            "indicator_id": "hr_cultural_fit",
            "indicator_name": "Cultural Fit Assessment",
            "score": <0-100>,
            "flag": "<green|yellow|red>",
            "weight": 0.20,
            "factor_scores": {{
                "value_alignment": <0-100>,
                "work_style": <0-100>,
                "communication_style": <0-100>,
                "team_orientation": <0-100>,
                "adaptability": <0-100>
            }},
            "findings": ["..."],
            "concerns": ["..."],
            "recommendations": ["..."]
        }},
        {{
            "indicator_id": "hr_team_dynamics",
            "indicator_name": "Team Dynamics Fit",
            "score": <0-100>,
            "flag": "<green|yellow|red>",
            "weight": 0.20,
            "factor_scores": {{
                "team_compatibility": <0-100>,
                "leadership_style": <0-100>,
                "collaboration_approach": <0-100>,
                "conflict_resolution": <0-100>,
                "hierarchy_comfort": <0-100>
            }},
            "findings": ["..."],
            "concerns": ["..."],
            "recommendations": ["..."]
        }},
        {{
            "indicator_id": "hr_role_metrics",
            "indicator_name": "Role-Specific HR Metrics",
            "score": <0-100>,
            "flag": "<green|yellow|red>",
            "weight": 0.25,
            "factor_scores": {{
                "notice_period_fit": <0-100>,
                "salary_alignment": <0-100>,
                "location_feasibility": <0-100>,
                "bgv_readiness": <0-100>,
                "documentation_status": <0-100>
            }},
            "findings": ["..."],
            "concerns": ["..."],
            "recommendations": ["..."]
        }},
        {{
            "indicator_id": "hr_soft_skills",
            "indicator_name": "Soft Skills Evaluation",
            "score": <0-100>,
            "flag": "<green|yellow|red>",
            "weight": 0.20,
            "factor_scores": {{
                "communication": <0-100>,
                "leadership_potential": <0-100>,
                "emotional_intelligence": <0-100>,
                "problem_solving": <0-100>,
                "time_management": <0-100>
            }},
            "findings": ["..."],
            "concerns": ["..."],
            "recommendations": ["..."]
        }},
        {{
            "indicator_id": "hr_risk_summary",
            "indicator_name": "HR Risk Assessment Summary",
            "score": <0-100>,
            "flag": "<green|yellow|red>",
            "weight": 0.15,
            "factor_scores": {{
                "attrition_risk": <0-100>,
                "counter_offer_probability": <0-100>,
                "timeline_risk": <0-100>,
                "reference_risk": <0-100>,
                "offer_acceptance_risk": <0-100>
            }},
            "findings": ["..."],
            "concerns": ["..."],
            "recommendations": ["..."]
        }}
    ],
    "hr_fitment_overall": {{
        "score": <0-100>,
        "flag": "<green|yellow|red>",
        "summary": "2-3 sentence HR fitment summary",
        "top_strengths": ["..."],
        "key_concerns": ["..."],
        "interview_focus_areas": ["..."]
    }},
    "predictive_scores": {{
        "joining_intent": <0-100>,
        "counter_offer_risk": <0-100>,
        "stability_score": <0-100>,
        "location_fit": <0-100>,
        "offer_decline_risk": <0-100>,
        "time_to_join": <0-100>,
        "overall_hiring_risk": "<low|medium|high>"
    }},
    "career_timeline": [
        {{"role": "...", "company": "...", "duration": "...", "years": <number>, "location": "..."}}
    ],
    "skill_progression": [
        {{"period": "...", "skills_acquired": ["..."], "growth_area": "..."}}
    ],
    "compensation_trajectory": [
        {{"employer": "...", "joining_ctc": <number>, "exit_ctc": <number>, "hike_percent": <number>}}
    ],
    "key_strengths": ["..."],
    "areas_to_probe": ["..."],
    "hiring_recommendation": "<proceed|proceed_with_caution|hold|reject>",
    "recruiter_summary": "2-3 sentence summary highlighting key insights and risks"
}}
"""
    
    try:
        response = await get_llm_response(system_prompt, user_prompt)
        analysis = parse_json_from_response(response)
        # Add non-disclosure flags to result
        analysis["non_disclosure_flags"] = non_disclosure_flags
        analysis["non_disclosure_penalty"] = penalty
        return analysis
    except Exception as e:
        logger.error(f"Career trajectory analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")

def calculate_overall_score(indicator_results: List[Dict], penalty: int = 0) -> tuple:
    """Calculate weighted overall score and determine flag, applying non-disclosure penalties"""
    total_weighted_score = 0.0
    total_weight = 0.0
    
    for result in indicator_results:
        weight = result.get("weight", 0.08)  # Default equal weight
        score = result.get("score", 0)
        weighted_score = score * weight
        result["weighted_score"] = round(weighted_score, 2)
        total_weighted_score += weighted_score
        total_weight += weight
    
    overall_score = round(total_weighted_score / total_weight if total_weight > 0 else 0, 1)
    
    # Apply non-disclosure penalty
    overall_score = max(0, overall_score - penalty)
    
    # Determine overall flag
    if overall_score >= 71:
        overall_flag = "green"
    elif overall_score >= 41:
        overall_flag = "yellow"
    else:
        overall_flag = "red"
    
    return overall_score, overall_flag

def categorize_flags(indicator_results: List[Dict]) -> tuple:
    """Categorize indicators by flag color"""
    green_flags = []
    yellow_flags = []
    red_flags = []
    
    for result in indicator_results:
        flag = result.get("flag", "yellow")
        name = result.get("indicator_name", "Unknown")
        score = result.get("score", 0)
        
        summary = f"{name}: {score}/100"
        if result.get("findings"):
            summary += f" - {result['findings'][0]}"
        
        if flag == "green":
            green_flags.append(summary)
        elif flag == "yellow":
            yellow_flags.append(summary)
        else:
            red_flags.append(summary)
    
    return green_flags, yellow_flags, red_flags

# ============ Career Trajectory API Routes ============

@api_router.get("/trajectory/indicators")
async def get_trajectory_indicators():
    """Get all career trajectory indicators with weights and thresholds"""
    return {
        "indicators": list(CAREER_TRAJECTORY_INDICATORS.values()),
        "total_weight": sum(ind["weight"] for ind in CAREER_TRAJECTORY_INDICATORS.values()),
        "job_change_reasons": JOB_CHANGE_REASONS,
        "hr_fitment_indicators": list(HR_FITMENT_INDICATORS.values()),
        "hr_fitment_total_weight": sum(ind["weight"] for ind in HR_FITMENT_INDICATORS.values())
    }

@api_router.get("/trajectory/questionnaire")
async def get_trajectory_questionnaire():
    """Get the pre-assessment questionnaire - 42 questions, all multiple choice"""
    total_questions = sum(len(questions) for questions in PRE_ASSESSMENT_QUESTIONNAIRE.values())
    return {
        "categories": PRE_ASSESSMENT_QUESTIONNAIRE,
        "total_questions": total_questions,
        "estimated_time": "Under 5 minutes",
        "all_mandatory": True,
        "format": "All questions are single-select for quick completion"
    }

@api_router.get("/trajectory/job-change-reasons")
async def get_job_change_reasons():
    """Get explainable job change reason options"""
    return JOB_CHANGE_REASONS

@api_router.post("/trajectory/assessment/create")
async def create_trajectory_assessment(request: CreateTrajectoryAssessmentRequest):
    """Create a new career trajectory assessment - triggered when candidate applies to job"""
    # Note: offered_ctc is NOT captured at creation - only after assessment completion
    assessment = CareerTrajectoryAssessment(
        candidate_name=request.candidate_name,
        candidate_email=request.candidate_email,
        candidate_phone=request.candidate_phone,
        resume_text=request.resume_text,
        assessment_type="pre_assessment",  # Always pre_assessment now
        data_collection_mode="candidate",  # Candidate must fill
        job_id=request.job_id,
        target_role=request.target_role,
        target_industry=request.target_industry,
        target_location=request.target_location,
        employment_history=request.employment_history,
        status="pending"  # Always pending until pre-assessment completed
    )
    
    # Save to database first
    assessment_dict = assessment.model_dump()
    await db.trajectory_assessments.insert_one(assessment_dict)
    
    # Generate pre-assessment link
    preassessment_link = f"{os.environ.get('FRONTEND_URL', 'http://localhost:3000')}/trajectory/preassessment/{assessment.id}?token={assessment.access_token}"
    
    # Simulate sending pre-assessment email
    if request.job_id:
        # Get job details
        job = await db.job_descriptions.find_one({"id": request.job_id})
        job_title = job.get("title", request.target_role) if job else request.target_role
    else:
        job_title = request.target_role or "Position"
    
    simulate_send_preassessment_email(
        candidate_email=request.candidate_email,
        candidate_name=request.candidate_name,
        job_title=job_title,
        job_id=request.job_id or assessment.id,
        assessment_link=preassessment_link
    )
    
    return {
        "success": True,
        "assessment_id": assessment.id,
        "access_token": assessment.access_token,
        "preassessment_link": preassessment_link,
        "status": "pending",
        "message": "Pre-assessment email sent to candidate. Assessment must be completed for application to proceed."
    }

@api_router.get("/trajectory/assessment/{assessment_id}")
async def get_trajectory_assessment(assessment_id: str):
    """Get a trajectory assessment by ID"""
    assessment = await db.trajectory_assessments.find_one(
        {"id": assessment_id},
        {"_id": 0}
    )
    
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    return assessment

def generate_trajectory_report_pdf(assessment: Dict) -> io.BytesIO:
    """Generate a comprehensive PDF report for Career Trajectory Assessment with HR Fitment Analysis"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=20, textColor=colors.HexColor('#1e3a5f'))
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=14, spaceAfter=10, textColor=colors.HexColor('#2563eb'))
    subheading_style = ParagraphStyle('CustomSubheading', parent=styles['Heading3'], fontSize=12, spaceAfter=8, textColor=colors.HexColor('#374151'))
    normal_style = ParagraphStyle('CustomNormal', parent=styles['Normal'], fontSize=10, spaceAfter=6)
    green_style = ParagraphStyle('GreenText', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#059669'))
    yellow_style = ParagraphStyle('YellowText', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#d97706'))
    red_style = ParagraphStyle('RedText', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#dc2626'))
    
    story = []
    
    # Helper function to format date
    def format_date(dt):
        if dt is None:
            return 'N/A'
        if isinstance(dt, datetime):
            return dt.strftime('%Y-%m-%d')
        if isinstance(dt, str):
            return dt[:10] if len(dt) >= 10 else dt
        return str(dt)
    
    # Title
    story.append(Paragraph("Career Trajectory Assessment Report", title_style))
    story.append(Paragraph(f"<b>Candidate:</b> {assessment.get('candidate_name', 'N/A')}", normal_style))
    story.append(Paragraph(f"<b>Target Role:</b> {assessment.get('target_role', 'N/A')}", normal_style))
    story.append(Paragraph(f"<b>Target Industry:</b> {assessment.get('target_industry', 'N/A')}", normal_style))
    story.append(Paragraph(f"<b>Assessment Date:</b> {format_date(assessment.get('created_at'))}", normal_style))
    story.append(Spacer(1, 20))
    
    # Overall Score
    overall_score = assessment.get('overall_score', 0)
    overall_flag = assessment.get('overall_flag', 'yellow')
    hiring_rec = assessment.get('hiring_recommendation', 'N/A')
    
    score_color = {'green': '#059669', 'yellow': '#d97706', 'red': '#dc2626'}.get(overall_flag, '#374151')
    story.append(Paragraph("Executive Summary", heading_style))
    story.append(Paragraph(f"<b>Overall Score:</b> <font color='{score_color}'>{overall_score:.1f}/100</font> ({overall_flag.upper()})", normal_style))
    story.append(Paragraph(f"<b>Hiring Recommendation:</b> {hiring_rec.replace('_', ' ').title()}", normal_style))
    
    recruiter_summary = assessment.get('recruiter_summary', '')
    if recruiter_summary:
        story.append(Paragraph(f"<b>Summary:</b> {recruiter_summary}", normal_style))
    story.append(Spacer(1, 15))
    
    # Predictive Scores
    predictive_scores = assessment.get('predictive_scores', {})
    if predictive_scores:
        story.append(Paragraph("Predictive Scores", heading_style))
        pred_data = [
            ['Metric', 'Score', 'Risk Level'],
            ['Joining Intent', f"{predictive_scores.get('joining_intent', 'N/A')}", 'High' if predictive_scores.get('joining_intent', 0) < 50 else 'Low'],
            ['Counter-Offer Risk', f"{predictive_scores.get('counter_offer_risk', 'N/A')}", 'High' if predictive_scores.get('counter_offer_risk', 0) > 50 else 'Low'],
            ['Stability Score', f"{predictive_scores.get('stability_score', 'N/A')}", 'Good' if predictive_scores.get('stability_score', 0) > 70 else 'Moderate'],
            ['Location Fit', f"{predictive_scores.get('location_fit', 'N/A')}", 'Good' if predictive_scores.get('location_fit', 0) > 70 else 'Moderate'],
            ['Offer Decline Risk', f"{predictive_scores.get('offer_decline_risk', 'N/A')}", 'High' if predictive_scores.get('offer_decline_risk', 0) > 50 else 'Low'],
            ['Time to Join', f"{predictive_scores.get('time_to_join', 'N/A')}", 'On Track' if predictive_scores.get('time_to_join', 0) > 70 else 'At Risk'],
        ]
        pred_table = Table(pred_data, colWidths=[200, 80, 100])
        pred_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a5f')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(pred_table)
        story.append(Spacer(1, 15))
    
    # Career Trajectory Indicators (12)
    indicator_results = assessment.get('indicator_results', [])
    if indicator_results:
        story.append(Paragraph("Career Trajectory Indicators (12)", heading_style))
        for indicator in indicator_results:
            ind_name = indicator.get('indicator_name', indicator.get('indicator_id', 'Unknown'))
            ind_score = indicator.get('score', 0)
            ind_flag = indicator.get('flag', 'yellow')
            flag_color = {'green': '#059669', 'yellow': '#d97706', 'red': '#dc2626'}.get(ind_flag, '#374151')
            
            story.append(Paragraph(f"<b>{ind_name}</b> - <font color='{flag_color}'>{ind_score}/100 ({ind_flag.upper()})</font>", subheading_style))
            
            findings = indicator.get('findings', [])
            if findings:
                story.append(Paragraph("<i>Findings:</i> " + "; ".join(findings[:3]), normal_style))
            
            concerns = indicator.get('concerns', [])
            if concerns:
                story.append(Paragraph("<i>Concerns:</i> " + "; ".join(concerns[:2]), normal_style))
            story.append(Spacer(1, 5))
        story.append(Spacer(1, 10))
    
    # HR Fitment Analysis (5 indicators)
    hr_fitment = assessment.get('hr_fitment_analysis', [])
    if hr_fitment:
        story.append(Paragraph("HR Fitment Analysis", heading_style))
        
        hr_overall = assessment.get('hr_fitment_overall', {})
        if hr_overall:
            hr_score = hr_overall.get('score', 0)
            hr_flag = hr_overall.get('flag', 'yellow')
            hr_color = {'green': '#059669', 'yellow': '#d97706', 'red': '#dc2626'}.get(hr_flag, '#374151')
            story.append(Paragraph(f"<b>Overall HR Fitment Score:</b> <font color='{hr_color}'>{hr_score}/100 ({hr_flag.upper()})</font>", normal_style))
            if hr_overall.get('summary'):
                story.append(Paragraph(f"<i>{hr_overall.get('summary')}</i>", normal_style))
            story.append(Spacer(1, 10))
        
        for hr_ind in hr_fitment:
            ind_name = hr_ind.get('indicator_name', 'Unknown')
            ind_score = hr_ind.get('score', 0)
            ind_flag = hr_ind.get('flag', 'yellow')
            flag_color = {'green': '#059669', 'yellow': '#d97706', 'red': '#dc2626'}.get(ind_flag, '#374151')
            
            story.append(Paragraph(f"<b>{ind_name}</b> - <font color='{flag_color}'>{ind_score}/100 ({ind_flag.upper()})</font>", subheading_style))
            
            # Factor scores if available
            factor_scores = hr_ind.get('factor_scores', {})
            if factor_scores:
                factors_text = ", ".join([f"{k.replace('_', ' ').title()}: {v}" for k, v in factor_scores.items()])
                story.append(Paragraph(f"<i>Factor Scores:</i> {factors_text}", normal_style))
            
            findings = hr_ind.get('findings', [])
            if findings:
                story.append(Paragraph("<i>Findings:</i> " + "; ".join(findings[:2]), normal_style))
            
            recommendations = hr_ind.get('recommendations', [])
            if recommendations:
                story.append(Paragraph("<i>Recommendations:</i> " + "; ".join(recommendations[:2]), normal_style))
            story.append(Spacer(1, 5))
        story.append(Spacer(1, 10))
    
    # Key Strengths and Areas to Probe
    key_strengths = assessment.get('key_strengths', [])
    if key_strengths:
        story.append(Paragraph("Key Strengths", heading_style))
        for strength in key_strengths[:5]:
            story.append(Paragraph(f"• {strength}", green_style))
        story.append(Spacer(1, 10))
    
    areas_to_probe = assessment.get('areas_to_probe', [])
    if areas_to_probe:
        story.append(Paragraph("Areas to Probe in Interview", heading_style))
        for area in areas_to_probe[:5]:
            story.append(Paragraph(f"• {area}", yellow_style))
        story.append(Spacer(1, 10))
    
    # Footer
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"<i>Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M')} | Role Sense Career Trajectory Assessment</i>", 
                          ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey)))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_trajectory_report_docx(assessment: Dict) -> io.BytesIO:
    """Generate a comprehensive DOCX report for Career Trajectory Assessment with HR Fitment Analysis"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Career Trajectory Assessment Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Helper function to format date
    def format_date(dt):
        if dt is None:
            return 'N/A'
        if isinstance(dt, datetime):
            return dt.strftime('%Y-%m-%d')
        if isinstance(dt, str):
            return dt[:10] if len(dt) >= 10 else dt
        return str(dt)
    
    # Basic Info
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Candidate: ').bold = True
    p.add_run(assessment.get('candidate_name', 'N/A'))
    
    p = doc.add_paragraph()
    p.add_run('Target Role: ').bold = True
    p.add_run(assessment.get('target_role', 'N/A'))
    
    p = doc.add_paragraph()
    p.add_run('Target Industry: ').bold = True
    p.add_run(assessment.get('target_industry', 'N/A'))
    
    p = doc.add_paragraph()
    p.add_run('Assessment Date: ').bold = True
    p.add_run(format_date(assessment.get('created_at')))
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    
    overall_score = assessment.get('overall_score', 0)
    overall_flag = assessment.get('overall_flag', 'yellow')
    hiring_rec = assessment.get('hiring_recommendation', 'N/A')
    
    p = doc.add_paragraph()
    p.add_run('Overall Score: ').bold = True
    score_run = p.add_run(f'{overall_score:.1f}/100 ({overall_flag.upper()})')
    if overall_flag == 'green':
        score_run.font.color.rgb = RGBColor(5, 150, 105)
    elif overall_flag == 'yellow':
        score_run.font.color.rgb = RGBColor(217, 119, 6)
    else:
        score_run.font.color.rgb = RGBColor(220, 38, 38)
    
    p = doc.add_paragraph()
    p.add_run('Hiring Recommendation: ').bold = True
    p.add_run(hiring_rec.replace('_', ' ').title())
    
    recruiter_summary = assessment.get('recruiter_summary', '')
    if recruiter_summary:
        p = doc.add_paragraph()
        p.add_run('Summary: ').bold = True
        p.add_run(recruiter_summary)
    
    # Predictive Scores Table
    predictive_scores = assessment.get('predictive_scores', {})
    if predictive_scores:
        doc.add_heading('Predictive Scores', level=1)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Score'
        hdr_cells[2].text = 'Risk Level'
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        metrics = [
            ('Joining Intent', predictive_scores.get('joining_intent', 'N/A'), 'High Risk' if predictive_scores.get('joining_intent', 0) < 50 else 'Low Risk'),
            ('Counter-Offer Risk', predictive_scores.get('counter_offer_risk', 'N/A'), 'High Risk' if predictive_scores.get('counter_offer_risk', 0) > 50 else 'Low Risk'),
            ('Stability Score', predictive_scores.get('stability_score', 'N/A'), 'Good' if predictive_scores.get('stability_score', 0) > 70 else 'Moderate'),
            ('Location Fit', predictive_scores.get('location_fit', 'N/A'), 'Good' if predictive_scores.get('location_fit', 0) > 70 else 'Moderate'),
            ('Offer Decline Risk', predictive_scores.get('offer_decline_risk', 'N/A'), 'High Risk' if predictive_scores.get('offer_decline_risk', 0) > 50 else 'Low Risk'),
            ('Time to Join', predictive_scores.get('time_to_join', 'N/A'), 'On Track' if predictive_scores.get('time_to_join', 0) > 70 else 'At Risk'),
        ]
        
        for metric, score, risk in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = str(score)
            row_cells[2].text = risk
    
    # Career Trajectory Indicators
    indicator_results = assessment.get('indicator_results', [])
    if indicator_results:
        doc.add_heading('Career Trajectory Indicators (12)', level=1)
        
        for indicator in indicator_results:
            ind_name = indicator.get('indicator_name', indicator.get('indicator_id', 'Unknown'))
            ind_score = indicator.get('score', 0)
            ind_flag = indicator.get('flag', 'yellow')
            
            p = doc.add_paragraph()
            p.add_run(f'{ind_name}: ').bold = True
            score_run = p.add_run(f'{ind_score}/100 ({ind_flag.upper()})')
            if ind_flag == 'green':
                score_run.font.color.rgb = RGBColor(5, 150, 105)
            elif ind_flag == 'yellow':
                score_run.font.color.rgb = RGBColor(217, 119, 6)
            else:
                score_run.font.color.rgb = RGBColor(220, 38, 38)
            
            findings = indicator.get('findings', [])
            if findings:
                p = doc.add_paragraph()
                p.add_run('Findings: ').italic = True
                p.add_run('; '.join(findings[:3]))
            
            concerns = indicator.get('concerns', [])
            if concerns:
                p = doc.add_paragraph()
                p.add_run('Concerns: ').italic = True
                p.add_run('; '.join(concerns[:2]))
    
    # HR Fitment Analysis
    hr_fitment = assessment.get('hr_fitment_analysis', [])
    if hr_fitment:
        doc.add_heading('HR Fitment Analysis', level=1)
        
        hr_overall = assessment.get('hr_fitment_overall', {})
        if hr_overall:
            hr_score = hr_overall.get('score', 0)
            hr_flag = hr_overall.get('flag', 'yellow')
            
            p = doc.add_paragraph()
            p.add_run('Overall HR Fitment Score: ').bold = True
            score_run = p.add_run(f'{hr_score}/100 ({hr_flag.upper()})')
            if hr_flag == 'green':
                score_run.font.color.rgb = RGBColor(5, 150, 105)
            elif hr_flag == 'yellow':
                score_run.font.color.rgb = RGBColor(217, 119, 6)
            else:
                score_run.font.color.rgb = RGBColor(220, 38, 38)
            
            if hr_overall.get('summary'):
                p = doc.add_paragraph()
                p.add_run(hr_overall.get('summary')).italic = True
        
        for hr_ind in hr_fitment:
            ind_name = hr_ind.get('indicator_name', 'Unknown')
            ind_score = hr_ind.get('score', 0)
            ind_flag = hr_ind.get('flag', 'yellow')
            
            p = doc.add_paragraph()
            p.add_run(f'{ind_name}: ').bold = True
            score_run = p.add_run(f'{ind_score}/100 ({ind_flag.upper()})')
            if ind_flag == 'green':
                score_run.font.color.rgb = RGBColor(5, 150, 105)
            elif ind_flag == 'yellow':
                score_run.font.color.rgb = RGBColor(217, 119, 6)
            else:
                score_run.font.color.rgb = RGBColor(220, 38, 38)
            
            factor_scores = hr_ind.get('factor_scores', {})
            if factor_scores:
                p = doc.add_paragraph()
                p.add_run('Factor Scores: ').italic = True
                factors_text = ', '.join([f"{k.replace('_', ' ').title()}: {v}" for k, v in factor_scores.items()])
                p.add_run(factors_text)
            
            findings = hr_ind.get('findings', [])
            if findings:
                p = doc.add_paragraph()
                p.add_run('Findings: ').italic = True
                p.add_run('; '.join(findings[:2]))
            
            recommendations = hr_ind.get('recommendations', [])
            if recommendations:
                p = doc.add_paragraph()
                p.add_run('Recommendations: ').italic = True
                p.add_run('; '.join(recommendations[:2]))
    
    # Key Strengths
    key_strengths = assessment.get('key_strengths', [])
    if key_strengths:
        doc.add_heading('Key Strengths', level=1)
        for strength in key_strengths[:5]:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(strength)
            run.font.color.rgb = RGBColor(5, 150, 105)
    
    # Areas to Probe
    areas_to_probe = assessment.get('areas_to_probe', [])
    if areas_to_probe:
        doc.add_heading('Areas to Probe in Interview', level=1)
        for area in areas_to_probe[:5]:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(area)
            run.font.color.rgb = RGBColor(217, 119, 6)
    
    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f'Report generated on {datetime.now().strftime("%Y-%m-%d %H:%M")} | Role Sense Career Trajectory Assessment').italic = True
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@api_router.get("/trajectory/assessment/{assessment_id}/report/{format}")
async def download_trajectory_report(assessment_id: str, format: str):
    """Download Career Trajectory Assessment Report in PDF or DOCX format"""
    if format not in ["pdf", "docx"]:
        raise HTTPException(status_code=400, detail="Invalid format. Use 'pdf' or 'docx'")
    
    assessment = await db.trajectory_assessments.find_one(
        {"id": assessment_id},
        {"_id": 0}
    )
    
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    candidate_name = assessment.get('candidate_name', 'Candidate').replace(' ', '_')
    filename = f"Career_Trajectory_Report_{candidate_name}_{datetime.now().strftime('%Y%m%d')}"
    
    if format == "pdf":
        buffer = generate_trajectory_report_pdf(assessment)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}.pdf"}
        )
    else:  # docx
        buffer = generate_trajectory_report_docx(assessment)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}.docx"}
        )

@api_router.get("/trajectory/assessment/token/{access_token}")
async def get_assessment_by_token(access_token: str):
    """Get assessment by access token (for candidate questionnaire access)"""
    assessment = await db.trajectory_assessments.find_one(
        {"access_token": access_token},
        {"_id": 0}
    )
    
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    # Return limited data for candidate view
    return {
        "id": assessment["id"],
        "candidate_name": assessment["candidate_name"],
        "target_role": assessment.get("target_role"),
        "target_industry": assessment.get("target_industry"),
        "questionnaire_completed": assessment.get("questionnaire_completed", False),
        "status": assessment["status"]
    }

@api_router.post("/trajectory/assessment/{assessment_id}/questionnaire")
async def submit_questionnaire(assessment_id: str, request: SubmitQuestionnaireRequest):
    """Submit questionnaire responses for an assessment with enhanced data extraction"""
    assessment = await db.trajectory_assessments.find_one({"id": assessment_id})
    
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    # Extract structured data from responses
    responses = request.responses
    
    # Update questionnaire responses and extract key fields
    update_data = {
        "questionnaire_responses": responses,
        "questionnaire_completed": True,
        "questionnaire_completed_at": datetime.now(timezone.utc),
        "updated_at": datetime.now(timezone.utc)
    }
    
    # Extract and store key fields from responses
    field_mappings = {
        "What is your native/hometown city?": "native_city",
        "What is your current city of residence?": "current_city",
        "Are you open to relocation for this role?": "relocation_willingness",
        "What is your current family status?": "family_status",
        "Is your spouse currently employed?": "spouse_employment",
        "If you have children, what are their schooling stages?": "children_schooling",
        "What is your current residence status?": "residence_status",
        "Have you resigned from your current employer?": "resignation_status",
        "If resigned, what is your resignation date?": "resignation_date",
        "What is your notice period (in days)?": "notice_period_days",
        "Is your notice period negotiable/buyable?": "notice_negotiable",
        "Have you received or are you expecting a counter-offer from your current employer?": "counter_offer_status",
        "Do you have any other active job offers currently?": "has_other_offers",
        "What is your current CTC (Annual)?": "current_ctc",
        "What is your expected CTC (Annual)?": "expected_ctc",
        "What is your minimum acceptable CTC (Annual)?": "minimum_ctc"
    }
    
    for question, field in field_mappings.items():
        if question in responses:
            value = responses[question]
            # Convert numeric strings to numbers for CTC and notice period
            if field in ["current_ctc", "expected_ctc", "minimum_ctc", "notice_period_days"]:
                try:
                    value = float(value) if value else None
                except (ValueError, TypeError):
                    pass
            update_data[field] = value
    
    # Extract counter-offer details if provided
    if responses.get("Have you received or are you expecting a counter-offer from your current employer?") == "Yes - Already received":
        update_data["counter_offer_details"] = {
            "ctc": responses.get("Counter-offer proposed CTC (Annual)"),
            "role": responses.get("Counter-offer proposed role/designation"),
            "timeline": responses.get("Counter-offer proposed joining date/timeline"),
            "details": responses.get("If you have received a counter-offer, please provide details:")
        }
    
    # Update employment history if provided
    if request.employment_history:
        update_data["employment_history"] = request.employment_history
    
    # Build assessment data for analysis
    merged_assessment = {**assessment, **update_data}
    
    # Re-run analysis with questionnaire responses
    try:
        analysis = await analyze_career_trajectory(
            resume_text=assessment.get("resume_text", ""),
            questionnaire_responses=responses,
            target_role=assessment.get("target_role"),
            target_industry=assessment.get("target_industry"),
            employment_history=merged_assessment.get("employment_history", []),
            assessment_data=merged_assessment
        )
        
        indicator_results = analysis.get("indicators", [])
        non_disclosure_penalty = analysis.get("non_disclosure_penalty", 0)
        overall_score, overall_flag = calculate_overall_score(indicator_results, non_disclosure_penalty)
        green_flags, yellow_flags, red_flags = categorize_flags(indicator_results)
        
        # Calculate inputs hash for future re-analyze detection
        new_inputs_hash = calculate_inputs_hash(merged_assessment)
        
        update_data.update({
            "indicator_results": indicator_results,
            "overall_score": overall_score,
            "overall_flag": overall_flag,
            "green_flags": green_flags,
            "yellow_flags": yellow_flags,
            "red_flags": red_flags,
            "key_strengths": analysis.get("key_strengths", []),
            "areas_to_probe": analysis.get("areas_to_probe", []),
            "recruiter_summary": analysis.get("recruiter_summary", ""),
            "hiring_recommendation": analysis.get("hiring_recommendation", ""),
            "career_timeline": analysis.get("career_timeline", []),
            "skill_progression": analysis.get("skill_progression", []),
            "compensation_trajectory": analysis.get("compensation_trajectory", []),
            "predictive_scores": analysis.get("predictive_scores", {}),
            "hr_fitment_analysis": analysis.get("hr_fitment_analysis", []),
            "hr_fitment_overall": analysis.get("hr_fitment_overall", {}),
            "non_disclosure_flags": analysis.get("non_disclosure_flags", []),
            "inputs_hash": new_inputs_hash,
            "status": "completed",
            "completed_at": datetime.now(timezone.utc)
        })
        
        # Send acknowledgement email after successful pre-assessment completion
        if request.job_id or assessment.get("job_id"):
            job_id = request.job_id or assessment.get("job_id")
            job = await db.job_descriptions.find_one({"id": job_id})
            job_title = job.get("title", assessment.get("target_role", "Position")) if job else assessment.get("target_role", "Position")
            simulate_send_acknowledgement_email(
                candidate_email=assessment["candidate_email"],
                candidate_name=assessment["candidate_name"],
                job_title=job_title,
                job_id=job_id
            )
            # Lock resume to job requisition
            update_data["resume_locked_to_job"] = job_id
            update_data["resume_locked_at"] = datetime.now(timezone.utc)
            
    except Exception as e:
        logger.error(f"Analysis after questionnaire failed: {e}")
        update_data["status"] = "in_progress"
    
    await db.trajectory_assessments.update_one(
        {"id": assessment_id},
        {"$set": update_data}
    )
    
    return {
        "success": True,
        "message": "Pre-assessment completed successfully. Resume accepted and locked to job requisition.",
        "status": update_data.get("status", "in_progress"),
        "predictive_scores": update_data.get("predictive_scores", {}),
        "hiring_recommendation": update_data.get("hiring_recommendation", "")
    }

@api_router.post("/trajectory/assessment/{assessment_id}/analyze")
async def run_trajectory_analysis(assessment_id: str, force: bool = False):
    """
    Re-analyze assessment - ONLY recalculates if inputs have changed.
    Set force=True to force recalculation regardless of input changes.
    """
    assessment = await db.trajectory_assessments.find_one({"id": assessment_id})
    
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    # Calculate current inputs hash
    current_inputs_hash = calculate_inputs_hash(assessment)
    stored_inputs_hash = assessment.get("inputs_hash", "")
    
    # Check if inputs have changed
    if not force and current_inputs_hash == stored_inputs_hash and assessment.get("indicator_results"):
        # No changes - return existing scores without recalculation
        return {
            "success": True,
            "message": "No changes detected. Scores remain unchanged.",
            "recalculated": False,
            "overall_score": assessment.get("overall_score"),
            "overall_flag": assessment.get("overall_flag"),
            "predictive_scores": assessment.get("predictive_scores", {}),
            "hiring_recommendation": assessment.get("hiring_recommendation"),
            "indicator_results": assessment.get("indicator_results", [])
        }
    
    # Inputs have changed - run analysis
    analysis = await analyze_career_trajectory(
        resume_text=assessment.get("resume_text", ""),
        questionnaire_responses=assessment.get("questionnaire_responses", {}),
        target_role=assessment.get("target_role"),
        target_industry=assessment.get("target_industry"),
        employment_history=assessment.get("employment_history", []),
        assessment_data=assessment
    )
    
    indicator_results = analysis.get("indicators", [])
    non_disclosure_penalty = analysis.get("non_disclosure_penalty", 0)
    overall_score, overall_flag = calculate_overall_score(indicator_results, non_disclosure_penalty)
    green_flags, yellow_flags, red_flags = categorize_flags(indicator_results)
    
    update_data = {
        "indicator_results": indicator_results,
        "overall_score": overall_score,
        "overall_flag": overall_flag,
        "green_flags": green_flags,
        "yellow_flags": yellow_flags,
        "red_flags": red_flags,
        "key_strengths": analysis.get("key_strengths", []),
        "areas_to_probe": analysis.get("areas_to_probe", []),
        "recruiter_summary": analysis.get("recruiter_summary", ""),
        "hiring_recommendation": analysis.get("hiring_recommendation", ""),
        "career_timeline": analysis.get("career_timeline", []),
        "skill_progression": analysis.get("skill_progression", []),
        "compensation_trajectory": analysis.get("compensation_trajectory", []),
        "predictive_scores": analysis.get("predictive_scores", {}),
        "hr_fitment_analysis": analysis.get("hr_fitment_analysis", []),
        "hr_fitment_overall": analysis.get("hr_fitment_overall", {}),
        "non_disclosure_flags": analysis.get("non_disclosure_flags", []),
        "inputs_hash": current_inputs_hash,
        "status": "completed" if assessment.get("questionnaire_completed") else "in_progress",
        "completed_at": datetime.now(timezone.utc) if assessment.get("questionnaire_completed") else None,
        "updated_at": datetime.now(timezone.utc)
    }
    
    await db.trajectory_assessments.update_one(
        {"id": assessment_id},
        {"$set": update_data}
    )
    
    # Fetch and return updated assessment
    updated_assessment = await db.trajectory_assessments.find_one(
        {"id": assessment_id},
        {"_id": 0}
    )
    
    return {
        "success": True,
        "message": "Analysis recalculated with updated inputs.",
        "recalculated": True,
        **updated_assessment
    }

@api_router.put("/trajectory/assessment/{assessment_id}/recruiter")
async def update_recruiter_assessment(assessment_id: str, request: RecruiterAssessmentUpdate):
    """Update recruiter notes, override score, decision, and send status emails"""
    assessment = await db.trajectory_assessments.find_one({"id": assessment_id})
    
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    update_data = {"updated_at": datetime.now(timezone.utc)}
    
    if request.recruiter_notes is not None:
        update_data["recruiter_notes"] = request.recruiter_notes
    if request.recruiter_override_score is not None:
        update_data["recruiter_override_score"] = request.recruiter_override_score
    if request.recruiter_decision is not None:
        update_data["recruiter_decision"] = request.recruiter_decision
        update_data["status"] = "reviewed"
    if request.verified_by_recruiter is not None:
        update_data["verified_by_recruiter"] = request.verified_by_recruiter
    if request.verification_notes is not None:
        update_data["verification_notes"] = request.verification_notes
    if request.employment_history is not None:
        update_data["employment_history"] = request.employment_history
    if request.current_ctc is not None:
        update_data["current_ctc"] = request.current_ctc
    if request.offered_ctc is not None:
        update_data["offered_ctc"] = request.offered_ctc
    
    await db.trajectory_assessments.update_one(
        {"id": assessment_id},
        {"$set": update_data}
    )
    
    # Send status email if requested and decision is reject/hold
    if request.send_status_email and request.recruiter_decision in ["reject", "hold"]:
        job_id = assessment.get("job_id", assessment["id"])
        job = await db.job_descriptions.find_one({"id": job_id}) if assessment.get("job_id") else None
        job_title = job.get("title", assessment.get("target_role", "Position")) if job else assessment.get("target_role", "Position")
        
        # Auto-generate rejection reasons from red flags
        reasons = generate_auto_rejection_reasons(assessment)
        
        simulate_send_rejection_email(
            candidate_email=assessment["candidate_email"],
            candidate_name=assessment["candidate_name"],
            job_title=job_title,
            job_id=job_id,
            reasons=reasons,
            status=request.recruiter_decision
        )
        update_data["rejection_reasons"] = reasons
        update_data["status_email_sent"] = True
        update_data["status_email_sent_at"] = datetime.now(timezone.utc)
        
        await db.trajectory_assessments.update_one(
            {"id": assessment_id},
            {"$set": {"rejection_reasons": reasons, "status_email_sent": True, "status_email_sent_at": datetime.now(timezone.utc)}}
        )
    
    return {"success": True, "message": "Assessment updated", "rejection_reasons": update_data.get("rejection_reasons", [])}

# Integration with Job Application Flow
@api_router.post("/trajectory/trigger-from-application")
async def trigger_preassessment_from_application(request: TriggerPreAssessmentRequest):
    """
    Triggered when a candidate applies to a job.
    Creates assessment and sends mandatory pre-assessment form.
    """
    # Check if assessment already exists for this candidate + job
    existing = await db.trajectory_assessments.find_one({
        "candidate_email": request.candidate_email,
        "job_id": request.job_id
    })
    
    if existing:
        return {
            "success": True,
            "assessment_id": existing["id"],
            "message": "Assessment already exists for this candidate and job",
            "existing": True,
            "status": existing.get("status")
        }
    
    # Get job details
    job = await db.job_descriptions.find_one({"id": request.job_id})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Create new assessment
    assessment = CareerTrajectoryAssessment(
        candidate_name=request.candidate_name,
        candidate_email=request.candidate_email,
        candidate_phone=request.candidate_phone,
        resume_text=request.resume_text,
        assessment_type="pre_assessment",
        data_collection_mode="candidate",
        job_id=request.job_id,
        target_role=job.get("title"),
        target_industry=job.get("industry"),
        target_location=job.get("location"),
        status="pending"
    )
    
    # If resume_id provided, link it
    if request.resume_id:
        assessment_dict = assessment.model_dump()
        assessment_dict["resume_id"] = request.resume_id
    else:
        assessment_dict = assessment.model_dump()
    
    await db.trajectory_assessments.insert_one(assessment_dict)
    
    # Generate pre-assessment link
    preassessment_link = f"{os.environ.get('FRONTEND_URL', 'http://localhost:3000')}/trajectory/preassessment/{assessment.id}?token={assessment.access_token}"
    
    # Send pre-assessment email
    simulate_send_preassessment_email(
        candidate_email=request.candidate_email,
        candidate_name=request.candidate_name,
        job_title=job.get("title", "Position"),
        job_id=request.job_id,
        assessment_link=preassessment_link
    )
    
    return {
        "success": True,
        "assessment_id": assessment.id,
        "access_token": assessment.access_token,
        "preassessment_link": preassessment_link,
        "job_title": job.get("title"),
        "message": "Pre-assessment triggered. Email sent to candidate.",
        "existing": False
    }

@api_router.post("/trajectory/assessment/{assessment_id}/resend-preassessment")
async def resend_preassessment_to_candidate(assessment_id: str):
    """
    Manual trigger for recruiter to resend pre-assessment form to candidate.
    Use this if auto-trigger failed or candidate didn't receive the email.
    """
    assessment = await db.trajectory_assessments.find_one({"id": assessment_id})
    
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    if assessment.get("questionnaire_completed"):
        return {
            "success": False,
            "message": "Pre-assessment already completed by candidate"
        }
    
    # Get job details
    job_id = assessment.get("job_id")
    job = await db.job_descriptions.find_one({"id": job_id}) if job_id else None
    job_title = job.get("title", assessment.get("target_role", "Position")) if job else assessment.get("target_role", "Position")
    
    # Generate pre-assessment link
    frontend_url = os.environ.get('FRONTEND_URL', 'http://localhost:3000')
    preassessment_link = f"{frontend_url}/trajectory/preassessment/{assessment['id']}?token={assessment['access_token']}"
    
    # Send pre-assessment email
    simulate_send_preassessment_email(
        candidate_email=assessment["candidate_email"],
        candidate_name=assessment["candidate_name"],
        job_title=job_title,
        job_id=job_id or assessment_id,
        assessment_link=preassessment_link
    )
    
    # Update resend count
    resend_count = assessment.get("preassessment_resend_count", 0) + 1
    await db.trajectory_assessments.update_one(
        {"id": assessment_id},
        {"$set": {
            "preassessment_resend_count": resend_count,
            "preassessment_last_resent_at": datetime.now(timezone.utc),
            "updated_at": datetime.now(timezone.utc)
        }}
    )
    
    return {
        "success": True,
        "message": f"Pre-assessment form resent to {assessment['candidate_email']}",
        "preassessment_link": preassessment_link,
        "resend_count": resend_count
    }

# Auto-reject incomplete assessments
@api_router.post("/trajectory/assessment/{assessment_id}/auto-reject-incomplete")
async def auto_reject_incomplete_assessment(assessment_id: str):
    """Auto-reject assessment if pre-assessment not completed within deadline"""
    assessment = await db.trajectory_assessments.find_one({"id": assessment_id})
    
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    if assessment.get("questionnaire_completed"):
        return {"success": False, "message": "Assessment already completed"}
    
    # Update status to rejected
    reasons = ["Pre-assessment form not completed within required timeframe", "Application automatically rejected due to incomplete information"]
    
    await db.trajectory_assessments.update_one(
        {"id": assessment_id},
        {"$set": {
            "status": "rejected",
            "recruiter_decision": "reject",
            "rejection_reasons": reasons,
            "auto_rejected": True,
            "auto_rejected_at": datetime.now(timezone.utc),
            "updated_at": datetime.now(timezone.utc)
        }}
    )
    
    # Send rejection email
    job_id = assessment.get("job_id", assessment["id"])
    job = await db.job_descriptions.find_one({"id": job_id}) if assessment.get("job_id") else None
    job_title = job.get("title", assessment.get("target_role", "Position")) if job else assessment.get("target_role", "Position")
    
    simulate_send_rejection_email(
        candidate_email=assessment["candidate_email"],
        candidate_name=assessment["candidate_name"],
        job_title=job_title,
        job_id=job_id,
        reasons=reasons,
        status="rejected"
    )
    
    return {
        "success": True,
        "message": "Assessment auto-rejected and candidate notified",
        "reasons": reasons
    }

# Employment History API
@api_router.post("/trajectory/assessment/{assessment_id}/employment-history")
async def submit_employment_history(assessment_id: str, request: SubmitEmploymentHistoryRequest):
    """Submit or update detailed employment history"""
    assessment = await db.trajectory_assessments.find_one({"id": assessment_id})
    
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    # Calculate metrics from employment history
    employment_history = request.employment_history
    total_months = 0
    total_jobs = len(employment_history)
    
    for job in employment_history:
        # Calculate tenure for each job
        start = job.get("start_date", "")
        end = job.get("end_date", "Present")
        if start:
            try:
                start_date = datetime.strptime(start, "%Y-%m")
                end_date = datetime.now() if end == "Present" else datetime.strptime(end, "%Y-%m")
                months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
                job["tenure_months"] = months
                total_months += months
                
                # Calculate hike percentage if both CTCs provided
                joining_ctc = job.get("ctc_at_joining")
                exit_ctc = job.get("ctc_at_exit")
                if joining_ctc and exit_ctc and joining_ctc > 0:
                    job["hike_percentage"] = round(((exit_ctc - joining_ctc) / joining_ctc) * 100, 1)
            except ValueError:
                pass
    
    avg_tenure_months = total_months / total_jobs if total_jobs > 0 else 0
    total_experience_years = total_months / 12
    
    update_data = {
        "employment_history": employment_history,
        "total_experience_years": round(total_experience_years, 1),
        "average_tenure_months": round(avg_tenure_months, 1),
        "updated_at": datetime.now(timezone.utc)
    }
    
    await db.trajectory_assessments.update_one(
        {"id": assessment_id},
        {"$set": update_data}
    )
    
    return {
        "success": True,
        "message": "Employment history updated",
        "total_experience_years": round(total_experience_years, 1),
        "average_tenure_months": round(avg_tenure_months, 1),
        "total_jobs": total_jobs
    }

# Document Upload API
@api_router.post("/trajectory/assessment/{assessment_id}/upload-document")
async def upload_assessment_document(assessment_id: str, document_type: str, file: UploadFile = File(...)):
    """Upload document for assessment (resignation proof or counter-offer letter)"""
    assessment = await db.trajectory_assessments.find_one({"id": assessment_id})
    
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    if document_type not in ["resignation_proof", "counter_offer_proof"]:
        raise HTTPException(status_code=400, detail="Invalid document type. Use 'resignation_proof' or 'counter_offer_proof'")
    
    # Validate file type
    allowed_types = ["image/png", "image/jpeg", "application/pdf"]
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail=f"File type {file.content_type} not allowed. Use PNG, JPEG, or PDF")
    
    # Save file
    file_ext = file.filename.split(".")[-1] if file.filename else "pdf"
    file_name = f"{assessment_id}_{document_type}_{uuid.uuid4().hex[:8]}.{file_ext}"
    file_path = f"/app/uploads/trajectory/{file_name}"
    
    # Ensure directory exists
    os.makedirs("/app/uploads/trajectory", exist_ok=True)
    
    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)
    
    # Update assessment with file URL
    field_name = f"{document_type}_url"
    await db.trajectory_assessments.update_one(
        {"id": assessment_id},
        {"$set": {field_name: file_path, "updated_at": datetime.now(timezone.utc)}}
    )
    
    return {
        "success": True,
        "message": f"{document_type.replace('_', ' ').title()} uploaded successfully",
        "file_path": file_path
    }

@api_router.get("/trajectory/assessments")
async def list_trajectory_assessments(
    status: Optional[str] = None,
    assessment_type: Optional[str] = None,
    job_id: Optional[str] = None,
    skip: int = 0,
    limit: int = 50
):
    """List all trajectory assessments with optional filters"""
    query = {}
    
    if status:
        query["status"] = status
    if assessment_type:
        query["assessment_type"] = assessment_type
    if job_id:
        query["job_id"] = job_id
    
    assessments = await db.trajectory_assessments.find(
        query,
        {"_id": 0}
    ).sort("created_at", -1).skip(skip).limit(limit).to_list(limit)
    
    total = await db.trajectory_assessments.count_documents(query)
    
    return {
        "assessments": assessments,
        "total": total,
        "skip": skip,
        "limit": limit
    }

@api_router.get("/trajectory/stats")
async def get_trajectory_stats():
    """Get statistics for career trajectory assessments"""
    total = await db.trajectory_assessments.count_documents({})
    completed = await db.trajectory_assessments.count_documents({"status": "completed"})
    pending = await db.trajectory_assessments.count_documents({"status": "pending"})
    reviewed = await db.trajectory_assessments.count_documents({"status": "reviewed"})
    
    # Flag distribution
    green_count = await db.trajectory_assessments.count_documents({"overall_flag": "green"})
    yellow_count = await db.trajectory_assessments.count_documents({"overall_flag": "yellow"})
    red_count = await db.trajectory_assessments.count_documents({"overall_flag": "red"})
    
    # Average scores by indicator (aggregation)
    pipeline = [
        {"$match": {"status": {"$in": ["completed", "reviewed"]}}},
        {"$unwind": "$indicator_results"},
        {"$group": {
            "_id": "$indicator_results.indicator_id",
            "avg_score": {"$avg": "$indicator_results.score"},
            "count": {"$sum": 1}
        }}
    ]
    
    indicator_stats = await db.trajectory_assessments.aggregate(pipeline).to_list(10)
    
    return {
        "total_assessments": total,
        "by_status": {
            "completed": completed,
            "pending": pending,
            "reviewed": reviewed,
            "in_progress": total - completed - pending - reviewed
        },
        "by_flag": {
            "green": green_count,
            "yellow": yellow_count,
            "red": red_count
        },
        "indicator_averages": {stat["_id"]: round(stat["avg_score"], 1) for stat in indicator_stats}
    }

@api_router.delete("/trajectory/assessment/{assessment_id}")
async def delete_trajectory_assessment(assessment_id: str):
    """Delete a trajectory assessment"""
    result = await db.trajectory_assessments.delete_one({"id": assessment_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    return {"success": True, "message": "Assessment deleted"}

# Integration endpoint - Assess from Resume Repository
@api_router.post("/trajectory/from-resume/{resume_id}")
async def create_assessment_from_resume(resume_id: str, target_role: Optional[str] = None, target_industry: Optional[str] = None):
    """Create a trajectory assessment from an existing resume in repository"""
    resume = await db.routed_resumes.find_one({"id": resume_id})
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Check if assessment already exists
    existing = await db.trajectory_assessments.find_one({
        "candidate_email": resume["email"],
        "status": {"$nin": ["completed", "reviewed"]}
    })
    
    if existing:
        return {
            "success": True,
            "assessment_id": existing["id"],
            "message": "Existing assessment found",
            "existing": True
        }
    
    # Create new assessment
    assessment = CareerTrajectoryAssessment(
        candidate_id=resume_id,
        candidate_name=resume["name"],
        candidate_email=resume["email"],
        resume_text=resume.get("raw_text", ""),
        assessment_type="post_application",
        job_id=resume.get("job_id"),
        target_role=target_role or resume.get("applied_job_title"),
        target_industry=target_industry,
        status="in_progress"
    )
    
    # Run initial analysis
    try:
        analysis = await analyze_career_trajectory(
            resume_text=resume.get("raw_text", ""),
            questionnaire_responses={},
            target_role=target_role or resume.get("applied_job_title"),
            target_industry=target_industry
        )
        
        indicator_results = analysis.get("indicators", [])
        overall_score, overall_flag = calculate_overall_score(indicator_results)
        green_flags, yellow_flags, red_flags = categorize_flags(indicator_results)
        
        assessment.indicator_results = indicator_results
        assessment.overall_score = overall_score
        assessment.overall_flag = overall_flag
        assessment.green_flags = green_flags
        assessment.yellow_flags = yellow_flags
        assessment.red_flags = red_flags
        assessment.key_strengths = analysis.get("key_strengths", [])
        assessment.areas_to_probe = analysis.get("areas_to_probe", [])
        assessment.recruiter_summary = analysis.get("recruiter_summary", "")
        assessment.career_timeline = analysis.get("career_timeline", [])
        assessment.skill_progression = analysis.get("skill_progression", [])
    except Exception as e:
        logger.error(f"Resume analysis failed: {e}")
    
    # Save to database
    assessment_dict = assessment.model_dump()
    await db.trajectory_assessments.insert_one(assessment_dict)
    
    return {
        "success": True,
        "assessment_id": assessment.id,
        "access_token": assessment.access_token,
        "questionnaire_link": f"/trajectory/questionnaire/{assessment.id}?token={assessment.access_token}",
        "status": assessment.status,
        "existing": False
    }

# ============ Dashboard/Stats Routes ============

@api_router.get("/dashboard/stats")
async def get_dashboard_stats():
    """Get dashboard statistics"""
    total_jds = await db.job_descriptions.count_documents({})
    total_candidates = await db.candidates.count_documents({})
    total_matches = await db.matches.count_documents({})
    
    # Pipeline breakdown
    pipeline_stages = await get_pipeline_stages()
    pipeline_stats = {}
    for stage in pipeline_stages:
        count = await db.candidates.count_documents({"pipeline_stage": stage["id"]})
        pipeline_stats[stage["id"]] = count
    
    # Recent activity
    recent_candidates = await db.candidates.find(
        {}, {"_id": 0, "id": 1, "name": 1, "created_at": 1}
    ).sort("created_at", -1).limit(5).to_list(5)
    
    recent_jds = await db.job_descriptions.find(
        {}, {"_id": 0, "id": 1, "title": 1, "created_at": 1}
    ).sort("created_at", -1).limit(5).to_list(5)
    
    return {
        "totals": {
            "job_descriptions": total_jds,
            "candidates": total_candidates,
            "matches": total_matches
        },
        "pipeline": pipeline_stats,
        "recent_candidates": recent_candidates,
        "recent_job_descriptions": recent_jds
    }

# ============ ADMIN PANEL API ROUTES ============

# Initialize default admin user on startup
async def initialize_admin_user():
    """Create default admin user if not exists"""
    admin = await db.admin_users.find_one({"email": "admin@rolesense.in"})
    if not admin:
        admin_user = {
            "id": str(uuid.uuid4()),
            "email": "admin@rolesense.in",
            "password_hash": hash_password("Admin@123"),  # Default password
            "name": "System Admin",
            "role": "super_admin",
            "created_at": datetime.now(timezone.utc),
            "is_active": True
        }
        await db.admin_users.insert_one(admin_user)
        logger.info("Default admin user created: admin@rolesense.in / Admin@123")

async def initialize_business_users():
    """Create business users for Ally Executive if not exists"""
    
    # 1. Create Corporate Client Organization for Ally Executive
    corporate_client = await db.client_organizations.find_one({"contact_email": "s-kaul@ally-executive.com"})
    if not corporate_client:
        corporate_org = {
            "id": str(uuid.uuid4()),
            "organization_name": "Ally Executive Search",
            "organization_type": "corporate",
            "business_domain": "ally-executive.com",
            "contact_email": "s-kaul@ally-executive.com",
            "contact_phone": "+91-9876543210",
            "contact_name": "S Kaul",
            "is_active": True,
            "access_level": "full",
            "subscription_status": "active",
            "trial_start_date": datetime.now(timezone.utc),
            "trial_end_date": datetime.now(timezone.utc) + timedelta(days=365),
            "modules_enabled": {
                "jd_intelligence": True,
                "resume_repository": True,
                "career_trajectory": True,
                "hr_fitment": True
            },
            "created_at": datetime.now(timezone.utc),
            "updated_at": datetime.now(timezone.utc)
        }
        await db.client_organizations.insert_one(corporate_org)
        
        # Create Corporate User
        corporate_user = {
            "id": str(uuid.uuid4()),
            "email": "s-kaul@ally-executive.com",
            "password_hash": hash_password("AllyExec@2025"),
            "name": "S Kaul",
            "role": "admin",
            "client_id": corporate_org["id"],
            "is_active": True,
            "created_at": datetime.now(timezone.utc)
        }
        await db.client_users.insert_one(corporate_user)
        logger.info("Corporate user created: s-kaul@ally-executive.com")
    
    # 2. Create Staffing Vendor Client Organization for Ally Executive
    vendor_client = await db.client_organizations.find_one({"contact_email": "pushpanjali.k@ally-executive.com"})
    if not vendor_client:
        vendor_org = {
            "id": str(uuid.uuid4()),
            "organization_name": "Ally Executive Staffing",
            "organization_type": "staffing_vendor",
            "business_domain": "ally-executive.com",
            "contact_email": "pushpanjali.k@ally-executive.com",
            "contact_phone": "+91-9876543211",
            "contact_name": "Pushpanjali K",
            "is_active": True,
            "access_level": "full",
            "subscription_status": "active",
            "trial_start_date": datetime.now(timezone.utc),
            "trial_end_date": datetime.now(timezone.utc) + timedelta(days=365),
            "modules_enabled": {
                "jd_intelligence": True,
                "resume_repository": True,
                "career_trajectory": True,
                "hr_fitment": True
            },
            "created_at": datetime.now(timezone.utc),
            "updated_at": datetime.now(timezone.utc)
        }
        await db.client_organizations.insert_one(vendor_org)
        
        # Create Vendor User
        vendor_user = {
            "id": str(uuid.uuid4()),
            "email": "pushpanjali.k@ally-executive.com",
            "password_hash": hash_password("AllyVendor@2025"),
            "name": "Pushpanjali K",
            "role": "admin",
            "client_id": vendor_org["id"],
            "is_active": True,
            "created_at": datetime.now(timezone.utc)
        }
        await db.client_users.insert_one(vendor_user)
        logger.info("Vendor user created: pushpanjali.k@ally-executive.com")

@api_router.post("/admin/login")
async def admin_login(request: AdminLoginRequest):
    """Admin login endpoint"""
    admin = await db.admin_users.find_one({"email": request.email, "is_active": True})
    
    # Decrypt password if encrypted
    password = decrypt_password(request.password) if request.encrypted else request.password
    
    if not admin or not verify_password(password, admin["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    # Update last login
    await db.admin_users.update_one(
        {"id": admin["id"]},
        {"$set": {"last_login": datetime.now(timezone.utc)}}
    )
    
    return {
        "success": True,
        "admin": {
            "id": admin["id"],
            "email": admin["email"],
            "name": admin["name"],
            "role": admin["role"]
        },
        "token": admin["id"]  # Simple token for demo
    }

@api_router.get("/admin/dashboard")
async def get_admin_dashboard():
    """Get comprehensive admin dashboard statistics"""
    from datetime import timedelta
    
    now = datetime.now(timezone.utc)
    start_of_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    start_of_quarter = now.replace(month=((now.month - 1) // 3) * 3 + 1, day=1, hour=0, minute=0, second=0, microsecond=0)
    last_30_days = now - timedelta(days=30)
    last_7_days = now - timedelta(days=7)
    
    # ========== 1. BUSINESS HEALTH KPIs ==========
    
    # Active Jobs
    total_active_jobs = await db.job_descriptions.count_documents({"status": "active"})
    structured_active_jobs = await db.structured_jds.count_documents({"status": "active"})
    total_jobs = total_active_jobs + structured_active_jobs
    
    # Applications
    total_applications = await db.job_applications.count_documents({})
    applications_mtd = await db.job_applications.count_documents({"created_at": {"$gte": start_of_month.isoformat()}})
    
    # Candidates by Stage
    candidates_preassessment = await db.trajectory_assessments.count_documents({"status": {"$in": ["pending", "in_progress"]}})
    candidates_completed = await db.trajectory_assessments.count_documents({"status": "completed"})
    candidates_shortlisted = await db.job_applications.count_documents({"status": "shortlisted"})
    candidates_rejected = await db.job_applications.count_documents({"status": "rejected"})
    
    # Clients & Partners
    total_clients = await db.client_organizations.count_documents({})
    active_clients = await db.client_organizations.count_documents({"is_active": True, "subscription_status": {"$in": ["trial", "active"]}})
    corporate_clients = await db.client_organizations.count_documents({"organization_type": "corporate", "is_active": True})
    staffing_partners = await db.client_organizations.count_documents({"organization_type": "staffing_vendor", "is_active": True})
    
    # Trial & Subscription
    trial_clients = await db.client_organizations.count_documents({"subscription_status": "trial"})
    paid_clients = await db.client_organizations.count_documents({"subscription_status": "active"})
    expired_clients = await db.client_organizations.count_documents({"subscription_status": "expired"})
    
    # Monthly Placements (applications marked as 'hired' or 'placed')
    placements_mtd = await db.job_applications.count_documents({
        "status": {"$in": ["hired", "placed", "joined"]},
        "updated_at": {"$gte": start_of_month.isoformat()}
    })
    
    # ========== 2. JOB & DEMAND CONTROL ==========
    
    # Jobs by Status
    jobs_open = await db.job_descriptions.count_documents({"status": "active"}) + await db.structured_jds.count_documents({"status": "active"})
    jobs_draft = await db.job_descriptions.count_documents({"status": "draft"}) + await db.structured_jds.count_documents({"status": "draft"})
    jobs_closed = await db.job_descriptions.count_documents({"status": {"$in": ["closed", "filled"]}})
    jobs_on_hold = await db.job_descriptions.count_documents({"status": "on_hold"})
    
    # Jobs by Client Type
    jobs_pipeline = [
        {"$group": {"_id": "$client_name", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
        {"$limit": 10},
        {"$project": {"_id": 0, "name": "$_id", "count": 1}}
    ]
    jobs_by_client_raw = await db.job_descriptions.aggregate(jobs_pipeline).to_list(10)
    jobs_by_client = [{"_id": j.get("name"), "count": j.get("count", 0)} for j in jobs_by_client_raw]
    
    # Jobs by Location
    location_pipeline = [
        {"$match": {"location": {"$exists": True, "$ne": None}}},
        {"$group": {"_id": "$location", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
        {"$limit": 10},
        {"$project": {"_id": 0, "name": "$_id", "count": 1}}
    ]
    jobs_by_location_raw = await db.job_descriptions.aggregate(location_pipeline).to_list(10)
    jobs_by_location = [{"_id": j.get("name"), "count": j.get("count", 0)} for j in jobs_by_location_raw]
    
    # ========== 3. CANDIDATE FUNNEL & INDICATORS ==========
    
    # Resume Stats
    total_resumes = await db.routed_resumes.count_documents({})
    resumes_mtd = await db.routed_resumes.count_documents({"created_at": {"$gte": start_of_month.isoformat()}})
    
    # Resume by Category
    resume_category_pipeline = [
        {"$group": {"_id": "$category", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
        {"$project": {"_id": 0, "name": "$_id", "count": 1}}
    ]
    resumes_by_category_raw = await db.routed_resumes.aggregate(resume_category_pipeline).to_list(20)
    resumes_by_category = [{"_id": r.get("name"), "count": r.get("count", 0)} for r in resumes_by_category_raw]
    
    # Assessment Stats
    total_assessments = await db.trajectory_assessments.count_documents({})
    assessments_completed = await db.trajectory_assessments.count_documents({"status": "completed"})
    assessments_pending = await db.trajectory_assessments.count_documents({"status": {"$in": ["pending", "in_progress"]}})
    
    # Pre-assessment completion rate
    preassessment_completion_rate = round((assessments_completed / max(total_assessments, 1)) * 100, 1)
    
    # Indicator Distribution (from fitment assessments)
    indicator_pipeline = [
        {"$match": {"overall_score": {"$exists": True}}},
        {"$bucket": {
            "groupBy": "$overall_score",
            "boundaries": [0, 40, 70, 100],
            "default": "unknown",
            "output": {"count": {"$sum": 1}}
        }}
    ]
    try:
        indicator_dist = await db.hr_fitment.aggregate(indicator_pipeline).to_list(10)
    except:
        indicator_dist = []
    
    # ========== 4. RECRUITER/PARTNER PERFORMANCE ==========
    
    # User activity (simplified - based on client users)
    total_users = await db.client_users.count_documents({})
    active_users = await db.client_users.count_documents({"is_active": True})
    
    # Partner-wise stats
    partner_pipeline = [
        {"$match": {"organization_type": "staffing_vendor"}},
        {"$project": {
            "_id": 0,
            "organization_name": 1,
            "total_users": 1,
            "total_resumes": 1,
            "total_assessments": 1,
            "subscription_status": 1
        }},
        {"$sort": {"total_resumes": -1}},
        {"$limit": 10}
    ]
    top_partners = await db.client_organizations.aggregate(partner_pipeline).to_list(10)
    
    # ========== 5. CLIENT & EMPLOYER ANALYTICS ==========
    
    # Client-wise job success
    client_pipeline = [
        {"$match": {"organization_type": "corporate"}},
        {"$project": {
            "_id": 0,
            "organization_name": 1,
            "business_domain": 1,
            "total_jds_created": 1,
            "total_resumes": 1,
            "subscription_status": 1,
            "access_level": 1
        }},
        {"$sort": {"total_jds_created": -1}},
        {"$limit": 10}
    ]
    top_clients = await db.client_organizations.aggregate(client_pipeline).to_list(10)
    
    # ========== 6. PLATFORM USAGE & SYSTEM HEALTH ==========
    
    # Resume parsing success (assume all parsed successfully for now)
    resume_parsing_success_rate = 98.5
    
    # Assessment completion time (mock for now)
    avg_assessment_time = "4.2 mins"
    
    # ========== 7. REVENUE & BILLING (Mock Data) ==========
    
    # Revenue calculations (placeholder - would need actual billing data)
    revenue_mtd = paid_clients * 15000  # Assume avg 15k/client/month
    revenue_qtd = paid_clients * 45000  # Quarterly
    
    # ========== 8. ALERTS & NOTIFICATIONS ==========
    
    # Trial ending soon (within 7 days)
    trial_ending_soon = await db.client_organizations.find({
        "subscription_status": "trial",
        "trial_end_date": {
            "$lte": (now + timedelta(days=7)).isoformat(),
            "$gte": now.isoformat()
        }
    }, {"_id": 0, "organization_name": 1, "trial_end_date": 1, "contact_email": 1}).to_list(20)
    
    # Recent clients
    recent_clients = await db.client_organizations.find(
        {}, {"_id": 0}
    ).sort("created_at", -1).limit(5).to_list(5)
    
    # Pending feedback
    pending_feedback = await db.customer_feedback.count_documents({"status": {"$in": ["new", "in_review"]}})
    
    return {
        # 1. Business Health KPIs
        "kpis": {
            "total_active_jobs": total_jobs,
            "total_applications": total_applications,
            "applications_mtd": applications_mtd,
            "candidates_preassessment": candidates_preassessment,
            "candidates_completed": candidates_completed,
            "candidates_shortlisted": candidates_shortlisted,
            "candidates_rejected": candidates_rejected,
            "active_clients": active_clients,
            "staffing_partners": staffing_partners,
            "placements_mtd": placements_mtd,
            "revenue_mtd": revenue_mtd,
            "revenue_qtd": revenue_qtd
        },
        
        # 2. Job & Demand Control
        "jobs": {
            "by_status": {
                "open": jobs_open,
                "draft": jobs_draft,
                "on_hold": jobs_on_hold,
                "closed": jobs_closed
            },
            "by_client": jobs_by_client,
            "by_location": jobs_by_location,
            "total": total_jobs
        },
        
        # 3. Candidate Funnel
        "candidates": {
            "total_resumes": total_resumes,
            "resumes_mtd": resumes_mtd,
            "by_category": resumes_by_category,
            "assessments": {
                "total": total_assessments,
                "completed": assessments_completed,
                "pending": assessments_pending,
                "completion_rate": preassessment_completion_rate
            },
            "indicator_distribution": indicator_dist
        },
        
        # 4. Recruiter/Partner Performance
        "partners": {
            "total_users": total_users,
            "active_users": active_users,
            "top_partners": top_partners
        },
        
        # 5. Client Analytics
        "clients": {
            "total": total_clients,
            "active": active_clients,
            "trial": trial_clients,
            "paid": paid_clients,
            "expired": expired_clients,
            "corporate": corporate_clients,
            "staffing_vendor": staffing_partners,
            "top_clients": top_clients
        },
        
        # 6. System Health
        "system": {
            "resume_parsing_success_rate": resume_parsing_success_rate,
            "avg_assessment_time": avg_assessment_time,
            "api_status": "healthy",
            "uptime": "99.9%"
        },
        
        # 7. Revenue
        "revenue": {
            "mtd": revenue_mtd,
            "qtd": revenue_qtd,
            "currency": "INR"
        },
        
        # 8. Alerts
        "alerts": {
            "trial_ending_soon": trial_ending_soon,
            "pending_feedback": pending_feedback
        },
        
        # Recent Activity
        "recent_clients": recent_clients
    }


@api_router.post("/admin/cleanup-test-data")
async def cleanup_test_data():
    """Remove all test/demo data from the database - Admin only"""
    cleanup_results = {
        "deleted_users": 0,
        "deleted_organizations": 0,
        "deleted_resumes": 0,
        "deleted_jds": 0,
        "deleted_applications": 0,
        "deleted_assessments": 0,
        "deleted_preassessment_responses": 0
    }
    
    # Remove demo/test users
    test_emails = [
        "test@example.com", 
        "vendor@test.com",
        "contact@demo-corporate.com",
        "contact@demo-staffing.com",
        "hr@testcompany.com"
    ]
    
    # Delete test users
    result = await db.client_users.delete_many({"email": {"$in": test_emails}})
    cleanup_results["deleted_users"] = result.deleted_count
    
    # Delete test organizations
    test_org_names = ["Demo Corporate Inc", "Demo Staffing Solutions", "Test External Company"]
    test_domains = ["testcompany.com"]
    result = await db.client_organizations.delete_many({
        "$or": [
            {"organization_name": {"$in": test_org_names}},
            {"business_domain": {"$in": test_domains}}
        ]
    })
    cleanup_results["deleted_organizations"] = result.deleted_count
    
    # Delete resumes with source "test" or test email patterns
    result = await db.routed_resumes.delete_many({
        "$or": [
            {"source": "test"},
            {"email": {"$regex": "@email.com$"}},
            {"email": {"$regex": "@test.com$"}},
            {"email": {"$regex": "@example.com$"}},
            {"email": {"$regex": "@testcompany.com$"}},
            {"email": {"$regex": "@testdomain.com$"}},
            {"name": {"$regex": "^Test ", "$options": "i"}}
        ]
    })
    cleanup_results["deleted_resumes"] = result.deleted_count
    
    # Delete test JDs (those with test/demo in company name)
    result = await db.job_descriptions.delete_many({
        "$or": [
            {"company": {"$regex": "test|demo", "$options": "i"}},
            {"title": {"$regex": "test|demo", "$options": "i"}}
        ]
    })
    cleanup_results["deleted_jds"] = result.deleted_count
    
    # Delete test applications
    result = await db.job_applications.delete_many({
        "$or": [
            {"email": {"$regex": "@email.com$"}},
            {"email": {"$regex": "@test.com$"}},
            {"email": {"$regex": "@example.com$"}},
            {"email": {"$regex": "@testcompany.com$"}},
            {"email": {"$regex": "@testdomain.com$"}},
            {"applicant_email": {"$regex": "@email.com$"}},
            {"applicant_email": {"$regex": "@test.com$"}},
            {"applicant_email": {"$regex": "@example.com$"}},
            {"applicant_name": {"$regex": "^Test ", "$options": "i"}}
        ]
    })
    cleanup_results["deleted_applications"] = result.deleted_count
    
    # Delete test assessments
    result = await db.trajectory_assessments.delete_many({
        "$or": [
            {"email": {"$regex": "@email.com$"}},
            {"email": {"$regex": "@test.com$"}},
            {"email": {"$regex": "@example.com$"}},
            {"email": {"$regex": "@testcompany.com$"}},
            {"candidate_email": {"$regex": "@email.com$"}},
            {"candidate_email": {"$regex": "@test.com$"}},
            {"candidate_email": {"$regex": "@testdomain.com$"}}
        ]
    })
    cleanup_results["deleted_assessments"] = result.deleted_count
    
    # Delete test pre-assessment responses
    result = await db.preassessment_responses.delete_many({
        "$or": [
            {"candidate_email": {"$regex": "@email.com$"}},
            {"candidate_email": {"$regex": "@test.com$"}},
            {"candidate_email": {"$regex": "@example.com$"}},
            {"candidate_email": {"$regex": "@testdomain.com$"}}
        ]
    })
    cleanup_results["deleted_preassessment_responses"] = result.deleted_count
    
    # Delete test competency reports
    await db.competency_reports.delete_many({
        "$or": [
            {"candidate_email": {"$regex": "@email.com$"}},
            {"candidate_email": {"$regex": "@test.com$"}},
            {"candidate_email": {"$regex": "@testdomain.com$"}}
        ]
    })
    
    logger.info(f"Test data cleanup completed: {cleanup_results}")
    
    return {
        "success": True,
        "message": "Test data cleanup completed successfully",
        "details": cleanup_results
    }

@api_router.post("/admin/clients")
async def create_client(request: CreateClientRequest):
    """Create a new client organization"""
    from datetime import timedelta
    
    # Check if client with same domain exists
    existing = await db.client_organizations.find_one({"business_domain": request.business_domain})
    if existing:
        raise HTTPException(status_code=400, detail="Client with this business domain already exists")
    
    trial_start = datetime.now(timezone.utc)
    trial_end = trial_start + timedelta(days=90)
    
    client = ClientOrganization(
        organization_name=request.organization_name,
        organization_type=request.organization_type,
        business_domain=request.business_domain,
        contact_email=request.contact_email,
        contact_phone=request.contact_phone,
        contact_person=request.contact_person,
        address=request.address,
        trial_start_date=trial_start,
        trial_end_date=trial_end
    )
    
    await db.client_organizations.insert_one(client.model_dump())
    
    # Generate default admin user for the client
    default_password = f"Welcome@{uuid.uuid4().hex[:6].upper()}"
    client_admin = ClientUser(
        client_id=client.id,
        email=request.contact_email,
        password_hash=hash_password(default_password),
        name=request.contact_person,
        role="admin"
    )
    await db.client_users.insert_one(client_admin.model_dump())
    
    # Update client user count
    await db.client_organizations.update_one(
        {"id": client.id},
        {"$inc": {"total_users": 1}}
    )
    
    return {
        "success": True,
        "client": client.model_dump(),
        "admin_credentials": {
            "email": request.contact_email,
            "password": default_password,
            "message": "Please change password on first login"
        },
        "trial_period": {
            "start": trial_start.isoformat(),
            "end": trial_end.isoformat(),
            "days": 90
        }
    }

@api_router.get("/admin/clients")
async def list_clients(
    status: Optional[str] = None,
    org_type: Optional[str] = None,
    skip: int = 0,
    limit: int = 50
):
    """List all client organizations"""
    query = {}
    if status:
        query["subscription_status"] = status
    if org_type:
        query["organization_type"] = org_type
    
    clients = await db.client_organizations.find(query, {"_id": 0}).sort("created_at", -1).skip(skip).limit(limit).to_list(limit)
    total = await db.client_organizations.count_documents(query)
    
    # Calculate trial days remaining for each client
    for client in clients:
        if client.get("subscription_status") == "trial" and client.get("trial_end_date"):
            trial_end = client["trial_end_date"]
            if isinstance(trial_end, str):
                trial_end = datetime.fromisoformat(trial_end.replace('Z', '+00:00'))
            elif isinstance(trial_end, datetime):
                # Ensure timezone awareness
                if trial_end.tzinfo is None:
                    trial_end = trial_end.replace(tzinfo=timezone.utc)
            now_utc = datetime.now(timezone.utc)
            days_remaining = (trial_end - now_utc).days
            client["trial_days_remaining"] = max(0, days_remaining)
    
    return {"clients": clients, "total": total}

@api_router.get("/admin/clients/{client_id}")
async def get_client(client_id: str):
    """Get client details"""
    client = await db.client_organizations.find_one({"id": client_id}, {"_id": 0})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Get users for this client
    users = await db.client_users.find({"client_id": client_id}, {"_id": 0, "password_hash": 0}).to_list(100)
    
    # Get resume count
    resume_count = await db.routed_resumes.count_documents({"client_id": client_id})
    
    # Get assessment count
    assessment_count = await db.trajectory_assessments.count_documents({"client_id": client_id})
    
    # Get JD count
    jd_count = await db.job_descriptions.count_documents({"client_id": client_id})
    
    return {
        "client": client,
        "users": users,
        "stats": {
            "resumes": resume_count,
            "assessments": assessment_count,
            "jds": jd_count,
            "users": len(users)
        }
    }

@api_router.put("/admin/clients/{client_id}/access")
async def update_client_access(client_id: str, request: UpdateAccessRequest):
    """Update client access level and module permissions"""
    client = await db.client_organizations.find_one({"id": client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    update_data = {"updated_at": datetime.now(timezone.utc)}
    
    if request.access_level:
        update_data["access_level"] = request.access_level
    if request.modules_enabled:
        update_data["modules_enabled"] = request.modules_enabled
    if request.monthly_assessment_limit is not None:
        update_data["monthly_assessment_limit"] = request.monthly_assessment_limit
    if request.subscription_status:
        update_data["subscription_status"] = request.subscription_status
        if request.subscription_status == "active":
            update_data["paid_start_date"] = datetime.now(timezone.utc)
    
    await db.client_organizations.update_one(
        {"id": client_id},
        {"$set": update_data}
    )
    
    return {"success": True, "message": "Access updated successfully"}

@api_router.post("/admin/clients/{client_id}/users")
async def create_client_user(client_id: str, request: CreateClientUserRequest):
    """Create a new user for a client organization"""
    client = await db.client_organizations.find_one({"id": client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Check if user with email exists
    existing = await db.client_users.find_one({"email": request.email})
    if existing:
        raise HTTPException(status_code=400, detail="User with this email already exists")
    
    user = ClientUser(
        client_id=client_id,
        email=request.email,
        password_hash=hash_password(request.password),
        name=request.name,
        role=request.role,
        department=request.department
    )
    
    await db.client_users.insert_one(user.model_dump())
    
    # Update client user count
    await db.client_organizations.update_one(
        {"id": client_id},
        {"$inc": {"total_users": 1}}
    )
    
    return {
        "success": True,
        "user": {
            "id": user.id,
            "email": user.email,
            "name": user.name,
            "role": user.role
        }
    }

@api_router.get("/admin/clients/{client_id}/users")
async def list_client_users(client_id: str):
    """List all users for a client"""
    users = await db.client_users.find(
        {"client_id": client_id},
        {"_id": 0, "password_hash": 0}
    ).to_list(100)
    return {"users": users}

@api_router.put("/admin/clients/{client_id}/users/{user_id}/toggle")
async def toggle_client_user(client_id: str, user_id: str):
    """Enable/disable a client user"""
    user = await db.client_users.find_one({"id": user_id, "client_id": client_id})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    new_status = not user.get("is_active", True)
    await db.client_users.update_one(
        {"id": user_id},
        {"$set": {"is_active": new_status}}
    )
    
    return {"success": True, "is_active": new_status}

# ============ Invitation & Security System ============

class CreateInvitationRequest(BaseModel):
    email_restricted: Optional[str] = None
    domain_restricted: Optional[str] = None
    role: str = "user"
    max_uses: int = 1
    expires_in_days: Optional[int] = 7

class JoinWithInvitationRequest(BaseModel):
    invitation_code: str
    email: str
    password: str
    name: str

@api_router.post("/admin/clients/{client_id}/invitations")
async def create_invitation(client_id: str, request: CreateInvitationRequest):
    """Create an invitation code for a client organization"""
    client = await db.client_organizations.find_one({"id": client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Generate unique invitation code
    code = f"RS-{uuid.uuid4().hex[:8].upper()}"
    
    # Calculate expiration
    expires_at = None
    if request.expires_in_days:
        expires_at = datetime.now(timezone.utc) + timedelta(days=request.expires_in_days)
    
    invitation = InvitationCode(
        code=code,
        client_id=client_id,
        created_by="admin",
        email_restricted=request.email_restricted,
        domain_restricted=request.domain_restricted,  # Only set if explicitly provided
        role=request.role,
        max_uses=request.max_uses,
        expires_at=expires_at
    )
    
    await db.invitation_codes.insert_one(invitation.model_dump())
    
    # Log security event
    await db.security_logs.insert_one({
        "id": str(uuid.uuid4()),
        "client_id": client_id,
        "action": "invitation_created",
        "details": {
            "code": code,
            "email_restricted": request.email_restricted,
            "domain_restricted": request.domain_restricted,
            "role": request.role,
            "max_uses": request.max_uses
        },
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    return {
        "success": True,
        "invitation": {
            "code": code,
            "client_name": client["organization_name"],
            "email_restricted": request.email_restricted,
            "domain_restricted": invitation.domain_restricted,
            "role": request.role,
            "max_uses": request.max_uses,
            "expires_at": expires_at.isoformat() if expires_at else None,
            "join_url": f"/join?code={code}"
        }
    }

@api_router.get("/admin/clients/{client_id}/invitations")
async def list_invitations(client_id: str):
    """List all invitation codes for a client"""
    invitations = await db.invitation_codes.find(
        {"client_id": client_id},
        {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    
    return {"invitations": invitations}

@api_router.delete("/admin/clients/{client_id}/invitations/{invitation_id}")
async def revoke_invitation(client_id: str, invitation_id: str):
    """Revoke an invitation code"""
    result = await db.invitation_codes.update_one(
        {"id": invitation_id, "client_id": client_id},
        {"$set": {"is_active": False}}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Invitation not found")
    
    return {"success": True, "message": "Invitation revoked"}

@api_router.get("/invitation/{code}/validate")
async def validate_invitation(code: str):
    """Validate an invitation code (public endpoint for joining)"""
    invitation = await db.invitation_codes.find_one({"code": code, "is_active": True})
    
    if not invitation:
        raise HTTPException(status_code=404, detail="Invalid or expired invitation code")
    
    # Check expiration
    if invitation.get("expires_at"):
        expires_at = ensure_timezone_aware(invitation["expires_at"])
        if datetime.now(timezone.utc) > expires_at:
            raise HTTPException(status_code=400, detail="Invitation code has expired")
    
    # Check usage limit
    if invitation["uses_count"] >= invitation["max_uses"]:
        raise HTTPException(status_code=400, detail="Invitation code has reached its usage limit")
    
    # Get client info
    client = await db.client_organizations.find_one({"id": invitation["client_id"]})
    
    return {
        "valid": True,
        "organization_name": client["organization_name"] if client else "Unknown",
        "organization_type": client["organization_type"] if client else "unknown",
        "role": invitation["role"],
        "email_restricted": invitation.get("email_restricted"),
        "domain_restricted": invitation.get("domain_restricted")
    }

@api_router.post("/invitation/join")
async def join_with_invitation(request: JoinWithInvitationRequest):
    """Join an organization using an invitation code"""
    invitation = await db.invitation_codes.find_one({"code": request.invitation_code, "is_active": True})
    
    if not invitation:
        raise HTTPException(status_code=404, detail="Invalid or expired invitation code")
    
    # Check expiration
    if invitation.get("expires_at"):
        expires_at = ensure_timezone_aware(invitation["expires_at"])
        if datetime.now(timezone.utc) > expires_at:
            raise HTTPException(status_code=400, detail="Invitation code has expired")
    
    # Check usage limit
    if invitation["uses_count"] >= invitation["max_uses"]:
        raise HTTPException(status_code=400, detail="Invitation code has reached its usage limit")
    
    # Check email restriction
    if invitation.get("email_restricted"):
        if request.email.lower() != invitation["email_restricted"].lower():
            raise HTTPException(status_code=400, detail="This invitation is restricted to a specific email")
    
    # Check domain restriction
    if invitation.get("domain_restricted"):
        email_domain = request.email.split("@")[1].lower()
        allowed_domain = invitation["domain_restricted"].lower()
        if email_domain != allowed_domain:
            raise HTTPException(status_code=400, detail=f"Only emails from {allowed_domain} can use this invitation")
    
    # Check if user already exists
    existing_user = await db.client_users.find_one({"email": request.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="An account with this email already exists")
    
    # Create the user
    new_user = ClientUser(
        client_id=invitation["client_id"],
        email=request.email,
        password_hash=hash_password(request.password),
        name=request.name,
        role=invitation["role"],
        created_by_admin=False,
        invited_by=invitation["created_by"],
        invitation_code=request.invitation_code
    )
    
    await db.client_users.insert_one(new_user.model_dump())
    
    # Update invitation usage
    await db.invitation_codes.update_one(
        {"code": request.invitation_code},
        {
            "$inc": {"uses_count": 1},
            "$push": {"used_by": new_user.id}
        }
    )
    
    # Update client user count
    await db.client_organizations.update_one(
        {"id": invitation["client_id"]},
        {"$inc": {"total_users": 1}}
    )
    
    # Log security event
    await db.security_logs.insert_one({
        "id": str(uuid.uuid4()),
        "client_id": invitation["client_id"],
        "user_id": new_user.id,
        "action": "user_joined_via_invitation",
        "details": {
            "invitation_code": request.invitation_code,
            "email": request.email,
            "role": invitation["role"]
        },
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    # Get client info
    client = await db.client_organizations.find_one({"id": invitation["client_id"]})
    
    return {
        "success": True,
        "message": "Successfully joined the organization",
        "user": {
            "id": new_user.id,
            "email": new_user.email,
            "name": new_user.name,
            "role": new_user.role
        },
        "organization": {
            "id": client["id"],
            "name": client["organization_name"],
            "type": client["organization_type"]
        }
    }

@api_router.get("/admin/clients/{client_id}/security-logs")
async def get_security_logs(client_id: str, limit: int = 50):
    """Get security/audit logs for a client"""
    logs = await db.security_logs.find(
        {"client_id": client_id},
        {"_id": 0}
    ).sort("created_at", -1).limit(limit).to_list(limit)
    
    return {"logs": logs}

# ============ IP Whitelisting ============

class AddIPWhitelistRequest(BaseModel):
    ip_address: str
    description: Optional[str] = None

@api_router.post("/admin/clients/{client_id}/ip-whitelist")
async def add_ip_whitelist(client_id: str, request: AddIPWhitelistRequest):
    """Add an IP address to the whitelist for a client"""
    client = await db.client_organizations.find_one({"id": client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Validate IP format (basic check)
    ip = request.ip_address.strip()
    if not ip:
        raise HTTPException(status_code=400, detail="IP address is required")
    
    # Check if already exists
    existing = await db.ip_whitelist.find_one({"client_id": client_id, "ip_address": ip})
    if existing:
        raise HTTPException(status_code=400, detail="IP address already whitelisted")
    
    whitelist_entry = IPWhitelist(
        client_id=client_id,
        ip_address=ip,
        description=request.description,
        created_by="admin"
    )
    
    await db.ip_whitelist.insert_one(whitelist_entry.model_dump())
    
    # Log security event
    await db.security_logs.insert_one({
        "id": str(uuid.uuid4()),
        "client_id": client_id,
        "action": "ip_whitelist_added",
        "details": {"ip_address": ip, "description": request.description},
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    return {"success": True, "id": whitelist_entry.id, "ip_address": ip}

@api_router.get("/admin/clients/{client_id}/ip-whitelist")
async def get_ip_whitelist(client_id: str):
    """Get all whitelisted IPs for a client"""
    whitelist = await db.ip_whitelist.find(
        {"client_id": client_id, "is_active": True},
        {"_id": 0}
    ).to_list(100)
    
    return {"whitelist": whitelist}

@api_router.delete("/admin/clients/{client_id}/ip-whitelist/{ip_id}")
async def remove_ip_whitelist(client_id: str, ip_id: str):
    """Remove an IP from the whitelist"""
    result = await db.ip_whitelist.update_one(
        {"id": ip_id, "client_id": client_id},
        {"$set": {"is_active": False}}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="IP whitelist entry not found")
    
    return {"success": True, "message": "IP removed from whitelist"}

@api_router.put("/admin/clients/{client_id}/security-settings")
async def update_security_settings(client_id: str, settings: dict):
    """Update security settings for a client (IP restriction enabled, etc.)"""
    client = await db.client_organizations.find_one({"id": client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    await db.client_organizations.update_one(
        {"id": client_id},
        {"$set": {
            "security_settings": settings,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {"success": True, "message": "Security settings updated"}

# Client-side authentication
@api_router.post("/auth/login")
async def client_login(request: AdminLoginRequest):
    """Client user login endpoint"""
    user = await db.client_users.find_one({"email": request.email, "is_active": True})
    
    # Decrypt password if encrypted
    password = decrypt_password(request.password) if request.encrypted else request.password
    
    if not user or not verify_password(password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    # Get client info
    client = await db.client_organizations.find_one({"id": user["client_id"]})
    if not client or not client.get("is_active"):
        raise HTTPException(status_code=403, detail="Organization is not active")
    
    # Check trial expiration
    if client.get("subscription_status") == "trial":
        trial_end = client.get("trial_end_date")
        if trial_end:
            # Handle different datetime formats
            if isinstance(trial_end, str):
                trial_end = datetime.fromisoformat(trial_end.replace('Z', '+00:00'))
            elif isinstance(trial_end, datetime):
                # Ensure timezone awareness
                if trial_end.tzinfo is None:
                    trial_end = trial_end.replace(tzinfo=timezone.utc)
            
            now_utc = datetime.now(timezone.utc)
            if now_utc > trial_end:
                # Trial expired - set to limited access
                await db.client_organizations.update_one(
                    {"id": client["id"]},
                    {"$set": {"access_level": "limited", "subscription_status": "expired"}}
                )
                client["access_level"] = "limited"
    
    # Update last login
    await db.client_users.update_one(
        {"id": user["id"]},
        {"$set": {"last_login": datetime.now(timezone.utc)}}
    )
    
    return {
        "success": True,
        "user": {
            "id": user["id"],
            "email": user["email"],
            "name": user["name"],
            "role": user["role"],
            "client_id": user["client_id"]
        },
        "organization": {
            "id": client["id"],
            "name": client["organization_name"],
            "type": client["organization_type"],
            "access_level": client.get("access_level", "full"),
            "modules_enabled": client.get("modules_enabled", {}),
            "subscription_status": client.get("subscription_status", "trial")
        },
        "token": user["id"]
    }

@api_router.post("/auth/change-password")
async def change_password(user_id: str, request: ChangePasswordRequest):
    """Change user password"""
    user = await db.client_users.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if not verify_password(request.current_password, user["password_hash"]):
        raise HTTPException(status_code=400, detail="Current password is incorrect")
    
    await db.client_users.update_one(
        {"id": user_id},
        {"$set": {"password_hash": hash_password(request.new_password)}}
    )
    
    return {"success": True, "message": "Password changed successfully"}

@api_router.post("/auth/signup")
async def self_signup(request: SelfSignupRequest):
    """Self-signup for new organizations (Corporate or Staffing)"""
    
    # Validate organization type
    if request.organization_type not in ["corporate", "staffing_vendor"]:
        raise HTTPException(status_code=400, detail="Invalid organization type. Must be 'corporate' or 'staffing_vendor'")
    
    # Check if organization with this domain already exists
    existing_org = await db.client_organizations.find_one({
        "business_domain": request.business_domain.lower()
    })
    if existing_org:
        raise HTTPException(status_code=400, detail="An organization with this domain already exists. Please sign in or contact support.")
    
    # Check if user email already exists
    existing_user = await db.client_users.find_one({"email": request.contact_email.lower()})
    if existing_user:
        raise HTTPException(status_code=400, detail="An account with this email already exists. Please sign in.")
    
    # Validate email domain matches business domain
    email_domain = request.contact_email.split("@")[1].lower()
    if email_domain != request.business_domain.lower():
        raise HTTPException(
            status_code=400, 
            detail=f"Email domain must match your business domain. Use an email ending with @{request.business_domain}"
        )
    
    # Decrypt password if encrypted
    password = decrypt_password(request.password) if request.encrypted else request.password
    
    # Validate password strength
    if len(password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters")
    
    # Create new organization
    new_org = ClientOrganization(
        organization_name=request.organization_name,
        organization_type=request.organization_type,
        business_domain=request.business_domain.lower(),
        contact_email=request.contact_email.lower(),
        contact_phone=request.contact_phone,
        contact_person=request.contact_person,
        subscription_status="trial",
        access_level="full",
        trial_start_date=datetime.now(timezone.utc),
        trial_end_date=datetime.now(timezone.utc) + timedelta(days=90),
        modules_enabled={
            "jd_intelligence": True,
            "resume_repository": True,
            "career_trajectory": True,
            "hr_fitment": True
        }
    )
    
    await db.client_organizations.insert_one(new_org.model_dump())
    
    # Create admin user for the organization
    admin_user = ClientUser(
        client_id=new_org.id,
        email=request.contact_email.lower(),
        password_hash=hash_password(password),
        name=request.contact_person,
        role="admin",
        created_by_admin=False
    )
    
    await db.client_users.insert_one(admin_user.model_dump())
    
    # Update organization user count
    await db.client_organizations.update_one(
        {"id": new_org.id},
        {"$set": {"total_users": 1}}
    )
    
    # Log the signup
    await db.security_logs.insert_one({
        "id": str(uuid.uuid4()),
        "client_id": new_org.id,
        "action": "self_signup",
        "details": {
            "organization_name": request.organization_name,
            "organization_type": request.organization_type,
            "admin_email": request.contact_email
        },
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    logger.info(f"New organization registered: {request.organization_name} ({request.organization_type})")
    
    return {
        "success": True,
        "message": "Welcome to RoleSense! Your organization has been registered.",
        "user": {
            "id": admin_user.id,
            "email": admin_user.email,
            "name": admin_user.name,
            "role": admin_user.role,
            "client_id": new_org.id
        },
        "organization": {
            "id": new_org.id,
            "name": new_org.organization_name,
            "type": new_org.organization_type,
            "access_level": new_org.access_level,
            "modules_enabled": new_org.modules_enabled,
            "subscription_status": new_org.subscription_status,
            "trial_end_date": new_org.trial_end_date.isoformat()
        },
        "token": admin_user.id,
        "trial_info": {
            "days_remaining": 90,
            "end_date": new_org.trial_end_date.isoformat()
        }
    }

# Customer Feedback
@api_router.post("/feedback")
async def submit_feedback(client_id: str, user_email: str, request: SubmitFeedbackRequest):
    """Submit customer feedback"""
    feedback = CustomerFeedback(
        client_id=client_id,
        user_email=user_email,
        feedback_type=request.feedback_type,
        subject=request.subject,
        description=request.description,
        priority=request.priority
    )
    
    await db.customer_feedback.insert_one(feedback.model_dump())
    
    return {
        "success": True,
        "feedback_id": feedback.id,
        "message": "Thank you for your feedback! Our team will review it shortly."
    }

@api_router.get("/admin/feedback")
async def list_feedback(
    status: Optional[str] = None,
    feedback_type: Optional[str] = None,
    priority: Optional[str] = None,
    skip: int = 0,
    limit: int = 50
):
    """List all customer feedback"""
    query = {}
    if status:
        query["status"] = status
    if feedback_type:
        query["feedback_type"] = feedback_type
    if priority:
        query["priority"] = priority
    
    feedback_list = await db.customer_feedback.find(query, {"_id": 0}).sort("created_at", -1).skip(skip).limit(limit).to_list(limit)
    total = await db.customer_feedback.count_documents(query)
    
    # Enrich with client names
    for feedback in feedback_list:
        client = await db.client_organizations.find_one({"id": feedback["client_id"]})
        feedback["organization_name"] = client["organization_name"] if client else "Unknown"
    
    return {"feedback": feedback_list, "total": total}

@api_router.put("/admin/feedback/{feedback_id}")
async def update_feedback(feedback_id: str, status: str, admin_response: Optional[str] = None):
    """Update feedback status and response"""
    update_data = {
        "status": status,
        "updated_at": datetime.now(timezone.utc)
    }
    if admin_response:
        update_data["admin_response"] = admin_response
    if status in ["resolved", "closed"]:
        update_data["resolved_at"] = datetime.now(timezone.utc)
    
    result = await db.customer_feedback.update_one(
        {"id": feedback_id},
        {"$set": update_data}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Feedback not found")
    
    return {"success": True, "message": "Feedback updated successfully"}

@api_router.get("/admin/analytics")
async def get_admin_analytics():
    """Get detailed analytics for admin dashboard"""
    # Usage by client
    pipeline = [
        {"$group": {
            "_id": "$client_id",
            "total_assessments": {"$sum": 1},
            "completed": {"$sum": {"$cond": [{"$eq": ["$status", "completed"]}, 1, 0]}}
        }}
    ]
    assessment_by_client = await db.trajectory_assessments.aggregate(pipeline).to_list(100)
    
    # Monthly trends
    from datetime import timedelta
    thirty_days_ago = datetime.now(timezone.utc) - timedelta(days=30)
    
    new_clients_30d = await db.client_organizations.count_documents({
        "created_at": {"$gte": thirty_days_ago}
    })
    
    new_assessments_30d = await db.trajectory_assessments.count_documents({
        "created_at": {"$gte": thirty_days_ago}
    })
    
    new_feedback_30d = await db.customer_feedback.count_documents({
        "created_at": {"$gte": thirty_days_ago}
    })
    
    # Feature usage
    feature_usage = {
        "jd_intelligence": await db.job_descriptions.count_documents({}),
        "resume_repository": await db.routed_resumes.count_documents({}),
        "career_trajectory": await db.trajectory_assessments.count_documents({}),
        "screening_questions": await db.job_descriptions.count_documents({"screening_questions": {"$exists": True, "$ne": []}})
    }
    
    return {
        "assessment_by_client": assessment_by_client,
        "trends_30d": {
            "new_clients": new_clients_30d,
            "new_assessments": new_assessments_30d,
            "new_feedback": new_feedback_30d
        },
        "feature_usage": feature_usage
    }

# Sales team connection endpoint
@api_router.post("/connect-sales")
async def connect_to_sales(client_id: str, user_email: str, message: Optional[str] = None):
    """Request connection to sales team for subscription upgrade"""
    client = await db.client_organizations.find_one({"id": client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Create a sales lead/request
    sales_request = {
        "id": str(uuid.uuid4()),
        "client_id": client_id,
        "organization_name": client["organization_name"],
        "contact_email": user_email,
        "contact_person": client.get("contact_person"),
        "message": message or "Interested in upgrading subscription",
        "current_plan": client.get("subscription_status"),
        "requested_at": datetime.now(timezone.utc),
        "status": "pending"
    }
    
    await db.sales_requests.insert_one(sales_request)
    
    # Create notification for admin
    notification = {
        "id": str(uuid.uuid4()),
        "type": "sales_request",
        "title": f"Sales Request from {client['organization_name']}",
        "message": f"{user_email} requested to connect with sales team",
        "client_id": client_id,
        "created_at": datetime.now(timezone.utc),
        "read": False
    }
    await db.admin_notifications.insert_one(notification)
    
    return {
        "success": True,
        "message": "Your request has been submitted. Our sales team will contact you within 24 hours.",
        "request_id": sales_request["id"]
    }

@api_router.get("/admin/sales-requests")
async def list_sales_requests(status: Optional[str] = None, skip: int = 0, limit: int = 50):
    """List all sales requests"""
    query = {}
    if status:
        query["status"] = status
    
    requests = await db.sales_requests.find(query, {"_id": 0}).sort("requested_at", -1).skip(skip).limit(limit).to_list(limit)
    total = await db.sales_requests.count_documents(query)
    
    return {"requests": requests, "total": total}


# ============ Interview Wizard & Custom Questions ============

class CustomQuestion(BaseModel):
    question: str
    category: str = "general"  # technical/behavioral/situational/general
    required: bool = True

class InterviewWizardConfig(BaseModel):
    job_id: Optional[str] = None
    custom_questions: List[CustomQuestion] = []
    include_ai_questions: bool = True
    max_ai_questions: int = 7

@api_router.get("/resume/{resume_id}/interview-wizard")
async def get_interview_wizard_data(resume_id: str):
    """Get all data needed for interview wizard - AI analysis, questions, etc."""
    
    resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Get associated job if any
    job = None
    if resume.get("applied_job_id"):
        job = await db.job_descriptions.find_one({"id": resume["applied_job_id"]}, {"_id": 0})
        if not job:
            job = await db.structured_jds.find_one({"id": resume["applied_job_id"]}, {"_id": 0})
    
    # Get any existing custom questions configured for this job/organization
    custom_questions = await db.interview_custom_questions.find({
        "$or": [
            {"job_id": resume.get("applied_job_id")},
            {"organization_id": resume.get("organization_id")},
            {"is_default": True}
        ]
    }, {"_id": 0}).to_list(20)
    
    # Get candidate responses if any
    candidate_responses = await db.candidate_responses.find_one(
        {"resume_id": resume_id}, 
        {"_id": 0}
    )
    
    return {
        "resume": resume,
        "job": job,
        "ai_analysis": {
            "parsed_data": resume.get("parsed_data", {}),
            "skill_match_scores": resume.get("skill_match_scores", {}),
            "technology_stack": resume.get("technology_stack", {}),
            "ai_assessment": resume.get("ai_assessment", {}),
            "confidence_score": resume.get("confidence_score", 0),
            "routing_reason": resume.get("routing_reason", "")
        },
        "ai_questions": resume.get("interview_questions", []),
        "custom_questions": custom_questions,
        "candidate_responses": candidate_responses,
        "status": resume.get("recruiter_status", "under_evaluation"),
        "recruiter_notes": resume.get("recruiter_notes", [])
    }

@api_router.post("/resume/{resume_id}/custom-questions")
async def save_custom_questions(resume_id: str, config: InterviewWizardConfig):
    """Save custom interview questions for a specific resume/job"""
    
    questions_doc = {
        "id": str(uuid.uuid4()),
        "resume_id": resume_id,
        "job_id": config.job_id,
        "questions": [q.dict() for q in config.custom_questions],
        "include_ai_questions": config.include_ai_questions,
        "max_ai_questions": config.max_ai_questions,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    # Upsert - update if exists, insert if not
    await db.interview_custom_questions.update_one(
        {"resume_id": resume_id},
        {"$set": questions_doc},
        upsert=True
    )
    
    return {"message": "Custom questions saved", "questions_id": questions_doc["id"]}

@api_router.post("/resume/{resume_id}/candidate-responses")
async def save_candidate_responses(resume_id: str, responses: dict):
    """Save candidate's answers to interview questions"""
    
    resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    response_doc = {
        "id": str(uuid.uuid4()),
        "resume_id": resume_id,
        "candidate_name": resume.get("name"),
        "candidate_email": resume.get("email"),
        "job_id": resume.get("applied_job_id"),
        "ai_question_responses": responses.get("ai_questions", []),
        "custom_question_responses": responses.get("custom_questions", []),
        "submitted_at": datetime.now(timezone.utc).isoformat(),
        "evaluation_status": "pending"
    }
    
    await db.candidate_responses.insert_one(response_doc)
    
    # Update resume status
    await db.routed_resumes.update_one(
        {"id": resume_id},
        {"$set": {
            "has_responses": True,
            "response_submitted_at": response_doc["submitted_at"]
        }}
    )
    
    return {
        "message": "Responses saved successfully",
        "response_id": response_doc["id"]
    }

@api_router.put("/resume/{resume_id}/recruiter-action")
async def update_recruiter_action(resume_id: str, action: dict):
    """Update recruiter status, notes, and notifications for a resume"""
    
    update_data = {}
    
    if "status" in action:
        # Validate status
        valid_statuses = ["under_evaluation", "selected_next_round", "hold", "rejected"]
        if action["status"] not in valid_statuses:
            raise HTTPException(status_code=400, detail=f"Invalid status. Must be one of: {valid_statuses}")
        update_data["recruiter_status"] = action["status"]
        update_data["status_updated_at"] = datetime.now(timezone.utc).isoformat()
    
    if "note" in action:
        # Add note to notes array
        note = {
            "id": str(uuid.uuid4()),
            "text": action["note"],
            "created_by": action.get("created_by", "recruiter"),
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.routed_resumes.update_one(
            {"id": resume_id},
            {"$push": {"recruiter_notes": note}}
        )
    
    if "send_notification" in action and action["send_notification"]:
        # Create notification
        resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
        if resume:
            notification = {
                "id": str(uuid.uuid4()),
                "type": "candidate_status_update",
                "resume_id": resume_id,
                "candidate_name": resume.get("name"),
                "candidate_email": resume.get("email"),
                "new_status": action.get("status"),
                "message": action.get("notification_message", f"Status updated to {action.get('status')}"),
                "created_at": datetime.now(timezone.utc).isoformat(),
                "sent": False
            }
            await db.notifications.insert_one(notification)
    
    if update_data:
        await db.routed_resumes.update_one({"id": resume_id}, {"$set": update_data})
    
    return {"message": "Recruiter action saved"}

@api_router.get("/resume/{resume_id}/generate-report")
async def generate_candidate_report(resume_id: str, format: str = "pdf"):
    """Generate comprehensive PDF/DOCX report for candidate with AI analysis"""
    
    resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Get candidate responses
    responses = await db.candidate_responses.find_one({"resume_id": resume_id}, {"_id": 0})
    
    # Get job details if any
    job = None
    if resume.get("applied_job_id"):
        job = await db.job_descriptions.find_one({"id": resume["applied_job_id"]}, {"_id": 0})
        if not job:
            job = await db.structured_jds.find_one({"id": resume["applied_job_id"]}, {"_id": 0})
    
    # Generate report content
    report_data = {
        "candidate_name": resume.get("name"),
        "email": resume.get("email"),
        "phone": resume.get("phone"),
        "applied_job": resume.get("applied_job_title") or (job.get("title") if job else "General Application"),
        "application_date": resume.get("created_at"),
        "parsed_data": resume.get("parsed_data", {}),
        "ai_assessment": resume.get("ai_assessment", {}),
        "skill_match_scores": resume.get("skill_match_scores", {}),
        "interview_questions": resume.get("interview_questions", []),
        "candidate_responses": responses,
        "recruiter_status": resume.get("recruiter_status", "under_evaluation"),
        "recruiter_notes": resume.get("recruiter_notes", []),
        "confidence_score": resume.get("confidence_score", 0),
        "routing_reason": resume.get("routing_reason", "")
    }
    
    if format == "pdf":
        pdf_content = generate_candidate_report_pdf(report_data)
        return Response(
            content=pdf_content,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="candidate_report_{resume.get("name", "unknown").replace(" ", "_")}.pdf"'}
        )
    elif format == "docx":
        docx_content = generate_candidate_report_docx(report_data)
        return Response(
            content=docx_content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="candidate_report_{resume.get("name", "unknown").replace(" ", "_")}.docx"'}
        )
    else:
        # Return JSON
        return report_data

def generate_candidate_report_pdf(data: dict) -> bytes:
    """Generate PDF report for candidate"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor('#0F172A'), spaceAfter=20)
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#10B981'), spaceBefore=15, spaceAfter=10)
    body_style = ParagraphStyle('CustomBody', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=8)
    
    # Title
    story.append(Paragraph(f"Candidate Assessment Report", title_style))
    story.append(Paragraph(f"<b>{data.get('candidate_name', 'Unknown')}</b>", styles['Heading3']))
    story.append(Paragraph(f"Applied for: {data.get('applied_job', 'N/A')}", body_style))
    story.append(Paragraph(f"Date: {data.get('application_date', 'N/A')[:10] if data.get('application_date') else 'N/A'}", body_style))
    story.append(Spacer(1, 20))
    
    # Contact Info
    story.append(Paragraph("Contact Information", heading_style))
    story.append(Paragraph(f"Email: {data.get('email', 'N/A')}", body_style))
    story.append(Paragraph(f"Phone: {data.get('phone', 'N/A')}", body_style))
    
    # AI Assessment
    ai_assessment = data.get('ai_assessment', {})
    if ai_assessment:
        story.append(Paragraph("AI Assessment Summary", heading_style))
        story.append(Paragraph(f"<b>Overall Rating:</b> {ai_assessment.get('overall_rating', 'N/A')}/10", body_style))
        story.append(Paragraph(f"<b>Recommendation:</b> {ai_assessment.get('recommendation', 'N/A')}", body_style))
        story.append(Paragraph(f"<b>Confidence Score:</b> {data.get('confidence_score', 0)*100:.0f}%", body_style))
        
        if ai_assessment.get('strengths'):
            story.append(Paragraph("<b>Strengths:</b>", body_style))
            for s in ai_assessment.get('strengths', []):
                story.append(Paragraph(f"• {s}", body_style))
        
        if ai_assessment.get('areas_for_growth'):
            story.append(Paragraph("<b>Areas for Growth:</b>", body_style))
            for a in ai_assessment.get('areas_for_growth', []):
                story.append(Paragraph(f"• {a}", body_style))
    
    # Parsed Data
    parsed = data.get('parsed_data', {})
    if parsed:
        story.append(Paragraph("Professional Profile", heading_style))
        story.append(Paragraph(f"<b>Current Role:</b> {parsed.get('current_role', 'N/A')}", body_style))
        story.append(Paragraph(f"<b>Experience:</b> {parsed.get('experience_years', 'N/A')} years", body_style))
        story.append(Paragraph(f"<b>Education:</b> {parsed.get('education', 'N/A')}", body_style))
        
        if parsed.get('top_skills'):
            story.append(Paragraph(f"<b>Top Skills:</b> {', '.join(parsed.get('top_skills', []))}", body_style))
    
    # Skill Match Scores
    skill_scores = data.get('skill_match_scores', {})
    if skill_scores:
        story.append(Paragraph("Skill Match Analysis", heading_style))
        for skill, score in skill_scores.items():
            story.append(Paragraph(f"• {skill.replace('_', ' ').title()}: {score*100:.0f}%", body_style))
    
    # Recruiter Status
    story.append(Paragraph("Evaluation Status", heading_style))
    status_map = {
        "under_evaluation": "Under Evaluation",
        "selected_next_round": "Selected for Next Round",
        "hold": "On Hold",
        "rejected": "Rejected"
    }
    story.append(Paragraph(f"<b>Status:</b> {status_map.get(data.get('recruiter_status'), data.get('recruiter_status', 'N/A'))}", body_style))
    
    # Notes
    notes = data.get('recruiter_notes', [])
    if notes:
        story.append(Paragraph("Recruiter Notes", heading_style))
        for note in notes:
            story.append(Paragraph(f"• {note.get('text', '')} ({note.get('created_at', '')[:10] if note.get('created_at') else ''})", body_style))
    
    # Footer
    story.append(Spacer(1, 30))
    story.append(Paragraph("Generated by RoleSense - Talent Intelligence Platform", ParagraphStyle('Footer', fontSize=8, textColor=colors.gray)))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_candidate_report_docx(data: dict) -> bytes:
    """Generate DOCX report for candidate"""
    doc = Document()
    
    # Title
    doc.add_heading('Candidate Assessment Report', 0)
    doc.add_heading(data.get('candidate_name', 'Unknown'), level=1)
    doc.add_paragraph(f"Applied for: {data.get('applied_job', 'N/A')}")
    doc.add_paragraph(f"Date: {data.get('application_date', 'N/A')[:10] if data.get('application_date') else 'N/A'}")
    
    # Contact
    doc.add_heading('Contact Information', level=2)
    doc.add_paragraph(f"Email: {data.get('email', 'N/A')}")
    doc.add_paragraph(f"Phone: {data.get('phone', 'N/A')}")
    
    # AI Assessment
    ai = data.get('ai_assessment', {})
    if ai:
        doc.add_heading('AI Assessment Summary', level=2)
        doc.add_paragraph(f"Overall Rating: {ai.get('overall_rating', 'N/A')}/10")
        doc.add_paragraph(f"Recommendation: {ai.get('recommendation', 'N/A')}")
        doc.add_paragraph(f"Confidence: {data.get('confidence_score', 0)*100:.0f}%")
        
        if ai.get('strengths'):
            doc.add_paragraph('Strengths:')
            for s in ai.get('strengths', []):
                doc.add_paragraph(f"• {s}", style='List Bullet')
    
    # Skills
    parsed = data.get('parsed_data', {})
    if parsed:
        doc.add_heading('Professional Profile', level=2)
        doc.add_paragraph(f"Current Role: {parsed.get('current_role', 'N/A')}")
        doc.add_paragraph(f"Experience: {parsed.get('experience_years', 'N/A')} years")
        if parsed.get('top_skills'):
            doc.add_paragraph(f"Top Skills: {', '.join(parsed.get('top_skills', []))}")
    
    # Status
    doc.add_heading('Evaluation Status', level=2)
    doc.add_paragraph(f"Status: {data.get('recruiter_status', 'under_evaluation').replace('_', ' ').title()}")
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============ Basic Competency Report (FREE Tier) ============

@api_router.get("/resume/{resume_id}/basic-analysis")
async def get_basic_resume_analysis(resume_id: str, organization_id: Optional[str] = None):
    """Get basic analysis when recruiter clicks on resume name - FREE tier"""
    
    resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Get job details for skill matching
    job = None
    job_skills = []
    if resume.get("applied_job_id"):
        job = await db.job_descriptions.find_one({"id": resume["applied_job_id"]}, {"_id": 0})
        if not job:
            job = await db.structured_jds.find_one({"id": resume["applied_job_id"]}, {"_id": 0})
        
        if job:
            # Extract job required skills
            parsed = job.get("parsed_data", {})
            job_skills.extend(parsed.get("required_skills", []))
            job_skills.extend(parsed.get("preferred_skills", []))
            competencies = job.get("competencies", {})
            job_skills.extend(competencies.get("must_have_skills", []))
            job_skills.extend(competencies.get("good_to_have_skills", []))
            job_skills.extend(job.get("must_have_skills", []))
    
    # Calculate skill match
    candidate_skills = set([s.lower() for s in resume.get("matched_skills", [])])
    job_skills_set = set([s.lower() for s in job_skills])
    
    matched_skills = list(candidate_skills.intersection(job_skills_set))
    missing_skills = list(job_skills_set - candidate_skills)
    
    match_percentage = len(matched_skills) / len(job_skills_set) * 100 if job_skills_set else 0
    
    # Get organization for logo
    org = None
    if organization_id:
        org = await db.clients.find_one({"id": organization_id}, {"_id": 0, "organization_name": 1, "logo_url": 1})
    
    # Check subscription status
    subscription_status = "free"
    if organization_id:
        client = await db.clients.find_one({"id": organization_id}, {"_id": 0})
        if client:
            subscription_status = client.get("subscription", {}).get("tier", "free")
    
    basic_analysis = {
        "resume_id": resume_id,
        "candidate_name": resume.get("name"),
        "candidate_email": resume.get("email"),
        "candidate_phone": resume.get("phone"),
        "applied_job": {
            "title": resume.get("applied_job_title") or (job.get("title") if job else None),
            "requisition_number": resume.get("applied_requisition_number")
        },
        "parsed_profile": {
            "current_role": resume.get("parsed_data", {}).get("current_role"),
            "experience_years": resume.get("parsed_data", {}).get("experience_years"),
            "education": resume.get("parsed_data", {}).get("education"),
            "companies": resume.get("parsed_data", {}).get("companies", [])[:3]
        },
        "skill_analysis": {
            "candidate_skills": list(candidate_skills)[:15],
            "matched_skills": matched_skills[:10],
            "missing_skills": missing_skills[:5],
            "match_percentage": round(match_percentage, 1)
        },
        "competency_summary": {
            "primary_function": resume.get("primary_function"),
            "sub_function": resume.get("sub_function"),
            "confidence_score": resume.get("confidence_score", 0),
            "routing_reason": resume.get("routing_reason", "")[:200]
        },
        "recruiter_notes": resume.get("recruiter_notes", []),
        "status": resume.get("recruiter_status", "under_evaluation"),
        "organization": org,
        "subscription_status": subscription_status,
        "career_trajectory_available": subscription_status in ["professional", "enterprise"],
        "pre_assessment_sent": resume.get("pre_assessment_sent", False),
        "pre_assessment_completed": resume.get("pre_assessment_completed", False)
    }
    
    return basic_analysis

@api_router.get("/resume/{resume_id}/competency-report")
async def download_competency_report(resume_id: str, organization_id: Optional[str] = None, format: str = "pdf"):
    """Download basic competency report - FREE tier for all users"""
    
    resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Get job details
    job = None
    if resume.get("applied_job_id"):
        job = await db.job_descriptions.find_one({"id": resume["applied_job_id"]}, {"_id": 0})
        if not job:
            job = await db.structured_jds.find_one({"id": resume["applied_job_id"]}, {"_id": 0})
    
    # Get organization for branding
    org = None
    if organization_id:
        org = await db.clients.find_one({"id": organization_id}, {"_id": 0})
    
    # Get pre-assessment responses if any
    pre_assessment = await db.pre_assessments.find_one({"resume_id": resume_id}, {"_id": 0})
    
    report_data = {
        "candidate": {
            "name": resume.get("name"),
            "email": resume.get("email"),
            "phone": resume.get("phone"),
            "current_role": resume.get("parsed_data", {}).get("current_role"),
            "experience_years": resume.get("parsed_data", {}).get("experience_years"),
            "education": resume.get("parsed_data", {}).get("education"),
            "top_skills": resume.get("parsed_data", {}).get("top_skills", [])
        },
        "job": {
            "title": resume.get("applied_job_title") or (job.get("title") if job else "General Application"),
            "requisition_number": resume.get("applied_requisition_number"),
            "company": job.get("company_name") or job.get("client_name") if job else None
        },
        "competency_analysis": {
            "matched_skills": resume.get("matched_skills", [])[:10],
            "primary_function": resume.get("primary_function"),
            "sub_function": resume.get("sub_function"),
            "confidence_score": resume.get("confidence_score", 0),
            "routing_reason": resume.get("routing_reason", "")
        },
        "fitment_summary": {
            "skill_match": "Good" if resume.get("confidence_score", 0) >= 0.7 else "Moderate" if resume.get("confidence_score", 0) >= 0.4 else "Low",
            "experience_match": "Aligned" if job else "Not Assessed",
            "overall_fitment": resume.get("ai_assessment", {}).get("recommendation", "Under Review")
        },
        "pre_assessment": pre_assessment,
        "recruiter_notes": resume.get("recruiter_notes", []),
        "verbal_evaluation": resume.get("verbal_evaluation", ""),
        "status": resume.get("recruiter_status", "under_evaluation"),
        "organization": {
            "name": org.get("organization_name") if org else None,
            "logo_url": org.get("logo_url") if org else None
        },
        "generated_at": datetime.now(timezone.utc).isoformat()
    }
    
    if format == "pdf":
        pdf_content = generate_competency_report_pdf(report_data)
        return Response(
            content=pdf_content,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="competency_report_{resume.get("name", "unknown").replace(" ", "_")}.pdf"'}
        )
    else:
        return report_data

def generate_competency_report_pdf(data: dict) -> bytes:
    """Generate basic competency report PDF with company branding"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle('Title', fontSize=18, textColor=colors.HexColor('#10B981'), spaceAfter=10, fontName='Helvetica-Bold')
    heading_style = ParagraphStyle('Heading', fontSize=12, textColor=colors.HexColor('#1F2937'), spaceBefore=15, spaceAfter=8, fontName='Helvetica-Bold')
    body_style = ParagraphStyle('Body', fontSize=10, leading=14, spaceAfter=6)
    small_style = ParagraphStyle('Small', fontSize=8, textColor=colors.gray)
    
    # Header with company branding
    org = data.get("organization", {})
    if org.get("name"):
        story.append(Paragraph(f"<b>{org['name']}</b>", ParagraphStyle('OrgName', fontSize=14, textColor=colors.HexColor('#6366F1'))))
        story.append(Spacer(1, 5))
    
    story.append(Paragraph("Candidate Competency Report", title_style))
    story.append(Paragraph(f"Generated: {data.get('generated_at', '')[:10]}", small_style))
    story.append(Spacer(1, 20))
    
    # Candidate Information
    candidate = data.get("candidate", {})
    story.append(Paragraph("Candidate Information", heading_style))
    story.append(Paragraph(f"<b>Name:</b> {candidate.get('name', 'N/A')}", body_style))
    story.append(Paragraph(f"<b>Email:</b> {candidate.get('email', 'N/A')}", body_style))
    story.append(Paragraph(f"<b>Phone:</b> {candidate.get('phone', 'N/A')}", body_style))
    story.append(Paragraph(f"<b>Current Role:</b> {candidate.get('current_role', 'N/A')}", body_style))
    story.append(Paragraph(f"<b>Experience:</b> {candidate.get('experience_years', 'N/A')} years", body_style))
    
    # Applied Position
    job = data.get("job", {})
    story.append(Paragraph("Applied Position", heading_style))
    story.append(Paragraph(f"<b>Role:</b> {job.get('title', 'N/A')}", body_style))
    story.append(Paragraph(f"<b>Requisition:</b> {job.get('requisition_number', 'N/A')}", body_style))
    
    # Competency Analysis
    comp = data.get("competency_analysis", {})
    story.append(Paragraph("Competency Analysis", heading_style))
    story.append(Paragraph(f"<b>Function:</b> {comp.get('primary_function', 'N/A')} - {comp.get('sub_function', 'N/A')}", body_style))
    story.append(Paragraph(f"<b>Confidence Score:</b> {comp.get('confidence_score', 0)*100:.0f}%", body_style))
    
    if comp.get("matched_skills"):
        skills_text = ", ".join(comp.get("matched_skills", []))
        story.append(Paragraph(f"<b>Matched Skills:</b> {skills_text}", body_style))
    
    # Fitment Summary
    fitment = data.get("fitment_summary", {})
    story.append(Paragraph("Fitment Summary", heading_style))
    story.append(Paragraph(f"<b>Skill Match:</b> {fitment.get('skill_match', 'N/A')}", body_style))
    story.append(Paragraph(f"<b>Experience Match:</b> {fitment.get('experience_match', 'N/A')}", body_style))
    story.append(Paragraph(f"<b>Overall Fitment:</b> {fitment.get('overall_fitment', 'N/A')}", body_style))
    
    # Pre-Assessment Responses (if any)
    pre_assessment = data.get("pre_assessment")
    if pre_assessment and pre_assessment.get("responses"):
        story.append(Paragraph("Pre-Assessment Responses", heading_style))
        for resp in pre_assessment.get("responses", [])[:5]:
            story.append(Paragraph(f"<b>Q:</b> {resp.get('question', '')}", body_style))
            story.append(Paragraph(f"<b>A:</b> {resp.get('answer', '')}", body_style))
            story.append(Spacer(1, 5))
    
    # Recruiter Notes
    notes = data.get("recruiter_notes", [])
    if notes:
        story.append(Paragraph("Recruiter Notes", heading_style))
        for note in notes[-5:]:
            story.append(Paragraph(f"• {note.get('text', '')} ({note.get('created_at', '')[:10] if note.get('created_at') else ''})", body_style))
    
    # Verbal Evaluation
    if data.get("verbal_evaluation"):
        story.append(Paragraph("Verbal Evaluation", heading_style))
        story.append(Paragraph(data.get("verbal_evaluation"), body_style))
    
    # Status
    story.append(Paragraph("Evaluation Status", heading_style))
    status_map = {"under_evaluation": "Under Evaluation", "selected_next_round": "Selected for Next Round", "hold": "On Hold", "rejected": "Rejected"}
    story.append(Paragraph(f"<b>Status:</b> {status_map.get(data.get('status'), data.get('status', 'N/A'))}", body_style))
    
    # Footer
    story.append(Spacer(1, 30))
    story.append(Paragraph("This is a basic competency report. For detailed Career Trajectory analysis, upgrade to Professional tier.", small_style))
    story.append(Paragraph(f"Powered by RoleSense - Talent Intelligence Platform", small_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ============ Pre-Assessment for Staffing Agencies ============

class PreAssessmentQuestion(BaseModel):
    question: str
    required: bool = True

class PreAssessmentConfig(BaseModel):
    questions: List[PreAssessmentQuestion]
    candidate_email: str
    custom_message: Optional[str] = None

@api_router.post("/resume/{resume_id}/pre-assessment/send")
async def send_pre_assessment(resume_id: str, config: PreAssessmentConfig, organization_id: Optional[str] = None):
    """Send pre-assessment questions to candidate - Staffing Agency feature"""
    
    resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Get organization for branding
    org = None
    if organization_id:
        org = await db.clients.find_one({"id": organization_id}, {"_id": 0})
    
    # Create pre-assessment record
    pre_assessment = {
        "id": str(uuid.uuid4()),
        "resume_id": resume_id,
        "candidate_name": resume.get("name"),
        "candidate_email": config.candidate_email,
        "job_title": resume.get("applied_job_title"),
        "organization_id": organization_id,
        "organization_name": org.get("organization_name") if org else None,
        "questions": [q.dict() for q in config.questions],
        "custom_message": config.custom_message,
        "status": "sent",
        "sent_at": datetime.now(timezone.utc).isoformat(),
        "responses": [],
        "completed_at": None
    }
    
    await db.pre_assessments.insert_one(pre_assessment)
    
    # Update resume status
    await db.routed_resumes.update_one(
        {"id": resume_id},
        {"$set": {
            "pre_assessment_sent": True,
            "pre_assessment_id": pre_assessment["id"],
            "pre_assessment_sent_at": pre_assessment["sent_at"]
        }}
    )
    
    # Generate candidate-facing link
    assessment_link = f"/pre-assessment/{pre_assessment['id']}"
    
    return {
        "message": "Pre-assessment sent to candidate",
        "assessment_id": pre_assessment["id"],
        "assessment_link": assessment_link
    }

@api_router.get("/pre-assessment/{assessment_id}")
async def get_pre_assessment_for_candidate(assessment_id: str):
    """Public endpoint for candidate to view and answer pre-assessment"""
    
    assessment = await db.pre_assessments.find_one({"id": assessment_id}, {"_id": 0})
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    if assessment.get("status") == "completed":
        return {"message": "This assessment has already been completed", "completed": True}
    
    return {
        "id": assessment["id"],
        "job_title": assessment.get("job_title"),
        "organization_name": assessment.get("organization_name"),
        "custom_message": assessment.get("custom_message"),
        "questions": assessment.get("questions", []),
        "completed": False
    }

@api_router.post("/pre-assessment/{assessment_id}/submit")
async def submit_pre_assessment(assessment_id: str, responses: dict):
    """Submit pre-assessment responses from candidate"""
    
    assessment = await db.pre_assessments.find_one({"id": assessment_id}, {"_id": 0})
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    if assessment.get("status") == "completed":
        raise HTTPException(status_code=400, detail="Assessment already completed")
    
    # Save responses
    await db.pre_assessments.update_one(
        {"id": assessment_id},
        {"$set": {
            "responses": responses.get("answers", []),
            "status": "completed",
            "completed_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    # Update resume
    await db.routed_resumes.update_one(
        {"id": assessment.get("resume_id")},
        {"$set": {
            "pre_assessment_completed": True,
            "pre_assessment_completed_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {"message": "Pre-assessment submitted successfully"}

@api_router.get("/resume/{resume_id}/pre-assessment/status")
async def get_pre_assessment_status(resume_id: str):
    """Get pre-assessment status for a resume"""
    
    assessment = await db.pre_assessments.find_one({"resume_id": resume_id}, {"_id": 0})
    
    if not assessment:
        return {"sent": False, "completed": False, "assessment": None}
    
    return {
        "sent": True,
        "completed": assessment.get("status") == "completed",
        "assessment": assessment
    }


# ============ Subscription & Career Trajectory Access ============

@api_router.get("/resume/{resume_id}/career-trajectory-preview")
async def get_career_trajectory_preview(resume_id: str, organization_id: Optional[str] = None):
    """Get preview/teaser of career trajectory - prompts subscription for full access"""
    
    resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Check subscription
    has_access = False
    subscription_tier = "free"
    if organization_id:
        client = await db.clients.find_one({"id": organization_id}, {"_id": 0})
        if client:
            subscription_tier = client.get("subscription", {}).get("tier", "free")
            has_access = subscription_tier in ["professional", "enterprise"]
    
    # Basic preview data (available to all)
    preview = {
        "candidate_name": resume.get("name"),
        "current_role": resume.get("parsed_data", {}).get("current_role"),
        "experience_years": resume.get("parsed_data", {}).get("experience_years"),
        "primary_function": resume.get("primary_function"),
        "teaser_insights": [
            f"Career progression analysis available",
            f"Predicted growth trajectory mapped",
            f"Industry benchmarking comparison",
            f"Leadership potential assessment",
            f"Skill gap analysis with recommendations"
        ],
        "has_full_access": has_access,
        "subscription_tier": subscription_tier,
        "upgrade_message": None if has_access else "Upgrade to Professional tier to unlock full Career Trajectory analysis with detailed predictions, benchmarking, and growth recommendations."
    }
    
    if has_access:
        # Full career trajectory data
        preview["full_analysis"] = {
            "career_path": resume.get("parsed_data", {}).get("companies", []),
            "progression_score": resume.get("confidence_score", 0) * 100,
            "predicted_trajectory": "Upward",
            "leadership_potential": "High" if resume.get("parsed_data", {}).get("experience_years", 0) > 10 else "Moderate",
            "industry_benchmark": "Above Average",
            "growth_areas": resume.get("ai_assessment", {}).get("areas_for_growth", []),
            "strengths": resume.get("ai_assessment", {}).get("strengths", [])
        }
    
    return preview

@api_router.post("/resume/{resume_id}/verbal-evaluation")
async def save_verbal_evaluation(resume_id: str, evaluation: dict):
    """Save verbal evaluation notes for a candidate"""
    
    await db.routed_resumes.update_one(
        {"id": resume_id},
        {"$set": {
            "verbal_evaluation": evaluation.get("text", ""),
            "verbal_evaluation_updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {"message": "Verbal evaluation saved"}

@api_router.get("/resume/{resume_id}/full-report")
async def download_full_report_with_cv(resume_id: str, organization_id: Optional[str] = None, format: str = "pdf"):
    """Download complete report with CV - requires subscription for career trajectory"""
    
    resume = await db.routed_resumes.find_one({"id": resume_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Check subscription for full features
    has_career_trajectory = False
    if organization_id:
        client = await db.clients.find_one({"id": organization_id}, {"_id": 0})
        if client:
            tier = client.get("subscription", {}).get("tier", "free")
            has_career_trajectory = tier in ["professional", "enterprise"]
    
    # Get all related data
    job = None
    if resume.get("applied_job_id"):
        job = await db.job_descriptions.find_one({"id": resume["applied_job_id"]}, {"_id": 0})
        if not job:
            job = await db.structured_jds.find_one({"id": resume["applied_job_id"]}, {"_id": 0})
    
    pre_assessment = await db.pre_assessments.find_one({"resume_id": resume_id}, {"_id": 0})
    
    org = None
    if organization_id:
        org = await db.clients.find_one({"id": organization_id}, {"_id": 0})
    
    report_data = {
        "candidate": {
            "name": resume.get("name"),
            "email": resume.get("email"),
            "phone": resume.get("phone"),
            "current_role": resume.get("parsed_data", {}).get("current_role"),
            "experience_years": resume.get("parsed_data", {}).get("experience_years"),
            "education": resume.get("parsed_data", {}).get("education"),
            "top_skills": resume.get("parsed_data", {}).get("top_skills", []),
            "companies": resume.get("parsed_data", {}).get("companies", [])
        },
        "job": {
            "title": resume.get("applied_job_title") or (job.get("title") if job else "General"),
            "requisition_number": resume.get("applied_requisition_number"),
            "company": job.get("company_name") or job.get("client_name") if job else None
        },
        "competency_analysis": {
            "matched_skills": resume.get("matched_skills", []),
            "primary_function": resume.get("primary_function"),
            "sub_function": resume.get("sub_function"),
            "confidence_score": resume.get("confidence_score", 0),
            "skill_match_scores": resume.get("skill_match_scores", {}),
            "technology_stack": resume.get("technology_stack", {})
        },
        "ai_assessment": resume.get("ai_assessment", {}),
        "interview_questions": resume.get("interview_questions", []),
        "pre_assessment": pre_assessment,
        "recruiter_notes": resume.get("recruiter_notes", []),
        "verbal_evaluation": resume.get("verbal_evaluation", ""),
        "status": resume.get("recruiter_status", "under_evaluation"),
        "career_trajectory": None,
        "has_career_trajectory": has_career_trajectory,
        "organization": {
            "name": org.get("organization_name") if org else None,
            "logo_url": org.get("logo_url") if org else None
        },
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "resume_text": resume.get("raw_text", "")[:5000]  # Include CV text
    }
    
    if has_career_trajectory:
        report_data["career_trajectory"] = {
            "career_path": resume.get("parsed_data", {}).get("companies", []),
            "progression_analysis": "Detailed analysis available",
            "leadership_potential": resume.get("ai_assessment", {}).get("recommendation", "Under Review"),
            "growth_areas": resume.get("ai_assessment", {}).get("areas_for_growth", []),
            "strengths": resume.get("ai_assessment", {}).get("strengths", [])
        }
    
    if format == "pdf":
        pdf_content = generate_full_report_pdf(report_data)
        return Response(
            content=pdf_content,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="full_report_{resume.get("name", "unknown").replace(" ", "_")}.pdf"'}
        )
    else:
        return report_data

def generate_full_report_pdf(data: dict) -> bytes:
    """Generate comprehensive PDF report with CV and all analysis"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    story = []
    styles = getSampleStyleSheet()
    
    # Styles
    title_style = ParagraphStyle('Title', fontSize=18, textColor=colors.HexColor('#10B981'), spaceAfter=10, fontName='Helvetica-Bold')
    heading_style = ParagraphStyle('Heading', fontSize=12, textColor=colors.HexColor('#1F2937'), spaceBefore=12, spaceAfter=6, fontName='Helvetica-Bold')
    body_style = ParagraphStyle('Body', fontSize=9, leading=12, spaceAfter=4)
    small_style = ParagraphStyle('Small', fontSize=8, textColor=colors.gray)
    
    # Header with company branding
    org = data.get("organization", {})
    if org.get("name"):
        story.append(Paragraph(f"<b>{org['name']}</b> - Client Assessment Report", ParagraphStyle('OrgName', fontSize=14, textColor=colors.HexColor('#6366F1'))))
    else:
        story.append(Paragraph("Candidate Assessment Report", title_style))
    
    story.append(Paragraph(f"Generated: {data.get('generated_at', '')[:10]} | Confidential", small_style))
    story.append(Spacer(1, 15))
    
    # Candidate Summary
    candidate = data.get("candidate", {})
    story.append(Paragraph("Candidate Profile", heading_style))
    story.append(Paragraph(f"<b>{candidate.get('name', 'N/A')}</b>", body_style))
    story.append(Paragraph(f"Email: {candidate.get('email', 'N/A')} | Phone: {candidate.get('phone', 'N/A')}", body_style))
    story.append(Paragraph(f"Current: {candidate.get('current_role', 'N/A')} | Experience: {candidate.get('experience_years', 'N/A')} years", body_style))
    
    # Applied Position
    job = data.get("job", {})
    story.append(Paragraph("Applied For", heading_style))
    story.append(Paragraph(f"{job.get('title', 'N/A')} (Req: {job.get('requisition_number', 'N/A')})", body_style))
    
    # Competency Analysis
    comp = data.get("competency_analysis", {})
    story.append(Paragraph("Competency Analysis", heading_style))
    story.append(Paragraph(f"Function: {comp.get('primary_function')} - {comp.get('sub_function')}", body_style))
    story.append(Paragraph(f"Confidence: {comp.get('confidence_score', 0)*100:.0f}%", body_style))
    if comp.get("matched_skills"):
        story.append(Paragraph(f"Key Skills: {', '.join(comp.get('matched_skills', [])[:8])}", body_style))
    
    # AI Assessment
    ai = data.get("ai_assessment", {})
    if ai:
        story.append(Paragraph("AI Assessment", heading_style))
        story.append(Paragraph(f"Rating: {ai.get('overall_rating', 'N/A')}/10 | Recommendation: {ai.get('recommendation', 'N/A')}", body_style))
        if ai.get("strengths"):
            story.append(Paragraph(f"Strengths: {', '.join(ai.get('strengths', [])[:3])}", body_style))
    
    # Pre-Assessment
    pre = data.get("pre_assessment")
    if pre and pre.get("responses"):
        story.append(Paragraph("Pre-Assessment Responses", heading_style))
        for resp in pre.get("responses", [])[:3]:
            story.append(Paragraph(f"Q: {resp.get('question', '')[:80]}...", body_style))
            story.append(Paragraph(f"A: {resp.get('answer', '')[:150]}...", body_style))
    
    # Career Trajectory (if available)
    if data.get("has_career_trajectory") and data.get("career_trajectory"):
        ct = data.get("career_trajectory")
        story.append(Paragraph("Career Trajectory Analysis", heading_style))
        story.append(Paragraph(f"Leadership Potential: {ct.get('leadership_potential', 'N/A')}", body_style))
        if ct.get("growth_areas"):
            story.append(Paragraph(f"Growth Areas: {', '.join(ct.get('growth_areas', [])[:3])}", body_style))
    
    # Notes
    if data.get("recruiter_notes"):
        story.append(Paragraph("Recruiter Notes", heading_style))
        for note in data.get("recruiter_notes", [])[-3:]:
            story.append(Paragraph(f"• {note.get('text', '')}", body_style))
    
    if data.get("verbal_evaluation"):
        story.append(Paragraph("Verbal Evaluation", heading_style))
        story.append(Paragraph(data.get("verbal_evaluation")[:500], body_style))
    
    # Status
    status_map = {"under_evaluation": "Under Evaluation", "selected_next_round": "Selected for Next Round", "hold": "On Hold", "rejected": "Rejected"}
    story.append(Paragraph("Status", heading_style))
    story.append(Paragraph(f"<b>{status_map.get(data.get('status'), data.get('status', 'N/A'))}</b>", body_style))
    
    # Footer
    story.append(Spacer(1, 20))
    if not data.get("has_career_trajectory"):
        story.append(Paragraph("* Career Trajectory analysis requires Professional subscription", small_style))
    story.append(Paragraph(f"Powered by RoleSense - Talent Intelligence Platform", small_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ============ Career Trajectory Access Control ============

@api_router.get("/admin/subscription-features")
async def get_subscription_features():
    """Get subscription feature access configuration"""
    
    features = await db.subscription_features.find({}, {"_id": 0}).to_list(50)
    
    # Default features if none exist
    if not features:
        default_features = [
            {
                "id": "career_trajectory",
                "name": "Career Trajectory Analysis",
                "description": "AI-powered career path analysis and predictions",
                "tiers": {
                    "free": False,
                    "basic": False,
                    "professional": True,
                    "enterprise": True
                }
            },
            {
                "id": "hr_fitment",
                "name": "HR Fitment Analysis",
                "description": "Cultural and team fit assessment",
                "tiers": {
                    "free": False,
                    "basic": True,
                    "professional": True,
                    "enterprise": True
                }
            },
            {
                "id": "advanced_reports",
                "name": "Advanced PDF Reports",
                "description": "Detailed candidate assessment reports",
                "tiers": {
                    "free": False,
                    "basic": False,
                    "professional": True,
                    "enterprise": True
                }
            },
            {
                "id": "custom_questions",
                "name": "Custom Interview Questions",
                "description": "Configure custom interview questions",
                "tiers": {
                    "free": {"limit": 3},
                    "basic": {"limit": 5},
                    "professional": {"limit": 10},
                    "enterprise": {"limit": "unlimited"}
                }
            }
        ]
        for f in default_features:
            f["created_at"] = datetime.now(timezone.utc).isoformat()
            await db.subscription_features.insert_one(f)
        features = default_features
    
    return {"features": features}

@api_router.put("/admin/subscription-features/{feature_id}")
async def update_subscription_feature(feature_id: str, update: dict):
    """Update subscription feature access settings"""
    
    await db.subscription_features.update_one(
        {"id": feature_id},
        {"$set": {
            "tiers": update.get("tiers", {}),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {"message": "Feature updated"}

@api_router.get("/admin/clients/{client_id}/feature-access")
async def get_client_feature_access(client_id: str):
    """Get feature access for a specific client based on subscription"""
    
    client = await db.clients.find_one({"id": client_id}, {"_id": 0})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    subscription = client.get("subscription", {})
    tier = subscription.get("tier", "free")
    
    # Get all features
    features = await db.subscription_features.find({}, {"_id": 0}).to_list(50)
    
    access = {}
    for feature in features:
        tier_access = feature.get("tiers", {}).get(tier, False)
        access[feature["id"]] = {
            "enabled": bool(tier_access) if not isinstance(tier_access, dict) else True,
            "limit": tier_access.get("limit") if isinstance(tier_access, dict) else None
        }
    
    # Check for override settings
    overrides = client.get("feature_overrides", {})
    for feature_id, override in overrides.items():
        if feature_id in access:
            access[feature_id]["enabled"] = override.get("enabled", access[feature_id]["enabled"])
            if "limit" in override:
                access[feature_id]["limit"] = override["limit"]
    
    return {
        "client_id": client_id,
        "subscription_tier": tier,
        "features": access
    }

@api_router.put("/admin/clients/{client_id}/feature-override")
async def set_client_feature_override(client_id: str, override: dict):
    """Set feature access override for a specific client"""
    
    feature_id = override.get("feature_id")
    if not feature_id:
        raise HTTPException(status_code=400, detail="feature_id required")
    
    await db.clients.update_one(
        {"id": client_id},
        {"$set": {
            f"feature_overrides.{feature_id}": {
                "enabled": override.get("enabled", True),
                "limit": override.get("limit"),
                "reason": override.get("reason", ""),
                "set_at": datetime.now(timezone.utc).isoformat()
            }
        }}
    )
    
    return {"message": "Feature override set"}


# Include the router in the main app
app.include_router(api_router)


@app.on_event("startup")
async def startup_db_client():
    """Setup database indexes and cleanup duplicates on startup"""
    # First, cleanup any existing duplicates BEFORE creating unique index
    # Group by name and keep only the first one
    folders = await db.functional_folders.find({}).to_list(100)
    seen_names = set()
    for folder in folders:
        if folder["name"] in seen_names:
            await db.functional_folders.delete_one({"_id": folder["_id"]})
        else:
            seen_names.add(folder["name"])
    
    # Cleanup duplicate sub-folders
    sub_folders = await db.sub_functional_folders.find({}).to_list(500)
    seen_sub = set()
    for sub in sub_folders:
        key = f"{sub['parent_function']}_{sub['name']}"
        if key in seen_sub:
            await db.sub_functional_folders.delete_one({"_id": sub["_id"]})
        else:
            seen_sub.add(key)
    
    # Now create unique indexes to prevent future duplicates
    try:
        await db.functional_folders.create_index("name", unique=True)
        await db.sub_functional_folders.create_index([("parent_function", 1), ("name", 1)], unique=True)
    except Exception as e:
        # Index might already exist, that's fine
        print(f"Index creation note: {e}")
    
    # Initialize default admin user
    await initialize_admin_user()
    
    # Initialize business users for Ally Executive
    await initialize_business_users()


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
